from __future__ import annotations

import io
from pathlib import Path
import sys

import pandas as pd
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))


def main() -> None:
    source = (ROOT / "app" / "Lmstudio_SSAI_chat_main.py").read_text(encoding="utf-8")

    greeting = source[source.index("def _generate_login_greeting"):source.index("def _render_login_greeting_banner")]
    assert "sims_user_name" in greeting
    assert "generated_identity_complete=False" in greeting
    assert "required_identity = (company_display_name, role_text, sims_user_name)" in greeting

    assert "def _tabular_text_for_attachment" in source
    assert 'frame.where(pd.notna(frame), "").to_string()' in source
    assert 'rows = 20 if preview else 200' in source
    assert "extracted = process_file(uf, preview=False)" in source
    assert "MAX_ATTACHMENT_ANALYSIS_CHARS = 20_000" in source

    frame = pd.DataFrame({"품목": ["A", None], "수량": [1, float("nan")]})
    rendered = frame.head(200).copy().astype("object").where(pd.notna(frame), "").to_string()
    assert "NaN" not in rendered and "nan" not in rendered

    doc = Document()
    doc.add_paragraph("DOCX-ATTACHMENT-CONTENT")
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    restored = Document(stream)
    assert restored.paragraphs[0].text == "DOCX-ATTACHMENT-CONTENT"

    requirements = (ROOT / "requirements.txt").read_text(encoding="utf-8")
    lock = (ROOT / "requirements.lock.txt").read_text(encoding="utf-8")
    assert "python-docx==1.1.2" in requirements
    assert "lxml==6.1.2" in requirements
    assert "python-docx==1.1.2" in lock
    assert "lxml==6.1.2" in lock
    print("RESULT OK tests=10")


if __name__ == "__main__":
    main()
