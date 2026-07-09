# 2025/09/18 DB 연결 하면서 모듈 개선
# VERSION = "chat_middleware/2025-11-08T-v3"
# 2026/05/03 업로드 파일명 시그니처로 변경 (크기 포함), 
#           SIMS 결과 영역 상태 초기화 함수 추가, 
#           SIMS 컨텍스트 추출 함수 개선, 
#           LLM 호출 함수 개선 (정규화 + 타임아웃/재시도), 안정 스크롤 JS 개선, 회로 차단기 헬퍼 추가,  
#           DB 진단 함수 추가, 
#           로그 tail 헬퍼 추가, SIMS 컨텍스트 주입용 함수 임포트 추가
# 2025/12/14 SIMS ↔ 채팅 브리지 미들웨어 연결, LM Studio용 메시지 정규화 함수 추가
# 2025/11/08 SIMS 컨텍스트 주입용 함수 임포트 추가, DB 진단 함수 추가, 로그 tail 헬퍼 추가
# 2025/10/30 안정 스크롤 JS 개선, LLM 호출 함수 개선 (정규화 + 타임아웃/재시도), 회로 차단기 헬퍼 추가
# 2025/10/28 SIMS 컨텍스트 추출 함수 개선, SIMS 결과 영역 상태 초기화 함수 추가
# 2025/10/27 업로드 파일명 시그니처로 변경 (크기 포함)
# 2025/10/24 환경변수 추가, LLM 재시도 안정화 (지터 + 명시적 에러)
# 2025/10/23 OCR 옵션 반영, LLM 재시도 안정화 (지터 + 명시적 에러)
# 2025/10/22 코드 최적화 (통합본)
# 2025/09/04 환경변수 , OCR 옵션 반영, LLM 재시도 안정화 (지터 + 명시적 에러),코드 최적화 (통합본)
# 2025/09/08 실전 안정화 패치 7개
# 2025/09/07 환경설정+헤더 정리(단일 소스 오브 트루스)
# 2026/06/07 현재료 관련 후속분석 action dispatcher 모듈 추가

#VERSION = "vendors/2025-12-14-001"
# 웹앱에 첨부 화일 업로드 기능 추가
# Lmstudio_SSAI_chat_main.py
# =========================================================
# Lmstudio_SSAI_chat_main.py  상단 부 정리본
# 순서: (1) 표준 import → (2) .env 로드 → (3) sys.path 설정
#      (4) 환경변수 가드 → (5) DB/UI 모듈 임포트 → (6) 기본 설정
# =========================================================
# =========================
# Standard library
# =========================
import os
import sys
import io
import re
import json
import uuid
import time
import html
import textwrap
import zipfile
import tempfile
import traceback

from pathlib import Path
from datetime import datetime
from typing import Any, Dict, List, Union
from openai import OpenAI 
import random
# (사용 여부 확인 후 유지/제거)
import hashlib

# =========================================================
# import guard: ensure project root on sys.path
# - 이 파일은 app/ 폴더 안에 있으므로, 프로젝트 루트는 parent.parent
# - 어떤 CWD에서 실행해도 "import app.***"가 안전하게 되도록 보장
# =========================================================
_THIS = Path(__file__).resolve()
_ROOT = _THIS.parent.parent  # .../LmStudion_project1
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))

# =========================================================
# 0) .env를 "가장 먼저" 로드 (DB/클라이언트 임포트 이전)
#    - 프로젝트 루트/.env, app/.env 순으로 탐색
# =========================================================
from dotenv import load_dotenv  # type: ignore
load_dotenv(override=True)

import os as _os_for_log_init, sys as _sys_for_log_init, logging as _logging_init
from app.utils.logging_setup import setup_rotating_logger

_level_name = (_os_for_log_init.getenv("LOG_LEVEL") or _os_for_log_init.getenv("SIMS_LOG_LEVEL") or "INFO").upper()
_level = getattr(_logging_init, _level_name, _logging_init.INFO)

# 파일+콘솔 회전 로거( logs/app.log )
_log_file = (
    _os_for_log_init.getenv("LOG_FILE")
    or _os_for_log_init.getenv("SIMS_LOG_FILE")
    or "logs/app.log"
)

log = setup_rotating_logger(
    name="ssai",
    level=_level,
    log_file=_log_file,
)

log.debug("LOG INIT → level=%s (.env loaded)", _level_name)

# (선택) SQLAlchemy 로깅 톤 정리는 초기화 이후에 설정
_logging_init.getLogger("sqlalchemy.engine").setLevel(_logging_init.WARNING)
_logging_init.getLogger("sqlalchemy.pool").setLevel(_logging_init.WARNING)
_logging_init.getLogger("sqlalchemy.dialects").setLevel(_logging_init.WARNING)

# =========================
# Third-party
# =========================
import streamlit.components.v1 as stc
import streamlit as st

# =========================================================
# Streamlit Page Config
# - 반드시 첫 번째 st.* 호출 전에 실행
# - .env는 위에서 이미 load_dotenv(override=True) 완료됨
# =========================================================
sidebar_state = os.getenv("SSAI_INITIAL_SIDEBAR_STATE", "collapsed").strip().lower()

if sidebar_state not in ("auto", "expanded", "collapsed"):
    sidebar_state = "collapsed"

st.set_page_config(
    page_title="SSAI LM Studio Chatbot",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state=sidebar_state,
)

import pandas as pd
from app.sims.nlq.nlq_router import try_handle_nlq

from app.ui.sims_entry import (
    sims_mode_selector,
    render_sims_sidebar_controls,
)
from app.ui.sims_panel import render_sims_main, set_run_flag
from app.ui.sims_hub import render_sims_hub
from app.db.db_config import load_mssql_config
from app.ui.chat_middleware import set_chat_render_anchor,render_sims_chat_item

# =========================
# 무거운 라이브러리는 지연 임포트 권장:
# import PyPDF2               # ← 사용 함수 내부에서 import 권장
# import docx                 # ← 사용 함수 내부에서 import 권장 (python-docx)

# 파일 상단 임포트 섹션
try:
    import pyodbc  # 선택적: 설치되어 있으면 상세 오류 타입 체크에 사용
except Exception:
    pyodbc = None  # 없으면 None으로 두고 가드 처리


# =========================================================
# 1) 패키지 루트(sys.path) 설정
#    - Streamlit 실행 경로 다양성 대응, app.* 임포트 보장
# =========================================================
THIS_DIR = os.path.dirname(os.path.abspath(__file__))   # .../app
PROJECT_ROOT = os.path.dirname(THIS_DIR)                # .../LmStudion_project1
if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)

# =========================================================
# 2) 환경변수 가드 (DB/클라이언트 임포트 이전)
#    - 누락 시 Streamlit에 친절히 안내 후 안전 중단
# =========================================================
from app.utils.config_guard import require_env  # app/utils/__init__.py 필요(빈 파일 OK)

try:
    require_env()  # 누락 시 SystemExit(2)
except SystemExit:
    st.error("필수 환경변수가 누락되었습니다. `.env.sample`을 참고하여 `.env`를 채워주세요.")
    st.stop()
except Exception as e:
    st.error(f"환경 검증 중 오류: {e}")
    st.stop()
# =========================================================
# 2-1) LLM 상태 체크 (임포트 전에)
# =========================================================
from app.services.llm_health import check_llm

LLM_STATUS = None
try:
    LLM_STATUS = check_llm()
except Exception as _e:
    LLM_STATUS = {"ok": False, "error": str(_e)}
# =========================================================
# 2-2) 로깅 설정 (DB/클라이언트 임포트 이전)
# =========================================================
# (정리) 로깅 초기화는 최상단에서 완료됨 → 여기서는 부팅 로그만 남김
if not st.session_state.get("__boot_logged"):
    log.info("==== 앱 시작 ====")
    st.session_state["__boot_logged"] = True
# =========================================================
# 3) DB 모듈 임포트 (정식 경로로 고정)
# =========================================================
try:
    from app.db.mssql_client import read_df, fetch_one, fetch_all, health_check
except Exception:
    log.exception("DB 모듈 임포트 실패(app.db.mssql_client)")
    st.stop()

# =========================================================
# 4) Sims ERP UI 모듈 임포트
# =========================================================
#from app.ui.sims_entry import (
#    sims_mode_selector,
#    render_sims_sidebar_controls,
#    render_sims_main,
#)
from app.ui.chat_middleware import wire_chat_context,render_sims_chat_item
from app.ui.chat_bridge import get_sims_context_text, get_sims_context_data  # ✅ SIMS 컨텍스트 주입용
from app.db.mssql_client import health_check


from app.ui.current_table_followups.action_dispatcher import handle_current_table_followup_by_action

from app.services.ssai_permission_policy import (
    describe_permission,
    get_required_permission,
)

from app.ui.ssai_login import (
    require_login,
    require_permission,
    get_current_user,
    get_selected_company,
)

from app.services.ssai_storage_service import (
    get_user_area_dir,
    make_safe_filename,
)

from app.services.ssai_audit_service import safe_log_audit_event

from app.ui.ssai_admin import (
    render_ssai_admin_page,
    render_ssai_admin_sidebar,
)

# =========================================================
# 4-0) 전역 CSS: SIMS popover / expander UI 스타일 보정
# =========================================================
def _inject_base_css_once() -> None:
    """Streamlit 전역 CSS를 1회만 주입한다.

    st.markdown(..., unsafe_allow_html=True) 방식은 일부 HTTPS/IIS 경유
    초기 렌더링에서 <style>...</style> 문자열이 화면에 노출되는 경우가 있어
    st.html() 우선, 미지원 버전은 st.markdown()으로 fallback 한다.
    """
    if st.session_state.get("__base_css_loaded"):
        return

    css = """
<style>
/* SIMS 팝오버 본문 최소 폭 확장 */
[data-testid="stPopover"] [data-testid="stPopoverBody"] {
    min-width: 640px;
}

/* expander 제목 줄 간격 보정 */
details[open] > summary {
    line-height: 1.1;
}

/* 패널 내부 여백/버튼 정렬 보정 */
[data-testid="stPopoverBody"] > div {
    padding: 0.75rem 1rem !important;
}
</style>
"""

    if hasattr(st, "html"):
        st.html(css)
    else:
        st.markdown(css, unsafe_allow_html=True)

    st.session_state["__base_css_loaded"] = True


_inject_base_css_once()

#
# =========================================================
# 4-0-1) 사이드바 환경/진단/디버그 도구 표시 여부.
# =========================================================

def _can_show_admin_diagnostics_sidebar() -> bool:
    """
    사이드바 환경/진단/디버그 도구 표시 여부.

    운영 사용자에게는 내부 환경정보, Health Check, Debug 메뉴를 노출하지 않는다.
    표시 대상:
    - SSART_ADMIN 만 표시

    숨김 대상:
    - SSART_USER
    - WHOLESALE_ADMIN
    - WHOLESALE_USER
    - 기타 일반 사용자
    """
    user = get_current_user()

    if not user:
        return False

    user_type = str(getattr(user, "user_type", "") or "").strip().upper()

    return user_type == "SSART_ADMIN"
#
# =========================================================
# 4-1) SIMS ↔ 채팅 브리지 미들웨어 연결
# =========================================================
# SIMS 실행 결과를 채팅으로 드레인하고, SIMS 컨텍스트를 system 메시지로 주입하며,
# '닫기' 입력 시 컨텍스트를 정리합니다.
# (wire_chat_context는 메인 렌더(prepass)에서 1회 호출됩니다.)

# =========================================================
# 5) Config / Client / Paths (single source of truth)
# =========================================================
APP_DIR = Path(__file__).parent
ENV_PATH = APP_DIR / ".env"  # 실제 운영은 프로젝트 루트 .env 권장

_DEFAULT_ENV_TEXT = textwrap.dedent("""\
# (옵션) LLM 호출 보호막
# LLM_TIMEOUT_S=25
# LLM_MAX_RETRY=2


# LM Studio API 연결
LMSTUDIO_BASE_URL=http://localhost:1234/v1
LMSTUDIO_API_KEY=lm-studio

# 파일/채팅 저장
CHAT_FILE=data/chat_rooms.json
UPLOAD_DIR=uploads

# 업로드 제한
MAX_FILE_SIZE_MB=25
MAX_PREVIEW_CHARS=4000
""")

# .env 파일이 없으면 생성 (권한 이슈 등은 조용히 무시)
def _ensure_dotenv_exists():
    try:
        if not ENV_PATH.exists():
            ENV_PATH.parent.mkdir(parents=True, exist_ok=True)
            ENV_PATH.write_text(_DEFAULT_ENV_TEXT, encoding="utf-8")
    except Exception:
        pass  # 권한 이슈 등은 조용히 무시

def _load_dotenv_if_possible():
    if load_dotenv is not None:
        try:
            load_dotenv(dotenv_path=ENV_PATH, override=False)  # 기존 env 보존
        except Exception:
            pass

_ensure_dotenv_exists()
_load_dotenv_if_possible()

# 안전한 설정 getter: OS env > st.secrets > default
def get_config(key: str, default=None, cast=str):
    v = os.environ.get(key)
    if v not in (None, ""):
        try:
            return cast(v) if cast is not str else v
        except Exception:
            return default
    try:
        v = st.secrets[key]  # secrets.toml이 없을 수 있으므로 try/except
        if v not in (None, ""):
            try:
                return cast(v) if cast is not str else v
            except Exception:
                return default
    except Exception:
        pass
    return default

# ===== 첨부 분석 작업(job) 상태 =====
st.session_state.setdefault("__an_busy", False)      # 현재 분석 작업 실행 중
st.session_state.setdefault("__an_job", None)        # {"id","room_id","sig","started_at"}
st.session_state.setdefault("__an_cancel", False)    # 사용자가 취소 버튼 클릭
st.session_state.setdefault("__sims_was_final", False)

def _sig_of_uploads(files):
    """업로드 목록의 변경 여부를 빠르게 판단하기 위한 시그니처"""
    parts = []
    for f in files or []:
        try:
            size = getattr(f, "size", None) or len(f.getvalue())
        except Exception:
            size = 0
        parts.append(f"{getattr(f,'name','unnamed')}:{size}")
    return "|".join(parts)

def _get_room_by_id(rid: str):
    for r in st.session_state.get("chat_rooms", []):
        if r.get("id") == rid:
            return r
    return None

def _safe_rerun():
    """Streamlit 버전별 rerun 호환 호출"""
    import streamlit as st
    if hasattr(st, "rerun"):
        st.rerun()
    else:
        # 구버전 호환 (experimental_rerun)
        if hasattr(st, "experimental_rerun"):
            st.experimental_rerun()
        else:
            # 최후의 안전장치: rerun 미지원이면 아무것도 하지 않음
            return

def _reset_sims_result_area_for_selection_change(prev_selected: dict | None, new_selected: dict | None) -> bool:
    prev_selected = prev_selected or {}
    new_selected = new_selected or {}

    if prev_selected == new_selected:
        return False

    # SIMS 결과 영역/컨텍스트만 비운다. 채팅 이력은 건드리지 않음.
    for k in [
        "__sims_context",
        "__sims_context_text",
        "__sims_context_note",
        "__sims_result",
        "__sims_last_push_sig",
        "__sims_last_push",
        "__sims_flash",
        "__sims_flash_close",
        "__sims_flash_csv",
        "__sims_flash_xlsx",
        "__sims_last_render_run_seq",
    ]:
        st.session_state.pop(k, None)

    # 결과 영역 상태 초기화
    st.session_state["__sims_open"] = False
    st.session_state["__sims_open_ui"] = False
    st.session_state["__sims_force_open"] = False
    st.session_state["__sims_panel_active"] = False
    st.session_state["__sims_rendered"] = False
    st.session_state["__sims_was_final"] = False
    st.session_state["__sims_run_flag"] = False
    st.session_state["__sims_inner_submit"] = False

    return True



def _sims_payload_company_sig(payload: dict | None) -> tuple[str, str]:
    """
    SIMS 패널 payload가 생성된 회사 식별자 추출.

    - 새 패치 이후 payload/meta에는 _ssai_company_id, _ssai_db_name을 기록한다.
    - 기존 payload와의 호환을 위해 값이 없으면 빈 문자열을 반환한다.
    """
    if not isinstance(payload, dict):
        return "", ""

    try:
        meta = payload.get("meta") if isinstance(payload.get("meta"), dict) else {}

        company_id = str(
            meta.get("_ssai_company_id")
            or payload.get("_ssai_company_id")
            or ""
        ).strip()

        db_name = str(
            meta.get("_ssai_db_name")
            or payload.get("_ssai_db_name")
            or ""
        ).strip()

        return company_id, db_name
    except Exception:
        return "", ""


def _sims_current_company_sig() -> tuple[str, str]:
    """현재 선택 회사의 company_id/db_name 서명."""
    try:
        company = get_selected_company()
    except Exception:
        company = None

    if not isinstance(company, dict):
        return "", ""

    return (
        str(company.get("company_id") or "").strip(),
        str(company.get("db_name") or "").strip(),
    )


def _sims_payload_matches_current_company(payload: dict | None) -> bool:
    """
    패널 payload가 현재 선택 회사와 일치하는지 검사한다.

    회사 서명이 없는 legacy payload는 여기서 막지 않는다.
    대신 회사 변경 clear에서 legacy payload/cache를 제거한다.
    """
    payload_company_id, payload_db_name = _sims_payload_company_sig(payload)
    if not payload_company_id and not payload_db_name:
        return True

    current_company_id, current_db_name = _sims_current_company_sig()

    if payload_company_id and current_company_id and payload_company_id != current_company_id:
        return False

    if payload_db_name and current_db_name and payload_db_name != current_db_name:
        return False

    return True


def _clear_sims_runtime_for_company_change(reason: str = "company_change") -> None:
    """
    회사 DB가 바뀔 때 이전 회사 기준의 현재표/컨텍스트/다운로드 캐시를 제거한다.

    핵심 목적:
    - company_id=4에서 조회한 현재표가 company_id=1로 전환된 직후 다시 stash되는 문제 차단
    - 채팅 이력 자체는 보존하되, volatile DataFrame 캐시와 current-table source만 제거
    - SIMS 패널은 열린 상태를 유지할 수 있지만, 자동 재렌더/재조회는 막고 사용자가 다시 실행하게 한다.
    """
    try:
        ss = st.session_state

        prev_source_key = str(ss.get("__sims_current_table_source_key") or "").strip()
        prev_last_key = str(ss.get("__sims_last_table_key") or "").strip()
        remove_table_keys = {k for k in (prev_source_key, prev_last_key) if k}

        # 1) 현재표/컨텍스트/마지막 결과/다운로드 관련 volatile key 제거
        explicit_keys = [
            "__sims_result",
            "__sims_context",
            "__sims_context_text",
            "__sims_context_note",
            "__sims_context_obj",
            "__sims_ctx",
            "__sims_ctx_hash",
            "__sims_analysis_ctx",
            "__sims_latest_analysis_key",
            "__sims_current_table_source_key",
            "__sims_current_table_source_action",
            "__sims_current_table_source_analysis_ctx",
            "__sims_last_table_key",
            "__sims_last_table_action",
            "__sims_last_push_sig",
            "__sims_last_push",
            "__sims_flash",
            "__sims_flash_close",
            "__sims_flash_csv",
            "__sims_flash_xlsx",
            "__sims_last_render_run_seq",
            "__sims_selected_snapshot",
            "__sims_panel_last_final_action",
            "__sims_panel_last_final_payload",
            "__sims_last_final_payload_for_chat",
            "__sims_last_final_payload_for_chat_action",
            "__sims_panel_source_promoted_sig",
            "__sims_panel_chat_push_sig",
            "__sims_form_submitted",
            "__sims_submitted_form_id",
            "__sims_widget_ns",
            "__sims_last_action",
            "__chat_rendered_ids_this_run",
        ]
        for key in explicit_keys:
            ss.pop(key, None)

        # 1-1) 액션별 view가 보관한 마지막 payload/widget 상태도 회사 DB 의존성이 있으므로 제거한다.
        #      예: 사용자목록 24건(company_id=1)이 company_id=4 전환 직후 다시 stash되는 문제 차단.
        view_prefixes = (
            "__users_",
            "__codes_",
            "__ven_",
            "__vendors_",
            "__road_addr_",
            "__analytics_",
            "__goods_",
            "__io",
            "__product_flow",
            "__rddbc_",
            "__m040_",
        )
        for key in list(ss.keys()):
            try:
                if str(key).startswith(view_prefixes):
                    ss.pop(key, None)
            except Exception:
                pass

        # 2) 회사 전환 시 현재표/다운로드용 DF 캐시는 전부 volatile로 보고 제거한다.
        #    채팅 이력은 보존되지만, 새 회사에서 이전 회사 DataFrame을 current source로 재사용하지 않는다.
        for store_key in (
            "sims_tables",
            "sims_export_tables",
            "__sims_export_tables_by_key",
            "__sims_export_tables",
        ):
            store = ss.get(store_key)
            if isinstance(store, dict):
                try:
                    remove_table_keys.update(str(k) for k in store.keys())
                    store.clear()
                except Exception:
                    for table_key in list(remove_table_keys):
                        store.pop(table_key, None)

        # 3) 패널 자동 재렌더 차단. 사용자가 새 회사 기준으로 다시 [실행]해야 한다.
        ss["__sims_panel_active"] = False
        ss["__sims_rendered"] = False
        ss["__sims_was_final"] = False
        ss["__sims_run_flag"] = False
        ss["__sims_inner_submit"] = False
        ss["__sims_force_open"] = False
        ss["__sims_ctx_dirty"] = True

        log.info(
            "[company.change.clear_sims] reason=%s removed_table_keys=%s open=%s selected=%r",
            reason,
            sorted(remove_table_keys),
            ss.get("__sims_open"),
            ss.get("__sims_selected"),
        )
    except Exception:
        log.exception("[company.change.clear_sims] failed reason=%s", reason)

def _check_sims_action_permission(selected: dict | None) -> bool:
    """
    SIMS Panel/Hub 실행 전 권한 확인.

    selected 예:
    {
        "category": "사용자",
        "action": "사용자목록 + 부서명"
    }
    """
    selected = selected or {}

    category = str(selected.get("category") or "").strip()
    action = str(selected.get("action") or "").strip()

    permission = get_required_permission(
        category=category,
        action=action,
    )

    if not permission:
        return True

    if require_permission(permission, show_error=False):
        return True

    st.error(
        f"이 SIMS 작업을 실행할 권한이 없습니다.\n\n"
        f"- 카테고리: {category or '-'}\n"
        f"- 작업: {action or '-'}\n"
        f"- 필요 권한: {permission} ({describe_permission(permission)})"
    )

    log.warning(
        "[auth.permission] blocked SIMS action category=%r action=%r permission=%s",
        category,
        action,
        permission,
    )

    return False

def _upload_unavailable_message() -> str:
    return "파일 첨부 기능은 현재 계정에서 사용할 수 없습니다. 필요하시면 관리자에게 요청해 주세요."


def _upload_unavailable_help() -> str:
    return "관리자 승인 후 파일 첨부와 분석 기능을 사용할 수 있습니다."

# =========================================================
# Debug 용
# =========================================================
def _safe_tail_log(log_path: Path, lines: int = 100) -> str:
    try:
        if not log_path.exists():
            return f"[INFO] 로그 파일이 없습니다: {log_path}"
        with log_path.open("r", encoding="utf-8", errors="ignore") as f:
            buf = f.readlines()
        buf = buf[-lines:]
        return "".join(buf).rstrip("\n")
    except Exception as e:
        return f"[ERROR] 로그 tail 중 오류 발생: {type(e).__name__}: {e}"

def _diagnose_db() -> dict:
    """
    MSSQL 환경/접속 진단.
    - 설정(info)는 load_mssql_config()에서 가져오고
    - 실제 연결 OK/FAIL 은 health_check() 결과에 100% 맞춘다.
    """
    try:
        cfg = load_mssql_config()
    except Exception as e:
        return {
            "ok": False,
            "status": "FAIL",
            "reason": "load_mssql_config 실패",
            "info": None,
            "detail": f"{type(e).__name__}: {e}",
        }

    info = {
        "host": cfg.host,
        "port": cfg.port,
        "db": cfg.database,
        "user": cfg.user,
        "encrypt": cfg.encrypt,
        "trust_cert": cfg.trust_cert,
        "timeout": cfg.timeout,
        # 있다면
        "odbc_driver": getattr(cfg, "odbc_driver", None),
    }

    result = {
        "ok": False,
        "status": "UNKNOWN",
        "reason": "",
        "info": info,
        "detail": "",
    }

    try:
        ok = health_check()
    except Exception as e:
        result["status"] = "FAIL"
        result["reason"] = "health_check() 예외 발생"
        result["detail"] = f"{type(e).__name__}: {e}"
        return result

    if ok:
        result["ok"] = True
        result["status"] = "OK"
        result["reason"] = "DB OK (health_check() → True)"
    else:
        result["ok"] = False
        result["status"] = "FAIL"
        result["reason"] = "DB FAIL (health_check() → False)"

    return result

def _extract_recent_sims_context() -> dict | None:
    """
    session_state 내에서 최근 SIMS 컨텍스트 후보를 찾아 요약용 구조로 리턴.
    - 키 이름이 바뀌어도 'sims' + 'ctx' 또는 'context' 를 포함하면 잡도록 완화.
    """
    ss = st.session_state

    # 1순위: 명시적인 키들
    for key in ("__sims_context", "sims_context", "SIMS_CONTEXT"):
        if key in ss:
            return {"key": key, "value": ss.get(key)}

    # 2순위: 패턴 검색
    candidates = []
    for k in ss.keys():
        lower = str(k).lower()
        if "sims" in lower and ("ctx" in lower or "context" in lower):
            candidates.append(k)

    if not candidates:
        return None

    key = candidates[-1]  # 가장 최근에 세팅된 것으로 가정
    return {"key": key, "value": ss.get(key)}

# 캐스터 헬퍼
cfg_str  = lambda k, d="":  get_config(k, d, cast=str)
cfg_int  = lambda k, d=0:   get_config(k, d, cast=int)
cfg_bool = lambda k, d=False: get_config(k, d, cast=lambda x: str(x).strip().lower() in ("1","true","yes","y","on"))


# --- LLM 호출 보호막 기본값 ---
LLM_TIMEOUT_S   = cfg_int("LLM_TIMEOUT_S", 90)   # 1회 요청 타임아웃(초) - Parallel 1 다중 사용자 대기 고려
LLM_MAX_RETRY   = cfg_int("LLM_MAX_RETRY", 1)    # 재시도 횟수(예: 1이면 총 2번 시도)
LLM_BACKOFF_SEQ = [0.6, 1.2, 2.0]                # 재시도 간 대기(초)

# --- 회로 차단기 기본값 ---
CB_FAIL_THRESHOLD = cfg_int("CB_FAIL_THRESHOLD", 3)   # 연속 실패 N회면 open
CB_OPEN_SECONDS   = cfg_int("CB_OPEN_SECONDS", 45)    # open 유지 시간(초)

# =========================
# 회로 차단기 (Circuit Breaker) 헬퍼
# =========================
def _cb_init():
    ss = st.session_state
    ss.setdefault("__cb_state", "closed")     # closed | open | half_open
    ss.setdefault("__cb_fail_count", 0)
    ss.setdefault("__cb_open_until", 0.0)

def _cb_now() -> float:
    return time.time()

def _cb_is_open() -> bool:
    _cb_init()
    return st.session_state["__cb_state"] == "open" and _cb_now() < st.session_state["__cb_open_until"]

def _cb_trip():
    _cb_init()
    st.session_state["__cb_state"] = "open"
    st.session_state["__cb_open_until"] = _cb_now() + CB_OPEN_SECONDS

def _cb_reset():
    _cb_init()
    st.session_state["__cb_state"] = "closed"
    st.session_state["__cb_fail_count"] = 0
    st.session_state["__cb_open_until"] = 0.0

def _cb_on_success(): _cb_reset()

def _cb_on_failure():
    _cb_init()
    # half_open에서 실패 → 바로 trip(open)
    if st.session_state["__cb_state"] == "half_open":
        _cb_trip()
        return

    st.session_state["__cb_fail_count"] += 1
    if st.session_state["__cb_fail_count"] >= CB_FAIL_THRESHOLD:
        _cb_trip()


# =========================
# 검색 헬퍼
# =========================
def _mk_search_pattern(q: str, *, regex: bool, case_sensitive: bool):
    if not q:
        return None
    flags = 0 if case_sensitive else re.IGNORECASE
    return re.compile(q if regex else re.escape(q), flags)

def _make_snippet(text: str, m: re.Match, ctx: int = 28) -> str:
    """매칭 앞뒤로 문맥을 잘라 <mark>로 하이라이트한 HTML 스니펫 반환"""
    s, e = m.start(), m.end()
    left  = max(0, s - ctx)
    right = min(len(text), e + ctx)
    head = "…" if left > 0 else ""
    tail = "…" if right < len(text) else ""
    esc = html.escape
    return f"{head}{esc(text[left:s])}<mark>{esc(text[s:e])}</mark>{esc(text[e:right])}{tail}"

def search_messages_in_room(room: dict, pattern: re.Pattern, roles: set[str]):
    """room['messages']에서 role 필터와 정규식 패턴으로 검색 결과 리스트 생성"""
    results = []
    for i, m in enumerate(room.get("messages", [])):
        role = m.get("role", "assistant")
        if role not in roles:
            continue
        content = (m.get("content") or "")
        mt = pattern.search(content)
        if not mt:
            continue
        results.append({
            "idx": i,
            "role": role,
            "time": m.get("time", ""),
            "snippet_html": _make_snippet(content, mt),
            "content": content,
        })
    return results
def _queue_search_reset():
    st.session_state["__search_reset"] = True
    st.rerun()

# =========================================================
# --- LM Studio용 메시지 정규화 ---
def _normalize_for_lmstudio_1(raw_messages: list[dict]) -> list[dict]:
    """
    LM Studio jinja 템플릿이 요구하는 형식으로 messages를 정리한다.
    - 최종적으로는 user / assistant / user / assistant ... 만 남기고
      system 역할은 첫 user 메시지 내용에 녹여 넣는다.
    - 선두는 반드시 user 로 시작
    - 연속된 동일 role 은 하나로 병합
    - 마지막도 user 로 끝나도록 보정 (질문형 호출 기준)
    """
    # 1) system 들의 내용을 한 덩어리로 합치기
    system_parts: list[str] = []
    for m in raw_messages:
        if m.get("role") == "system":
            c = str(m.get("content", "")).strip()
            if c:
                system_parts.append(c)
    system_block = "\n\n".join(system_parts) if system_parts else ""

    # 2) user / assistant 메시지만 추출
    dialog: list[dict] = []
    for m in raw_messages:
        role = m.get("role")
        if role not in ("user", "assistant"):
            continue
        content = str(m.get("content", "")).strip()
        if not content:
            continue
        dialog.append({"role": role, "content": content})

    # 3) 앞쪽에 assistant 가 나오면 user 가 나올 때까지 버림
    while dialog and dialog[0]["role"] != "user":
        dialog.pop(0)

    # 4) dialog 가 비어 있고 system_block 만 있으면, system 내용을 user 로 바꿔서 1개 생성
    if not dialog:
        if system_block:
            return [{"role": "user", "content": system_block}]
        return []

    # 5) 첫 user 메시지에 system 내용을 합쳐 넣기
    if system_block:
        dialog[0]["content"] = (system_block.rstrip() + "\n\n" + dialog[0]["content"]).strip()

    # 6) user / assistant 번갈이 + 연속 role 병합
    fixed: list[dict] = []
    for m in dialog:
        if not fixed:
            fixed.append(m)
            continue
        if fixed[-1]["role"] == m["role"]:
            fixed[-1]["content"] = (fixed[-1]["content"].rstrip() + "\n\n" + m["content"]).strip()
        else:
            fixed.append(m)

    # 7) 마지막이 assistant 로 끝나면 user 가 나올 때까지 뒤에서 제거
    while fixed and fixed[-1]["role"] != "user":
        fixed.pop()

    return fixed

# =========================================================
# --- LM Studio용 메시지 정규화 (채팅 완료 직전용) ---
def _normalize_for_lmstudio(messages: list[dict]) -> list[dict]:
    """
    LM Studio jinja 템플릿이 요구하는 형식으로 messages 를 정리한다.
    - system role 제거 (내용은 첫 user 메세지 앞에 붙임)
    - user / assistant 만 남김
    - 선두는 user 로 시작
    - user / assistant 번갈이
    - 마지막도 user 로 끝나게 조정
    """
    # 1) system 모으기 + user/assistant 분리
    system_chunks: list[str] = []
    dialog: list[dict] = []
    for m in messages or []:
        role = m.get("role")
        content = str(m.get("content", "")).strip()
        if not content:
            continue
        if role == "system":
            system_chunks.append(content)
        elif role in ("user", "assistant"):
            dialog.append({"role": role, "content": content})

    system_text = "\n\n".join(system_chunks).strip() or None

    # 2) 앞쪽이 assistant 이면 user 가 나올 때까지 버림
    while dialog and dialog[0]["role"] != "user":
        dialog.pop(0)

    # 3) dialog 비었으면 system 만 있는 상황 → user 하나로 만들어줌
    if not dialog:
        if system_text:
            return [{"role": "user", "content": system_text}]
        return []

    # 4) system_text 가 있으면 첫 user 앞에 붙여줌
    if system_text and dialog[0]["role"] == "user":
        dialog[0]["content"] = (system_text + "\n\n" + dialog[0]["content"]).strip()

    # 5) user/assistant 번갈이 + 같은 role 연속이면 병합
    fixed: list[dict] = []
    for m in dialog:
        if not fixed:
            fixed.append(m)
            continue
        if fixed[-1]["role"] == m["role"]:
            fixed[-1]["content"] = (fixed[-1]["content"].rstrip() + "\n\n" + m["content"]).strip()
        else:
            fixed.append(m)

    # 6) 마지막이 assistant 로 끝나면 user 가 나올 때까지 뒤에서 제거
    while fixed and fixed[-1]["role"] != "user":
        fixed.pop()

    return fixed

# =========================================================
# --- LLM 호출: 타임아웃 + 재시도 (회로차단 없음) ---
# =========================================================
def call_chat_with_retry(
    *,
    messages,
    model=None,
    temperature: float = 0.2,
    stream: bool = False,
    timeout_s: int | None = None,         # per-call 타임아웃(선택)
    max_retry: int | None = None,         # per-call 재시도 횟수(선택, 성공 포함 X)
    backoff_seq: list[float] | None = None,  # per-call 백오프 시퀀스(선택)
):
    """
    - stream=True  -> 스트리밍 generator 반환 (create 호출 성공 시에만 반환)
    - stream=False -> 완료 응답 객체 반환
    """
    # 0) LM Studio 템플릿용 role 정규화
    try:
        fixed_messages = _normalize_for_lmstudio(messages or [])
    except Exception:
        # 문제가 생기면 원본이라도 사용 (최후 안전장치)
        fixed_messages = messages or []

    # 전역 기본값과 per-call 오버라이드 병합
    eff_timeout = int(timeout_s if timeout_s is not None else LLM_TIMEOUT_S)
    eff_retry   = int(max_retry if max_retry is not None else LLM_MAX_RETRY)
    eff_backoff = list(backoff_seq) if backoff_seq is not None else list(LLM_BACKOFF_SEQ or [0.6, 1.2, 2.0])

    # per-call 타임아웃 부여 (드라이버가 지원하면 with_options 사용)
    cli = getattr(CLIENT, "with_options", lambda **kw: CLIENT)(timeout=eff_timeout)

    model_id = model or (st.session_state.get("selected_model") or cfg_str("LMSTUDIO_MODEL", "local-model"))

    try:
        prompt_chars = sum(len(str(m.get("content") or "")) for m in fixed_messages if isinstance(m, dict))
    except Exception:
        prompt_chars = 0

    log.info(
        "[llm.request] start model=%s stream=%s timeout_s=%s max_retry=%s messages=%s prompt_chars=%s",
        model_id,
        stream,
        eff_timeout,
        eff_retry,
        len(fixed_messages or []),
        prompt_chars,
    )

    last_err = None
    total_tries = eff_retry + 1  # 최초 1회 + 재시도 수
    total_t0 = time.perf_counter()

    for attempt in range(total_tries):
        attempt_t0 = time.perf_counter()
        try:
            resp = cli.chat.completions.create(
                model=model_id,
                messages=fixed_messages,
                temperature=temperature,
                stream=stream,
                # 주의: LM Studio 계열은 request_timeout 전달 비권장. with_options(timeout=...)만 사용.
            )

            elapsed_ms = int((time.perf_counter() - attempt_t0) * 1000)
            total_ms = int((time.perf_counter() - total_t0) * 1000)
            log.info(
                "[llm.request] create_ok model=%s stream=%s attempt=%s/%s elapsed_ms=%s total_ms=%s",
                model_id,
                stream,
                attempt + 1,
                total_tries,
                elapsed_ms,
                total_ms,
            )
            return resp

        except Exception as e:
            last_err = e
            elapsed_ms = int((time.perf_counter() - attempt_t0) * 1000)

            if attempt >= total_tries - 1:
                break

            sleep_s = eff_backoff[attempt] if attempt < len(eff_backoff) else eff_backoff[-1]
            log.warning(
                "[llm.request] retry model=%s attempt=%s/%s elapsed_ms=%s sleep_s=%.2f error=%s: %s",
                model_id,
                attempt + 1,
                total_tries,
                elapsed_ms,
                sleep_s,
                type(e).__name__,
                e,
            )
            time.sleep(sleep_s + random.uniform(0, 0.3))

    total_ms = int((time.perf_counter() - total_t0) * 1000)
    log.warning(
        "[llm.request] failed model=%s stream=%s tries=%s total_ms=%s error=%s: %s",
        model_id,
        stream,
        total_tries,
        total_ms,
        type(last_err).__name__,
        last_err,
    )
    raise RuntimeError(f"LLM 호출 실패(최종): {type(last_err).__name__}: {last_err}") from last_err

# =========================================================
# --- 안정 스크롤 헬퍼: DOM 붙을 때까지 기다리며 여러 번 시도 ---
# =========================================================
def _scroll_to_anchor_js(anchor_id: str, *, center: bool = True) -> None:
    try:
        import json as _json
    except Exception:
        _json = None

    block = "center" if center else "start"
    # anchor_id에 특수문자 있을 수 있으니 JS에 안전하게 박아 넣기
    js_id   = _json.dumps(anchor_id) if _json else f'"{anchor_id}"'
    js_block= _json.dumps(block) if _json else f'"{block}"'

    stc.html(f"""
    <script>
    (function(){{
      const ID = {js_id};
      const block = {js_block};

      // 중복 실행 방지 락
      if (window.__jumpLock === ID) return;
      window.__jumpLock = ID;

      function cssEscapeSafe(s){{
        try {{ return CSS.escape(s); }} catch(e) {{
          return String(s).replace(/[^a-zA-Z0-9_-]/g, '\\\\$&');
        }}
      }}
      function getDoc(){{
        try {{
          if (window.parent && window.parent.document) return window.parent.document;
        }} catch (e) {{}}
        return document;
      }}
      const doc = getDoc();

      function scroll(el){{
        try {{
          el.scrollIntoView({{ behavior: "smooth", block }});
        }} catch (e) {{
          el.scrollIntoView();
        }}
      }}

      function tryOnce(){{
        const el = doc.getElementById(ID) || doc.querySelector("#" + cssEscapeSafe(ID));
        if (el) {{ scroll(el); return true; }}
        return false;
      }}

      // 즉시/지연/RAF/Observer까지 총동원
      if (tryOnce()) return;

      let tries = 0, max = 120; // 약 2초(60fps 가정)
      function rafTick(){{
        if (tryOnce()) return;
        if (tries++ < max) requestAnimationFrame(rafTick);
      }}
      requestAnimationFrame(rafTick);

      setTimeout(tryOnce, 100);
      setTimeout(tryOnce, 300);
      setTimeout(tryOnce, 700);
      setTimeout(tryOnce, 1200);

      const obs = new MutationObserver(() => {{ if (tryOnce()) obs.disconnect(); }});
      obs.observe(doc.body, {{ childList: true, subtree: true }});
      setTimeout(() => obs.disconnect(), 2500);
    }})();
    </script>
    """, height=0)
## ===================================================
## CSV/엑셀 파일명에 액션명 반영
## =====================================================
#def _slug_name(s: str) -> str:
#    # 파일명 안전 슬러그 (한글/영문/숫자/언더스코어/하이픈만 남김)
#    s = re.sub(r"[^0-9A-Za-z가-힣_\-]+", "_", str(s or "")).strip("_")
#    return s[:60] or "sims_result"

## 액션명 후보: 플래시가 갖고 있으면 최우선 → 사이드바(패널/허브) 선택값 → 제목 → 기본값
#_action_guess = (
#    (_flash.get("action") if isinstance(_flash, dict) else None)
#    or st.session_state.get("__sims_selected_action")
#    or st.session_state.get("__sims_action")
#    or st.session_state.get("__sims_hub_action")
#    or _flash.get("title")
#    or "SIMS_결과"
#)
#base = _slug_name(_action_guess)
#stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
#csv_name  = f"{base}_{stamp}.csv"
#xlsx_name = f"{base}_{stamp}.xlsx"

# =========================================================
# 값 로드
# =========================================================
CHAT_FILE         = cfg_str("CHAT_FILE", "data/chat_rooms.json")
UPLOAD_DIR        = cfg_str("UPLOAD_DIR", "uploads")
MAX_FILE_SIZE_MB  = cfg_int("MAX_FILE_SIZE_MB", 25)
MAX_PREVIEW_CHARS = cfg_int("MAX_PREVIEW_CHARS", 4000)

# ✅ 타임아웃(초) 통일: LLM_TIMEOUT_S가 우선, 없으면 OPENAI_TIMEOUT(백워드 호환)
#LLM_TIMEOUT_S   = cfg_int("LLM_TIMEOUT_S", cfg_int("OPENAI_TIMEOUT", 25))
#LLM_MAX_RETRY   = cfg_int("LLM_MAX_RETRY", 2)
#LLM_BACKOFF_SEQ = [0.6, 1.2, 2.0]
#
# (선택) 과거 코드 호환을 위한 별칭 — 남아있는 사용처가 있다면 임시 유지
#CLIENT_TIMEOUT = LLM_TIMEOUT_S

# ========================================================= 
# 경로 준비
# =========================================================
Path(CHAT_FILE).parent.mkdir(parents=True, exist_ok=True)
Path(UPLOAD_DIR).mkdir(parents=True, exist_ok=True)

# =========================================================
# Phase 3) SS AI 로그인 게이트
# - 로그인 전: 로그인 화면만 표시
# - 로그인 후: 회사 선택 완료 시 기존 메인 화면 진입
# =========================================================
if not require_login():
    st.stop()

# =========================================================
# Phase 3) SS AI 관리자 화면
# - 권한 있는 사용자에게만 사이드바 관리자 메뉴 표시
# - 관리자 페이지가 열려 있으면 일반 채팅 화면 대신 관리자 화면 렌더
# =========================================================
render_ssai_admin_sidebar()

if render_ssai_admin_page():
    st.stop()


# ---------------------------------------------------
# OpenAI(LM Studio) 클라이언트 초기화 (단 한번)
CLIENT = OpenAI(
    base_url=cfg_str("LMSTUDIO_BASE_URL", "http://localhost:1234/v1"),
    api_key=cfg_str("LMSTUDIO_API_KEY", "lm-studio"),
)

def _fallback_login_greeting(profile: dict[str, Any]) -> str:
    company_name = str(profile.get("company_name") or "").strip()
    duty_name = str(profile.get("duty_name") or "").strip()
    sims_user_name = str(profile.get("sims_user_name") or "").strip()
    user_type = str(profile.get("user_type") or "").strip()

    if user_type.startswith("SSART"):
        return "안녕하세요. 저는 유통 ERP를 학습하고 있는 SSAI입니다. SIMS 업무를 더 잘 이해하고 도와드리겠습니다."

    title = duty_name or "담당자"
    name_part = f" {sims_user_name}" if sims_user_name else ""

    if company_name:
        return (
            f"안녕하세요. {company_name} {title}{name_part}님, "
            f"저는 유통 ERP를 학습하고 있는 {company_name}의 SSAI입니다. "
            "유통분야 전문 AI가 되어 업무를 도와드리겠습니다."
        )

    return "안녕하세요. 저는 유통 ERP를 학습하고 있는 SSAI입니다. 유통분야 전문 AI가 되어 업무를 도와드리겠습니다."

def _clean_generated_greeting(text: str, fallback: str) -> str:
    """
    LLM이 '예시 1', '옵션 2', 설명문 형태로 답한 경우 차단한다.
    로그인 인사말은 화면에 1개 문장만 보여야 한다.
    """
    raw = str(text or "").strip()

    if not raw:
        return fallback

    bad_words = (
        "예시",
        "옵션",
        "출력",
        "인사말",
        "가능한",
        "다음",
        "아래",
        "후보",
        "샘플",
        "1:",
        "2:",
        "예 1",
        "예 2",
        "SSART_ADMIN",
        "WHOLESALE_USER",
        "WHOLESALE_ADMIN",
    )

    if any(w in raw for w in bad_words):
        return fallback

    # markdown/quote 제거
    raw = raw.replace("```", "").strip()
    raw = re.sub(r"^\s*[>\-\*\d\.\)]\s*", "", raw).strip()
    raw = raw.strip("\"'“”‘’ ")

    # 여러 줄이면 첫 문장만 쓰되, 줄이 너무 많으면 fallback
    lines = [x.strip() for x in raw.splitlines() if x.strip()]
    if len(lines) >= 3:
        return fallback

    if lines:
        raw = " ".join(lines)

    raw = re.sub(r"\s+", " ", raw).strip()

    if not raw:
        return fallback

    if len(raw) > 180:
        raw = raw[:180].rstrip() + "..."

    return raw

def _generate_login_greeting(profile: dict[str, Any]) -> str:
    fallback = _fallback_login_greeting(profile)

    try:
        company_name = str(profile.get("company_name") or "").strip()
        duty_name = str(profile.get("duty_name") or "").strip()
        sims_user_name = str(profile.get("sims_user_name") or "").strip()
        user_type = str(profile.get("user_type") or "").strip()

        role_text = "관리자" if user_type.startswith("SSART") else (duty_name or "담당자")

        if user_type.startswith("SSART"):
            meaning = (
                "안녕하세요. 저는 유통 ERP를 학습하고 있는 SSAI입니다. "
                "SIMS 업무를 더 잘 이해하고 도와드리겠습니다."
            )
        else:
            meaning = (
                f"안녕하세요. {company_name} {role_text}님, "
                f"저는 유통분야 전문 AI가 되려는 {company_name}의 SSAI입니다."
            )

        prompt = f"""
한국어 업무용 로그인 인사말을 작성하세요.

출력 규칙:
- 인사말 본문 1개만 출력
- 예시, 옵션, 번호, 제목, 설명문 출력 금지
- 따옴표 출력 금지
- 1~2문장
- 친절하고 전문적인 말투
- 비밀번호, ID, DB명, 사용자 유형 코드는 절대 언급하지 말 것
- SSART_ADMIN 같은 내부 권한명은 절대 언급하지 말 것

화면 표시용 정보:
회사명: {company_name}
직책/역할: {role_text}
사용자명: {sims_user_name}

반드시 유지할 의미:
{meaning}
""".strip()

        resp = call_chat_with_retry(
            messages=[{"role": "user", "content": prompt}],
            model=st.session_state.get("selected_model") or "local-model",
            temperature=0.7,
            stream=False,
            timeout_s=10,
            max_retry=0,
        )

        text = (resp.choices[0].message.content or "").strip()
        text = re.sub(r"\s+", " ", text).strip()

        if not text:
            return fallback

        # 화면 인사말은 너무 길지 않게 제한
        if len(text) > 220:
            text = text[:220].rstrip() + "..."

        return _clean_generated_greeting(text, fallback)

    except Exception:
        return fallback


def _render_login_greeting_banner() -> None:
    profile = st.session_state.get("__ssai_login_profile")

    if not isinstance(profile, dict):
        return

    sig = "|".join(
        [
            str(profile.get("user_id") or ""),
            str(profile.get("company_id") or ""),
            str(profile.get("sims_user_name") or ""),
            str(profile.get("duty_name") or ""),
            str(profile.get("profile_source") or ""),
        ]
    )

    if st.session_state.get("__ssai_login_greeting_sig") != sig:
        st.session_state["__ssai_login_greeting"] = _generate_login_greeting(profile)
        st.session_state["__ssai_login_greeting_sig"] = sig

    greeting = str(st.session_state.get("__ssai_login_greeting") or "").strip()
    if greeting:
        st.info(greeting)


# 업로드 폴더 쓰기 가드
try:
    testfile = Path(UPLOAD_DIR) / ".write_test"
    testfile.write_text("ok", encoding="utf-8")
    testfile.unlink(missing_ok=True)
except Exception as e:
    st.error(f"📁 업로드 폴더 쓰기 실패: {UPLOAD_DIR} ({e}) — 권한/경로를 확인하세요.")

# =========================
# 의존성 감지
# =========================
try:
    import pytesseract
    _TESS_AVAILABLE = True
except Exception:
    _TESS_AVAILABLE = False

try:
    from PIL import Image, ImageOps, ImageFilter, ImageEnhance
    from PIL import ImageFile
    ImageFile.LOAD_TRUNCATED_IMAGES = True
    Image.MAX_IMAGE_PIXELS = 50_000_000  # 초고해상도 이미지 폭탄 방어
    _PIL_AVAILABLE = True
except Exception:
    _PIL_AVAILABLE = False

try:
    import magic  # python-magic (선택)
    _MAGIC_AVAILABLE = True
except Exception:
    _MAGIC_AVAILABLE = False

try:
    import chardet  # (선택) CSV 인코딩 탐지
    _CHARDET_AVAILABLE = True
except Exception:
    _CHARDET_AVAILABLE = False

# =========================
# SIMS 컨텍스트 메타 질문 답변 시도
# =========================
#def _try_answer_ctx_meta_question(user_input: str, room: dict) -> bool:
#    import re
#    ss = st.session_state
#    ctx = ss.get("__sims_ctx") or ss.get("__sims_context_obj")  # chat_middleware가 쓰는 키들(SSOT)
#    if not isinstance(ctx, dict):
#        return False
#
#    text = (user_input or "").strip()
#    if not text:
#        return False
#
#    # "컨텍스트/현재 데이터/행 수/컬럼" 계열만 먼저 처리
#    pat = r"(컨텍스트|현재\s*데이터|지금\s*데이터|행\s*수|몇\s*건|컬럼|열\s*목록|필드)"
#    if not re.search(pat, text):
#        return False
#
#    data = ctx.get("data") if isinstance(ctx.get("data"), dict) else ctx
#    meta = (data.get("meta") or {}) if isinstance(data, dict) else {}
#    cols = data.get("columns") or []
#    action = meta.get("action") or ctx.get("action") or "(unknown)"
#    row_total = meta.get("row_count_total") or meta.get("row_count") or "?"
#
    # 너무 길면 컬럼은 앞부분만
#    cols_preview = ", ".join(cols[:30]) + (f" ... (+{max(0, len(cols)-30)}개)" if len(cols) > 30 else "")
#
#    answer = (
#        f"현재 SIMS 컨텍스트 요약\n"
#        f"- action: {action}\n"
#        f"- rows: {row_total}\n"
#        f"- cols: {len(cols)}\n"
#        f"- columns: {cols_preview if cols else '(없음)'}"
#    )
#
#    room.setdefault("messages", []).append({
#        "id": str(uuid.uuid4()),
#        "seq": _next_seq(),
#        "role": "assistant",
#        "content": answer,
#        "time": make_ts(),
#    })
#    save_chat_rooms()
#    return True

# =========================
# 공통 유틸
# =========================
def _truncate(s: str, limit: int = MAX_PREVIEW_CHARS) -> str:
    s = s or ""
    return (s[:limit] + f"\n\n... (미리보기 {limit}자 제한으로 이후 생략)") if len(s) > limit else s

def _out(text: str, preview: bool = True) -> str:
    return _truncate(text) if preview else (text or "")

def _bytes_len(uploaded_file) -> int:
    pos = uploaded_file.tell()
    uploaded_file.seek(0, os.SEEK_END)
    size = uploaded_file.tell()
    uploaded_file.seek(pos)
    return size

def sanitize_filename(name: str) -> str:
    return re.sub(r'[^a-zA-Z0-9가-힣._-]', '_', name)

def _effective_upload_dir() -> Path:
    """
    실제 첨부 저장 폴더.

    우선순위:
    1. 로그인 사용자 + 선택 회사가 있으면 사용자별 uploads 폴더
       SSAI_STORAGE_ROOT/company_N/user_N/uploads
    2. 실패 시 기존 UPLOAD_DIR fallback
    """
    try:
        user = get_current_user()
        company = get_selected_company()

        if user and isinstance(company, dict):
            company_id = company.get("company_id")
            user_id = getattr(user, "user_id", None)

            if company_id and user_id:
                return get_user_area_dir(
                    company_id=company_id,
                    user_id=user_id,
                    area="uploads",
                    create=True,
                )
    except Exception as e:
        try:
            log.warning(
                "[upload.storage] user upload dir fallback: %s: %s",
                type(e).__name__,
                e,
            )
        except Exception:
            pass

    fallback = Path(UPLOAD_DIR)
    fallback.mkdir(parents=True, exist_ok=True)
    return fallback


def _make_upload_save_path(original_name: str, file_hash: str) -> Path:
    """
    첨부 파일 저장 경로 생성.
    동일 파일명 충돌 방지를 위해 hash 12자리 suffix 사용.
    """
    safe_name = make_safe_filename(original_name or "upload")
    base, ext = os.path.splitext(safe_name)

    base = sanitize_filename(base).strip("._ ") or "upload"
    ext = sanitize_filename(ext).strip()

    if len(base) > 120:
        base = base[:120]

    if len(ext) > 20:
        ext = ext[:20]

    upload_dir = _effective_upload_dir()
    return upload_dir / f"{base}__{str(file_hash or uuid.uuid4().hex)[:12]}{ext}"


def _sniff_mime(file_bytes: bytes, fallback: str) -> str:
    if _MAGIC_AVAILABLE:
        try:
            return magic.from_buffer(file_bytes[:2048], mime=True) or fallback
        except Exception:
            pass
    return fallback or "application/octet-stream"

def _sha256_of_filelike(f) -> str:
    pos = f.tell()
    f.seek(0)
    h = hashlib.sha256()
    for chunk in iter(lambda: f.read(8192), b""):
        h.update(chunk)
    f.seek(pos)
    return h.hexdigest()

if "extraction_cache" not in st.session_state:
    st.session_state.extraction_cache = {}  # {(file_hash, preview, ocr_conf): str}

def cleanup_uploads(max_bytes=2_000_000_000, upload_dir: str | Path | None = None):  # 2GB
    """
    업로드 폴더 용량 정리.

    기본은 현재 로그인 사용자의 uploads 폴더 기준.
    """
    try:
        root = Path(upload_dir) if upload_dir else _effective_upload_dir()
        root.mkdir(parents=True, exist_ok=True)

        files = [
            (
                p.name,
                p.stat().st_size,
                p.stat().st_mtime,
            )
            for p in root.iterdir()
            if p.is_file()
        ]

        total = sum(s for _, s, _ in files)

        if total <= max_bytes:
            return

        for name, size, _ in sorted(files, key=lambda x: x[2]):
            try:
                (root / name).unlink(missing_ok=True)
            except Exception:
                pass

            total -= size

            if total <= max_bytes:
                break

    except Exception:
        pass


# =========================
# OCR
# =========================
def ocr_image_pil(
    img: "Image.Image",
    langs: list,
    psm: int = 3,
    oem: int = 3,
    upscale: bool = True,
    binarize: bool = True,
    denoise: bool = False,
) -> str:
    if not _TESS_AVAILABLE or not _PIL_AVAILABLE:
        return ""

    proc = img.convert("RGB")

    # 업스케일
    w, h = proc.size
    if upscale and max(w, h) < 1600:
        scale = max(1.0, 1600 / max(w, h))
        proc = proc.resize((int(w * scale), int(h * scale)), Image.LANCZOS)

    # 이진화
    if binarize:
        g = ImageOps.grayscale(proc)
        proc = g.point(lambda x: 255 if x > 128 else 0, mode="1").convert("L")

    if denoise:
        proc = proc.filter(ImageFilter.MedianFilter(size=3))

    try:
        proc = ImageEnhance.Contrast(proc).enhance(1.2)
    except Exception:
        pass

    lang = "+".join(langs) if langs else "kor"
    config = f"--psm {psm} --oem {oem}"

    try:
        text = pytesseract.image_to_string(proc, lang=lang, config=config)
        return text.strip()
    except Exception:
        return ""

def _ocr_conf_tuple() -> tuple:
    return (
        tuple(st.session_state.get("ocr_langs", ["kor","eng"])),
        int(st.session_state.get("ocr_psm", 3)),
        int(st.session_state.get("ocr_oem", 3)),
        bool(st.session_state.get("ocr_upscale", True)),
        bool(st.session_state.get("ocr_binarize", True)),
        bool(st.session_state.get("ocr_denoise", False)),
    )

# === 시간/순서 유틸 ===
import re
st.session_state.setdefault("__seq", 0)

def _next_seq():
    st.session_state["__seq"] += 1
    return st.session_state["__seq"]

def make_ts() -> str:
    # 모든 메시지를 초 단위로 통일
    from datetime import datetime
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

def normalize_ts(ts: str) -> str:
    """렌더용 표시 통일: 분만 있으면 ':00' 붙임, ISO 'T' 제거, 마이크로초 제거"""
    if not ts:
        return ""
    ts = str(ts).strip().replace("T", " ")
    ts = ts.split(".", 1)[0]  # .123456 제거
    if re.fullmatch(r"\d{4}-\d{2}-\d{2} \d{2}:\d{2}", ts):
        ts += ":00"
    return ts

# =========================
# 요약 파이프라인
# =========================
MAX_CHUNK_CHARS = 4000
CHUNK_OVERLAP   = 200
DEFAULT_SUMMARY_TARGET = 1500

def _split_into_chunks(text: str, chunk_size: int = MAX_CHUNK_CHARS, overlap: int = CHUNK_OVERLAP) -> list[str]:
    text = text or ""
    if len(text) <= chunk_size:
        return [text]
    chunks, i = [], 0
    step = max(1, chunk_size - overlap)
    while i < len(text):
        chunks.append(text[i:i+chunk_size])
        i += step
    return chunks

def summarize_text_long(
    text: str,
    target_chars: int = DEFAULT_SUMMARY_TARGET,
    system_prompt: str | None = "너는 문서 요약 전문가다. 수치/날짜/단위는 보존하고 군더더기를 제거하라."
) -> str:
    if not text or not text.strip():
        return ""

    def _first_content(resp) -> str:
        try:
            return (resp.choices[0].message.content or "").strip()
        except Exception:
            return ""

    model_id = st.session_state.get("selected_model") or "local-model"

    # 1) 청크 요약
    parts: list[str] = []
    for idx, ck in enumerate(_split_into_chunks(text), start=1):
        msgs = [
            {"role": "user",
             "content": f"아래 텍스트의 핵심만 간결히 요약해줘(불필요한 개행/잡음 제거):\n\n{ck}"}
        ]
        if system_prompt:
            msgs.insert(0, {"role": "system", "content": system_prompt})

        try:
            resp = call_chat_with_retry(
                messages=msgs,
                model=model_id,
                temperature=0.2,
                stream=False,
            )
            parts.append(_first_content(resp))
        except Exception as e:
            parts.append(f"[청크 요약 실패 {idx}: {e}]")

    # 2) 메타 요약 (부분 요약 통합)
    merged = "\n\n".join(parts)
    merged = _clip_for_model(merged, limit=120_000)

    msgs2 = [{
        "role": "user",
        "content": (
            f"아래 부분 요약들을 하나로 통합해 핵심 위주로 정리해줘. "
            f"최종 길이는 약 {target_chars}자 내외. 항목형/문단형 혼합 OK, 수치/날짜 보존:\n\n{merged}"
        )
    }]
    if system_prompt:
        msgs2.insert(0, {"role": "system", "content": system_prompt})

    try:
        resp2 = call_chat_with_retry(
            messages=msgs2,
            model=model_id,
            temperature=0.2,
            stream=False,
        )
        final = _first_content(resp2)
    except Exception as e:
        return f"[최종 요약 실패] {e}"

    # 3) 길면 한 번 더 압축
    if len(final) > target_chars + 400:
        msgs3 = [{"role": "user", "content": f"다음 요약을 {target_chars}자 내로 더 압축해줘:\n\n{final}"}]
        if system_prompt:
            msgs3.insert(0, {"role": "system", "content": system_prompt})
        try:
            resp3 = call_chat_with_retry(
                messages=msgs3,
                model=model_id,
                temperature=0.2,
                stream=False,
            )
            final = _first_content(resp3) or final
        except Exception as e:
            final = f"[재압축 실패] {e}\n\n{final}"

    return final

# =========================
# LLM 호출 헬퍼/가드
# =========================
BASE_SYSTEM_PROMPT = (
    "너는 유통 ERP인 SIMS의 문서를 안전하게 요약·분석하는 어시스턴트다. "
    "테이블의 관계에 충실하고, 숫자/단위/날짜를 보존하라. "
    "SIMS 컨텍스트에 JSON이 주어지면 meta/aggregations/records 구조를 따른다. "
    "✅ aggregations(특히 basis='DB')가 있으면 그것이 '전체(정답) 집계'이므로 최우선으로 사용하라. "
    "records는 샘플일 수 있다(meta.sample_rows < meta.row_count_total이면 샘플). "
    "샘플 records로 전체 건수/전체 distinct를 추정해선 안 된다. "
    "aggregations가 없을 때만 records 범위 내에서 직접 그룹핑/합계를 계산하라. "
    "부서별 인원 수, 사용자 수, 코드별 집계 등과 같이 특정 항목을 물어보면 "
    "가능하면 aggregations를 우선 사용하고(없으면 records로 계산) 답하라. "
    "데이터에 실제로 존재하는 컬럼을 '없다'고 말해서는 안 된다."
    "목록을 나열해야 하면 최대 20개까지만 보여주고, 더 많으면 '총 n건' 요약 후 조건 조회(패널 필터)로 안내하라."
    "SIMS 컨텍스트가 제공된 상태에서는 '정보가 부족'하다는 답변을 하지 말고, "
    "meta(row_count/cols/action)와 records/aggregations를 근거로 반드시 답하라. "
    "사용자가 '컨텍스트/현재 데이터/컬럼/행 수/몇 건'을 물으면 meta/columns를 근거로 즉시 설명하라."
    "사용자가 '이 결과 분석/요약/정리'처럼 최신 조회 결과에 대한 분석을 요청하면 "
    "meta.summary_md, meta.message, aggregations, *_counts, columns, records 순서로 근거를 사용하라. "
    "records는 샘플일 수 있으므로 전체 분포/전체 합계는 meta나 aggregations에 있을 때만 단정하라. "
    "분석 답변은 ①핵심 요약 ②주요 수치 ③주의/확인할 점 ④다음 조회 제안 순서로 짧게 정리하라."
    "SIMS_JSON.kind가 SIMS_ANALYSIS_CONTEXT_V1이면 이것은 샘플 표가 아니라 최신 SIMS 조회 결과 전체를 Python으로 집계한 분석 컨텍스트다. "
    "이 경우 summary, shortage_grade_counts, forecast_grade_counts, trend_judge_counts는 전체 결과 기준이므로 샘플 데이터라고 표현하지 마라. "
    "risk_products_top과 grade_samples는 품목 설명용 대표 목록이며, 전체 건수 판단은 summary와 *_counts를 우선 사용하라. "
    "SIMS/ERP 데이터의 건수/컬럼/샘플은 실제 컨텍스트(SIMS_JSON/SIMS CONTEXT)나 NLQ 조회 결과가 있을 때만 말하고, "
    "없으면 추측/창작하지 말고 “현재 컨텍스트가 없어 확인 불가”라고 답하라."
)
GENERAL_SYSTEM_PROMPT = (
    "너는 유능한 한국어 AI 어시스턴트다. "
    "사용자의 일반 지식 질문(상식/설명/개념/글쓰기/요약 등)에 친절하고 정확하게 답하라. "
    "확실하지 않으면 그 점을 명확히 말하고, 필요하면 확인 질문을 하라."
    "사용자가 SIMS/ERP/거래처/사용자/코드 등 업무 데이터를 직접 언급하지 않았다면, "
    "SIMS 컨텍스트/ERP/데이터 부재 같은 표현을 꺼내지 말고 일반 지식으로 답하라."
)

_SIMS_HINT_RE = re.compile(
    r"(SIMS|ERP|거래처|거래처명|거래처코드|거래처마스터|마스터"
    r"|대표자|대표자명|사업자|사업자번호|사업자등록번호"
    r"|전화|전화번호|연락처|휴대폰|핸드폰|팩스|팩스번호"
    r"|영업사원|영업사원명"
    r"|단가적용처|단가적용처명|재고적용처|재고적용처명"
    r"|사용자|사용자명|사용자코드|사용자ID|사번|부서|직책|영업지역"
    r"|코드|코드마스터|업무코드|그룹코드|상세코드|코드종류|코드명|한글명"
    r"|품목|제품|상품|제품코드|상품코드|제품명|상품명|보험코드|바코드|제약사|제조사|제조사명"
    r"|입고|출고|재고|매출|매입|수금|지급|전표"
    r"|거래명세서|세금계산서|거래명세서공통|세금계산서공통"
    r"|거래명세서순번|세금계산서순번|거래명세서구분|세금계산서구분"
    r"|검증|불일치|실재고월집계|장부재고월집계"
    r"|주소|소재지|상세주소|도로명주소|도로명|시도명|시구군명|법정읍면동명|법정동명|지역"
    r"|등록자|등록자명|등록일자|등록일|등록한|등록된"
    r"|수정자|수정자명|수정일자|수정일|수정한|수정된"
    r"|최근입사자|최근\s*입사자"
    r"|Rddbc0\d+|rddbc0\d+)",
    re.IGNORECASE,
)

def _looks_like_explicit_sims_nlq_command(text: str) -> bool:
    """
    명시적인 SIMS/NLQ 조회 명령인지 판단한다.
    이런 문장은 최신 결과 후속 질문이 아니라 NLQ router로 보내야 한다.
    """
    t = (text or "").strip()
    if not t:
        return False

    # "조회결과" 같은 후속 질문은 제외하고,
    # 문장 끝이 조회/검색/목록 계열인 실제 명령만 막는다.
    if not re.search(
        r"((조회|검색|목록)(해줘|해주세요|해 주세요)?|보여\s*(줘|주세요)?|알려\s*(줘|주세요)?)[.!? ]*$",
        t,
    ):
        return False

    explicit_actions = (
        "품목별 매출 추세 요약표",
        "품목별 매출 추세",
        "품목별 매출 예상",
        "품목별 재고부족현황",
        "입고명세",
        "출고명세",
        "매입명세",
        "매출명세",
        "매입거래명세서",
        "매출거래명세서",
        "매입세금계산서",
        "매출세금계산서",        
        "거래명세서 공통",
        "세금계산서 공통",
        "거래명세서",
        "세금계산서",        
        "입고 거래명세서 불일치",
        "입고 세금계산서 불일치",
        "출고 거래명세서 불일치",
        "출고 세금계산서 불일치",
        "실재고월집계",
        "장부재고월집계",
        "제품수불현황",
        "제품수불부",
        "제품재고현황",
        "제품재고장",
        "거래처 조회",
        "제품 조회",
        "사용자 조회",
        "업무코드 조회",
        "도로명주소",
    )

    return any(a in t for a in explicit_actions)

def _is_doc_validation_nlq_text(text: str) -> bool:
    """
    입고/출고 ↔ 거래명세서/세금계산서 검증 직접조회 문장인지 판정한다.

    이런 문장은 거래명세서/세금계산서 공통 조회 보정 대상이 아니다.
    예:
    - 입고↔세금계산서 검증 2025 조회
    - 출고↔세금계산서 검증 2023 ~2026 조회
    - 입고 세금계산서 불일치 20230101~20260608 조회
    """
    t = str(text or "").strip()
    if not t:
        return False

    compact = re.sub(r"\s+", "", t)

    has_check = any(w in compact for w in ("검증", "불일치", "누락"))
    has_side = any(w in compact for w in ("입고", "매입", "출고", "매출"))
    has_doc = any(w in compact for w in ("거래명세서", "세금계산서"))

    return bool(has_check and has_side and has_doc)


# NLQ 문장 보정: 짧은 명령형 문장을 공통 조회 형태로 보정해서 NLQ router로 보내는 함수
# 예: "세금계산서 2020 조회" → "세금계산서 공통 2020 조회"
def _normalize_doc_common_nlq_text(text: str) -> str:
    """
    짧은 문서 조회 문장을 공통 조회 문장으로 보정한다.

    예:
    - 세금계산서 2017 ~ 2026 조회
      → 세금계산서 공통 2017 ~ 2026 조회

    - 계산서 거래처 화순 조회
      → 세금계산서 공통 거래처명 화순 조회

    - 거래명세서 2026 조회
      → 거래명세서 공통 2026 조회
    """
    raw = str(text or "").strip()
    if not raw:
        return raw

    # 검증 직접조회는 거래명세서/세금계산서 공통 조회로 보정하면 안 된다.
    # 예: "입고↔세금계산서 검증 2025 조회"를
    #     "세금계산서 공통 매입분 거래처명 ↔ 검증 2025 조회"로 바꾸면 안 된다.
    if _is_doc_validation_nlq_text(raw):
        return raw

    # 순번을 명시한 문장은 건드리지 않는다.
    if "순번" in raw:
        return raw

    # 조회성 문장만 처리
    if not re.search(r"(조회|검색|보여\s*줘|보여주세요|알려\s*줘|알려주세요)\s*$", raw):
        return raw

    t = raw

    if "거래명세서" in t:
        if "공통" not in t:
            t = t.replace("거래명세서", "거래명세서 공통", 1)

    elif "세금계산서" in t:
        if "공통" not in t:
            t = t.replace("세금계산서", "세금계산서 공통", 1)

    elif re.search(r"(^|\s)계산서", t):
        # 사용자가 "계산서"라고만 하면 세금계산서 공통으로 본다.
        if "공통" not in t:
            t = re.sub(r"(^|\s)계산서", r"\1세금계산서 공통", t, count=1)

    else:
        return raw

    # 거래처 화순 → 거래처명 화순
    t = re.sub(r"(거래처|매입처|매출처)\s+(?!명)", "거래처명 ", t)

    # 공통 매입 / 공통 매출 → 공통 매입분 / 공통 매출분
    t = re.sub(r"공통\s+매입(?!분)", "공통 매입분", t)
    t = re.sub(r"공통\s+매출(?!분)", "공통 매출분", t)

    return re.sub(r"\s+", " ", t).strip()

# 거래처 조건이 있는 문장은 공통 조회로 보정해서 NLQ router로 보내는 함수
# 예: "거래명세서 매입 거래처명 인천약품 조회" → "거래명세서 공통 매입분 거래처명 인천약품 조회"
def _normalize_doc_vendor_nlq_text(text: str) -> str:
    """
    거래명세서/세금계산서 + 거래처 조건 문장이
    거래처 마스터 조회로 빠지지 않도록 IO 문장으로 정규화한다.

    예:
    - 거래명세서 매입 거래처명 인천약품
      → 거래명세서 공통 매입분 거래처명 인천약품 조회

    - 거래명세서 매입 거래처 인천약품 조회
      → 거래명세서 공통 매입분 거래처명 인천약품 조회

    - 거래명세서 매입 인천약품 2026 조회
      → 거래명세서 공통 매입분 거래처명 인천약품 2026 조회
    """
    raw = str(text or "").strip()
    if not raw:
        return raw

    # 검증 직접조회는 거래처 조건 보정 대상이 아니다.
    # "↔ 검증" 같은 토큰을 거래처명으로 만들면 안 된다.
    if _is_doc_validation_nlq_text(raw):
        return raw

    t = raw
    
    if not re.search(r"(거래명세서|세금계산서)", t):
        return raw

    # 조회성 문장이 아니면 건드리지 않는다.
    if not re.search(r"(조회|검색|보여\s*줘|보여주세요|알려\s*줘|알려주세요)\s*$", t):
        # 끝에 조회가 없어도 문서 + 매입/매출 + 거래처명이면 조회로 보정
        if not re.search(r"(거래명세서|세금계산서).*(매입|매출).*(거래처|거래처명|매입처|매출처)", t):
            return raw

    doc_word = "세금계산서" if "세금계산서" in t else "거래명세서"

    if re.search(r"매입분?|입고", t):
        div_word = "매입분"
    elif re.search(r"매출분?|출고", t):
        div_word = "매출분"
    else:
        return raw

    # 조회/검색 등 끝 단어 제거
    body = re.sub(r"(조회|검색|보여\s*줘|보여주세요|알려\s*줘|알려주세요)\s*$", "", t).strip()

    # 문서명/공통/구분 제거
    body = re.sub(r"(거래명세서|세금계산서)", "", body).strip()
    body = re.sub(r"공통", "", body).strip()
    body = re.sub(r"(매입분?|매출분?|입고|출고)", "", body).strip()

    # 거래처 라벨 제거
    body = re.sub(r"(거래처명|거래처|매입처명|매입처|매출처명|매출처)", "", body).strip()

    if not body:
        return raw

    # 날짜/연도 토큰 분리
    date_tokens = re.findall(
        r"(?:20\d{2}(?:[-./]?\d{2}){0,2}|20\d{6}|20\d{4})",
        body,
    )

    name_part = body
    for d in date_tokens:
        name_part = name_part.replace(d, " ")

    name_part = re.sub(r"\s+", " ", name_part).strip()
    date_part = " ".join(date_tokens).strip()

    # 이름 없이 연도만 있는 경우는 거래처 조건으로 만들면 안 된다.
    if not name_part:
        return raw

    normalized = f"{doc_word} 공통 {div_word} 거래처명 {name_part}"
    if date_part:
        normalized += f" {date_part}"
    normalized += " 조회"

    return normalized


def is_sims_result_followup_question(text: str) -> bool:
    """
    최신 SIMS 조회 결과에 대한 후속 질문인지 판단한다.

    중요:
    - True이면 SIMS_JSON은 LLM에 붙여야 한다.
    - 하지만 NLQ router로 보내면 안 된다.
    - 예: "현재 조회결과 알려줘", "이 결과 요약해줘", "방금 표 분석해줘"
    """
    t = (text or "").strip()
    if not t:
        return False

    # 명시적인 NLQ 조회 명령은 후속 질문으로 보지 않는다.
    # 예: "품목별 매출 추세 요약표 2025년 조회"
    if _looks_like_explicit_sims_nlq_command(t):
        return False

    # 최신 SIMS 컨텍스트가 없으면 후속 분석 질문으로 보지 않는다.
    try:
        has_sims_ctx = bool(
            st.session_state.get("__sims_ctx")
            or st.session_state.get("__sims_context")
            or st.session_state.get("__sims_context_obj")
            or st.session_state.get("__sims_context_text")
        )
    except Exception:
        has_sims_ctx = False

    if not has_sims_ctx:
        return False

    # ------------------------------------------------------------
    # 최신 SIMS 결과를 대상으로 한 암시적 후속분석/상세질문
    # 예:
    # - 제품별 매출 금액 및 수량을 조회하고, 상위 20개 제품을 보여줘
    # - 거래처별 매출 금액 및 수량을 조회하고, 상위 20개 거래처를 보여줘
    # - 대학약국 거래 내역 상세 조회
    #
    # 단, "출고명세 거래처명 대학약국 조회"처럼 명시적 SIMS 조회명령은
    # 위의 _looks_like_explicit_sims_nlq_command()에서 이미 False 처리되어
    # NLQ router로 내려가게 한다.
    # ------------------------------------------------------------
    current_mark_words = (
        "현재표",
        "현재 표",
        "현재조회",
        "현재 조회",
        "현재조회결과",
        "현재 조회결과",
        "이표",
        "이 표",
        "위표",
        "위 표",
        "방금표",
        "방금 표",
    )

    implicit_domain_words = (
        "제품별",
        "품목별",
        "거래처별",
        "재고위치별",
        "담당자별",
        "영업사원별",
        "제품명",
        "거래처명",
        "재고위치",
        "거래 내역",
        "거래내역",
        "상세 내역",
        "상세내역",
        "매출",
        "매입",
        "입고",
        "출고",
        "수량",
        "금액",
        "공급가액",
        "세액",
        "합계금액",
    )

    implicit_request_words = (
        "조회",
        "검색",
        "보여",
        "알려",
        "분석",
        "요약",
        "정리",
        "상세",
        "내역",
        "목록",
        "상위",
        "하위",
        "TOP",
        "top",
        "합계",
        "평균",
        "건수",
        "비중",
        "분포",
    )

    if any(w in t for w in current_mark_words) and any(w in t for w in implicit_request_words):
        return True

    if any(w in t for w in implicit_domain_words) and any(w in t for w in implicit_request_words):
        return True

    # "조회결과"처럼 붙여 쓰는 경우까지 포함
    followup_words = (        
        "이 결과",
        "위 결과",
        "현재 결과",
        "조회 결과",
        "조회결과",
        "현재 조회",
        "현재 조회결과",
        "방금",
        "방금 표",
        "이 표",
        "위 표",
        "현재 표",
        "현재 데이터",
        "현재 컨텍스트",
        "sims 결과",
        "SIMS 결과",
        "주요수치",
        "주요 수치",
        "확인할점",
        "확인할 점",
        "다음조회",
        "다음 조회",
        "제안",
        "표",
        "표로",

    )

    request_words = (
        "알려",
        "말해",
        "보여",
        "분석",
        "요약",
        "정리",
        "해석",
        "설명",
        "인사이트",
        "특징",
        "문제점",
        "이상",
        "위험",
        "주의",
        "상위",
        "하위",
        "합계",
        "평균",
        "건수",
        "몇 건",
        "분포",
        "추세",
        "예상",
        "부족",
        "비교",
        "중요",
        "핵심",
    )

    if any(w in t for w in followup_words) and any(w in t for w in request_words):
        return True

    # 분석/KPI 결과에 대한 후속 질문 보강
    # 예:
    # - "부족등급 분포 기준으로 위험한 제품을 설명해줘"
    # - "예상등급 분포 알려줘"
    # - "감소예상 제품 정리해줘"
    # 이런 문장은 "제품 조회"가 아니라 최신 SIMS 결과 분석 요청이다.
    analytics_followup_words = (
        "부족등급",
        "예상등급",
        "추세판정",
        "재고커버",
        "재고커버월수",
        "분포",
        "위험",
        "위험한 제품",
        "주의",
        "주의 제품",
        "부족 제품",
        "재고없음",
        "수요관찰",
        "감소예상",
        "상승예상",
        "반품주의",
        "자료부족",
        "정상 제품",
        "핵심요약",
        "핵심 요약",
        "주요수치",
        "주요 수치",
        "확인할점",
        "확인할 점",
        "다음조회",
        "다음 조회",
        "다음 조회 제안",
        "제안",
        "표로",
        "shortage_grade",
        "forecast_grade",
        "trend_judge",
        "shortage_grade_counts",
        "forecast_grade_counts",
        "trend_judge_counts",
    )

    if has_sims_ctx and any(w in t for w in analytics_followup_words) and any(w in t for w in request_words):
        return True

    # 짧은 단독 명령도 허용: "분석해줘", "요약해줘", "정리해줘"

    if re.fullmatch(
        r"(분석|요약|정리|해석|설명)(해줘|해 주세요|해주시겠어요|해봐|해 봐)?[.!? ]*",
        t,
    ):
        return True

    return False


_ANALYTICS_KPI_SOURCE_ACTIONS = {
    "품목별 매출 추세 분석",
    "품목별 매출 추세 요약표",
    "품목별 매출 예상",
    "품목별 재고부족현황",
}


def _current_source_is_analytics_kpi() -> bool:
    """
    현재표 후속분석 기준 source가 분석/KPI 계열인지 판정한다.

    주의:
    - 현재표 후속표가 한 번 생성되면 __sims_last_table_action 이
      "현재표 제품별 예상 TOP 20" 같은 파생 action으로 바뀔 수 있다.
    - 그래도 __sims_current_table_source_action / 패널 마지막 final action /
      현재 선택 action을 함께 보면 원본이 분석/KPI인지 판단할 수 있다.
    """
    try:
        ss = st.session_state
        candidates: list[str] = []

        for key in (
            "__sims_current_table_source_action",
            "__sims_last_table_action",
            "__sims_panel_last_final_action",
        ):
            v = str(ss.get(key) or "").strip()
            if v:
                candidates.append(v)

        sel = ss.get("__sims_selected") or {}
        if isinstance(sel, dict):
            category = str(sel.get("category") or "").strip()
            action = str(sel.get("action") or "").strip()
            if category == "분석/KPI" and action:
                candidates.append(action)

        for action in candidates:
            # 정확히 원본 action인 경우
            if action in _ANALYTICS_KPI_SOURCE_ACTIONS:
                return True

            # "분석/KPI::품목별 매출 예상" 형태로 저장된 경우
            if any(src in action for src in _ANALYTICS_KPI_SOURCE_ACTIONS):
                return True

            # 현재표 파생 결과 action인 경우
            if action.startswith("현재표") and any(
                w in action
                for w in (
                    "매출추세", "추세판정", "추세",
                    "매출예상", "예상", "예상등급",
                    "재고부족", "부족", "부족등급",
                )
            ):
                return True

        return False
    except Exception:
        return False


def _looks_like_explicit_analytics_base_nlq(text: str) -> bool:
    """
    분석/KPI의 새 조회 문장인지 판정한다.

    현재 분석/KPI 표가 떠 있어도 아래 문장들은 현재표 후속분석이 아니라
    새 NLQ 조회로 보내야 한다.
    """
    t = str(text or "").strip()
    if not t:
        return False

    compact = re.sub(r"\s+", "", t)

    if any(w in compact for w in ("현재표", "현재조회결과", "현재결과")):
        return False

    # 조회라는 단어가 없어도 분석/KPI의 명확한 action 명칭이면
    # 현재표 implicit follow-up으로 보내지 않고 새 조회 action으로 보낸다.
    explicit_action_keywords = (
        "품목별매출추세요약",
        "품목별추세요약",
        "매출추세요약",
        "추세요약",
        "품목별매출예상",
        "매출예상",
        "예상매출",
        "품목별재고부족현황",
        "재고부족현황",
    )

    if any(key in compact for key in explicit_action_keywords):
        return True

    # 여기부터는 조회성 단어가 있는 문장만 새 조회 후보로 본다.
    if not any(w in compact for w in ("조회", "검색", "보여줘", "만들어줘", "출력")):
        return False

    # 추세판정 필터가 있으면 현재표 후속질문일 가능성이 높다.
    # 예: 품목 매출추세 감소 조회
    trend_words = (
        "감소",
        "증가",
        "안정",
        "자료부족",
        "반품주의",
        "신규증가",
        "신규/증가",
        "감소예상",
        "상승예상",
        "안정예상",
    )
    compact_no_slash = compact.replace("/", "")
    trend_words_no_slash = tuple(x.replace("/", "") for x in trend_words)

    if any(w in compact_no_slash for w in trend_words_no_slash):
        return False

    base_patterns = (
        "매출추세조회",
        "품목별매출추세조회",
        "품목별매출추세분석조회",

        "품목별매출추세요약표조회",
        "품목별매출추세요약조회",
        "품목별추세요약표조회",
        "품목별추세요약조회",
        "매출추세요약표조회",
        "매출추세요약조회",
        "추세요약표조회",
        "추세요약조회",

        "품목별매출예상조회",
        "매출예상조회",
        "예상매출조회",

        "품목별재고부족현황조회",
        "재고부족현황조회",
    )

    if any(p in compact for p in base_patterns):
        return True

    # "품목별 매출 추세" + 조회성 문장은 새 분석/KPI 조회로 본다.
    # 단, TOP/상위/상세/목록은 현재표 후속질문일 수 있으므로 제외한다.
    compact_upper = compact.upper()

    if (
        "품목별매출추세" in compact
        and any(w in compact for w in ("조회", "검색", "보여줘", "만들어줘", "출력"))
        and not any(w in compact_upper for w in ("TOP", "상위", "상세", "목록"))
    ):
        return True

    return False


def _looks_like_master_nlq_guard(text: str) -> bool:
    """
    분석/KPI 현재표가 떠 있어도 마스터 조회 문장은 현재표 후속분석으로 가로채지 않는다.

    예:
    - 제품 조회
    - 제품코드 목록
    - 거래처 조회
    - 사용자 조회
    - 업무코드 조회
    - 도로명주소 조회
    """
    t = str(text or "").strip()
    if not t:
        return False

    compact = re.sub(r"\s+", "", t)

    # 사용자가 명시적으로 현재표라고 한 경우는 보호하지 않는다.
    if any(w in compact for w in ("현재표", "현재조회결과", "현재결과", "현재조회자료")):
        return False

    # 조회성 문장이 아니면 마스터 조회 보호 대상이 아니다.
    if not any(w in compact for w in ("조회", "검색", "목록", "상세", "보여줘", "보여주세요")):
        return False

    # 분석/KPI 후속분석 의도가 명확한 경우는 보호하지 않는다.
    analytics_intent_words = (
        "제품별", "품목별", "거래처별", "제조사별", "매입처별", "재고적용처별",
        "예상", "예상등급", "부족", "부족등급", "추세", "추세판정",
        "TOP", "top", "상위", "요약", "분석", "집계",
        "다음월", "3개월", "6개월",
    )
    if any(w in compact for w in analytics_intent_words):
        return False

    master_patterns = (
        "제품조회",
        "제품목록",
        "제품상세",
        "제품코드조회",
        "제품코드목록",
        "제품코드상세",

        "거래처조회",
        "거래처목록",
        "거래처상세",

        "사용자조회",
        "사용자목록",
        "사용자상세",

        "업무코드조회",
        "코드조회",
        "코드명검색",
        "그룹코드조회",

        "도로명주소조회",
        "도로명주소검색",
    )

    if any(p in compact for p in master_patterns):
        return True

    # "제품 한미 조회", "거래처 대학약국 조회", "사용자 김 조회" 같은 자연어 형태 보호
    if re.search(r"^(제품|제품코드|거래처|사용자|업무코드|코드|도로명주소).*(조회|검색|목록|상세)$", compact):
        return True

    return False

def _current_analytics_source_action_for_followup() -> str:
    """
    암시적 분석/KPI 현재표 후속분석은 실제 current source action만 기준으로 판단한다.

    주의:
    - 사용자가 패널 선택만 바꾼 상태에서는 selected action이 새 값이어도
      current source df는 아직 이전 표일 수 있다.
    - 따라서 여기서는 selected action을 기준으로 삼지 않는다.
    """
    try:
        ss = st.session_state

        # 암시적 분석/KPI 현재표 후속질문은 "실제 현재표 source"만 기준으로 한다.
        # 패널 선택값/패널 마지막 action은 사용하지 않는다.
        # 이유: 중간에 제품/거래처/사용자/업무코드 마스터 NLQ를 실행하면
        # 현재표 source는 마스터 표로 바뀌지만 패널 selected action은 분석/KPI로 남아 있을 수 있다.
        candidates = [
            str(ss.get("__sims_current_table_source_action") or "").strip(),
            str(ss.get("__sims_last_table_action") or "").strip(),
        ]

        for action in candidates:
            if not action:
                continue

            # 마스터/NLQ 표가 현재 source이면 분석/KPI 암시적 후속분석 금지
            if action in {
                "거래처 목록",
                "제품코드 목록",
                "제품코드 상세",
                "사용자목록 + 부서명",
                "업무코드 조회",
                "코드명 검색",
                "그룹코드조회",
            }:
                return ""

            if action in _ANALYTICS_KPI_SOURCE_ACTIONS:
                return action

            for src in _ANALYTICS_KPI_SOURCE_ACTIONS:
                if src in action:
                    return src

        return ""
    except Exception:
        return ""

def _implicit_analytics_query_matches_source(text: str, source_action: str) -> bool:
    """
    암시적 현재표 후속분석 질문이 현재 source_action과 맞는지 확인한다.

    예:
    - 예상등급별 요약      → 품목별 매출 예상에서만 허용
    - 부족등급별 요약      → 품목별 재고부족현황에서만 허용
    - 추세판정별 요약      → 품목별 매출 추세 분석/요약표에서만 허용
    - 제조사별 매출 분석   → 매출 추세/매출 예상 source에서만 허용
    """
    compact = re.sub(r"\s+", "", str(text or "").strip())
    src = str(source_action or "").strip()

    if not compact or not src:
        return False

    is_forecast_q = any(
        w in compact
        for w in (
            "예상",
            "매출예상",
            "예상매출",
            "예상등급",
            "예상등급별",
            "예상등급요약",
            "예상등급분석",
            "제품별예상",
            "품목별예상",
            "다음월예상",
            "3개월예상",
            "6개월예상",
        )
    )

    is_shortage_q = any(
        w in compact
        for w in (
            "재고부족",
            "부족",
            "부족수량",
            "부족등급",
            "부족등급별",
            "부족등급요약",
            "부족등급분석",
        )
    )

    is_trend_q = any(
        w in compact
        for w in (
            "매출추세",
            "추세판정",
            "추세판정별",
            "추세판정요약",
            "추세판정분석",
            "추세별",
        )
    )

    is_sales_group_q = (
        any(w in compact for w in ("제조사별", "거래처별", "매입처별", "제품별", "품목별"))
        and any(w in compact for w in ("매출", "금액", "분석", "요약", "집계", "TOP", "top", "상위"))
    )

    if is_shortage_q:
        return src == "품목별 재고부족현황"

    if is_forecast_q:
        return src == "품목별 매출 예상"

    if is_trend_q:
        return src in {
            "품목별 매출 추세 분석",
            "품목별 매출 추세 요약표",
        }

    if is_sales_group_q:
        return src in {
            "품목별 매출 추세 분석",
            "품목별 매출 추세 요약표",
            "품목별 매출 예상",
        }

    return False

def _looks_like_implicit_analytics_current_followup(text: str) -> bool:
    """
    현재 source가 분석/KPI 표일 때,
    사용자가 '현재표'를 생략한 후속질문을 현재표 후속분석으로 본다.

    예:
    - 품목 매출추세 감소 조회
    - 품목 매출추세 감소 품목 top 20
    - 감소 품목 상세히 보여줘
    - 제조사별 매출 분석
    """
    source_action = _current_analytics_source_action_for_followup()
    if not source_action:
        return False

    # 현재표 source가 분석/KPI가 아니면 암시적 "현재표 ..." 변환 금지
    if source_action not in _ANALYTICS_KPI_SOURCE_ACTIONS:
        return False

    t = str(text or "").strip()
    if not t:
        return False

    compact = re.sub(r"\s+", "", t)

    if not _implicit_analytics_query_matches_source(t, source_action):
        return False

    # 마스터 조회 문장은 분석/KPI 현재표 후속분석으로 가로채지 않는다.
    if _looks_like_master_nlq_guard(t):
        return False

    # 분석/KPI 기본 조회 문장은 현재표 후속분석으로 가로채면 안 된다.
    # 예: "품목별 매출 추세 조회"는 새 NLQ 조회로 보내야 1,500행 월별 원자료가 생성된다.
    if _looks_like_explicit_analytics_base_nlq(t):
        return False

    if any(w in compact for w in ("현재표", "현재조회결과", "현재결과")):
        return False

    # 기간/연도/날짜가 들어가면 새 NLQ 조회일 가능성이 높으므로 제외
    if re.search(r"(20\d{2}|20\d{6}|\d{8}|~)", compact):
        return False

    analytics_words = (
        # 매출추세/추세판정
        "매출추세",
        "추세판정",
        "추세",
        "감소",
        "증가",
        "안정",
        "자료부족",
        "반품주의",

        # 매출예상/예상등급
        "매출예상",
        "예상매출",
        "예상",
        "예상등급",
        "예상등급별",
        "예상등급요약",
        "예상등급분석",
        "제품별예상",
        "품목별예상",
        "다음월예상",
        "3개월예상",
        "6개월예상",

        # 재고부족/부족등급
        "재고부족",
        "부족",
        "부족등급",
        "부족등급별",
        "부족등급요약",
        "부족수량",

        # 차원/요청어
        "제품별",
        "품목별",
        "제조사별",
        "거래처별",
        "매입처별",
        "재고적용처별",        
        "TOP",
        "top",
        "상위",
        "목록",
        "상세",
        "집계",
        "요약",
        "분석",
    )

    return any(w in compact for w in analytics_words)


def _normalize_implicit_analytics_current_followup(text: str) -> str:
    t = str(text or "").strip()
    if not t:
        return t
    return f"현재표 {t}"


def _current_table_norm_text(value: Any) -> str:
    return str(value or "").strip()

def _current_table_to_num(series: pd.Series) -> pd.Series:
    try:
        return pd.to_numeric(
            series.astype(str)
            .str.replace(",", "", regex=False)
            .str.replace("원", "", regex=False)
            .str.replace("개", "", regex=False)
            .str.replace("건", "", regex=False)
            .str.strip(),
            errors="coerce",
        ).fillna(0)
    except Exception:
        return pd.Series([0] * len(series), index=series.index)


def _current_table_find_col(
    df: pd.DataFrame,
    *,
    exact: tuple[str, ...] = (),
    include_any: tuple[str, ...] = (),
    exclude_any: tuple[str, ...] = (),
) -> str:
    if not isinstance(df, pd.DataFrame) or df.empty:
        return ""

    cols = [str(c) for c in df.columns]

    for name in exact:
        if name in df.columns:
            return name

    for col in cols:
        s = str(col).replace(" ", "")
        if exclude_any and any(x in s for x in exclude_any):
            continue
        if include_any and any(x in s for x in include_any):
            return col

    return ""




def _current_table_get_latest_df() -> tuple[pd.DataFrame | None, str]:
    """
    현재표 후속표 생성 기준 DF를 가져온다.

    핵심 원칙:
    - 화면의 최신 표(__sims_last_table_key)는 파생 TOP 표일 수 있다.
    - 후속 집계/상세표 기준은 별도 source key(__sims_current_table_source_key)를 우선 사용한다.
    - 같은 table_key에 UI용 200건과 export 전체 DF가 같이 있으면 행 수가 가장 큰 DF를 사용한다.
    """
    try:
        ss = st.session_state

        source_table_key = str(ss.get("__sims_current_table_source_key") or "").strip()
        last_table_key = str(ss.get("__sims_last_table_key") or "").strip()

        def _best_df_for_key(key: str) -> pd.DataFrame | None:
            best_df = None
            best_rows = -1

            for store_name in ("__sims_export_tables_by_key", "sims_export_tables", "sims_tables"):
                store = ss.get(store_name)
                if not isinstance(store, dict):
                    continue

                df = store.get(key)
                if isinstance(df, pd.DataFrame) and not df.empty:
                    rows = int(len(df))
                    if rows > best_rows:
                        best_df = df
                        best_rows = rows

            if isinstance(best_df, pd.DataFrame) and not best_df.empty:
                log.info(
                    "[chat.followup_table] source df selected table_key=%s rows=%s source_key=%s last_key=%s",
                    key,
                    len(best_df),
                    source_table_key,
                    last_table_key,
                )
                return best_df.copy()

            return None

        preferred_keys: list[str] = []
        for key in (source_table_key, last_table_key):
            if key and key not in preferred_keys:
                preferred_keys.append(key)

        for key in preferred_keys:
            df = _best_df_for_key(key)
            if isinstance(df, pd.DataFrame) and not df.empty:
                return df, key

        # source/last key가 모두 비어 있으면 임의로 과거 테이블을 고르지 않는다.
        # 표 없는 조회(0건/안내문) 이후 "현재표 ..." 질문이 들어왔을 때
        # 이전 표를 현재표로 오인하는 것을 막기 위함이다.
        return None, ""

    except Exception:
        log.exception("[chat.followup_table] source df lookup failed")

    return None, ""


def _current_table_should_block_llm_fallback(text: str) -> bool:
    """
    현재표 후속질문 중 LLM으로 임의 답변하면 안 되는 정형 표 요청 판정.

    정책:
    - TOP/목록/표/월별/일자별/거래처별/제품별/조건 필터는 pandas handler가 처리해야 한다.
    - handler가 처리하지 못하면 명확한 안내를 반환하고 LLM 분석으로 넘기지 않는다.
    - "분석해줘/요약해줘/의미 설명" 같은 서술형 질문만 LLM fallback 허용 후보가 된다.
    """
    compact = re.sub(r"\s+", "", str(text or ""))
    if not compact:
        return True

    hard_keywords = (
        "목록", "상세", "상세표", "표", "테이블", "TOP", "top", "상위",
        "월별", "일자별", "날짜별", "요일별",
        "거래처별", "제품별", "품목별", "매입처별", "매출처별", "제조사별", "재고위치별",
        "이상", "이하", "초과", "미만", "같음", "동일", "=",
        "불일치목록", "차이금액", "계산서금액", "거래금액",
    )
    return any(k in compact for k in hard_keywords)


def _push_no_current_table_notice(source_query: str) -> bool:
    """현재표 source가 없을 때 이전 SIMS 컨텍스트로 LLM 답변하지 않도록 안내만 반환한다."""
    return _current_table_push_notice(
        title="현재표 후속분석 불가",
        action="현재표 후속분석 불가",
        message=(
            "현재표로 사용할 조회 결과가 없습니다.\n\n"
            "방금 실행한 조회가 0건 또는 안내 결과로 끝났거나, 현재표 원본이 비어 있습니다.\n"
            "이전 조회표를 잘못 분석하지 않도록 LLM 분석으로 넘기지 않았습니다.\n\n"
            "조회 결과가 있는 화면을 다시 실행한 뒤 현재표 질문을 해 주세요."
        ),
        query_summary="현재표 / 후속분석 불가 / 현재표 원본 없음",
        source_query=str(source_query or ""),
    )

def _current_table_is_blankish_value(v: Any) -> bool:
    """
    현재표 후속표에서 빈값으로 볼 값 판정.
    None, NaN, 문자열 None/nan/<NA> 등을 모두 빈값으로 본다.
    """
    try:
        if pd.isna(v):
            return True
    except Exception:
        pass

    s = str(v).strip()
    return s in ("", "None", "none", "NONE", "nan", "NaN", "NAN", "<NA>", "NaT", "NULL", "null")


def _current_table_fill_alias_values(df: pd.DataFrame) -> pd.DataFrame:
    """
    원본표에는 값이 있는데 후속표 앞쪽 별칭 컬럼이 None으로 보이는 경우 보정.

    예:
    - 제조사명 이 비어 있고 제조사 에 값이 있으면 제조사명 채움
    - 재고위치명 이 비어 있고 재고위치 에 값이 있으면 재고위치명 채움

    실제 값을 임의로 만들지는 않고, 같은 행의 다른 별칭 컬럼 값만 사용한다.
    """
    if not isinstance(df, pd.DataFrame) or df.empty:
        return df

    out = df.copy()

    alias_pairs = [
        ("제조사명", "제조사"),
        ("제조사", "제조사명"),

        ("재고위치명", "재고위치"),
        ("재고위치", "재고위치명"),

        ("거래처명", "매출처명"),
        ("거래처명", "매입처명"),
        ("매출처명", "거래처명"),
        ("매입처명", "거래처명"),

        ("제품명", "품목명"),
        ("제품명", "상품명"),
        ("품목명", "제품명"),
        ("상품명", "제품명"),
    ]

    for target, source in alias_pairs:
        if target not in out.columns or source not in out.columns:
            continue

        try:
            mask = out[target].map(_current_table_is_blankish_value)
            out.loc[mask, target] = out.loc[mask, source]
        except Exception:
            pass

    return out


def _current_table_clean_none_for_display(df: pd.DataFrame) -> pd.DataFrame:
    """
    현재표 후속표 최종 표시 정리.

    주의:
    - 값이 있는 컬럼은 건드리지 않는다.
    - None/nan/<NA>/NaT 같은 표시값만 빈칸으로 정리한다.
    - 재고수량 같은 조건 기준 컬럼은 numeric filter 쪽에서 별도 0 보정한다.
    """
    if not isinstance(df, pd.DataFrame) or df.empty:
        return df

    out = df.copy()

    none_tokens = {
        "None",
        "none",
        "NONE",
        "nan",
        "NaN",
        "NAN",
        "<NA>",
        "NaT",
        "NULL",
        "null",
    }

    for c in out.columns:
        try:
            if str(c).strip() == "순번":
                continue

            s = out[c].astype("object")

            # 실제 None/NaN 계열
            s = s.where(pd.notna(s), "")

            # 문자열 None/nan 계열
            s = s.map(lambda v: "" if str(v).strip() in none_tokens else v)

            out[c] = s
        except Exception:
            pass

    return out

def _current_table_numeric_filter_op(text: str) -> tuple[str, float, str]:
    """
    현재표 숫자 조건 해석.

    지원 예:
    - 0 이하  -> <= 0
    - 0 미만  -> < 0
    - 0 이상  -> >= 0
    - 0 초과  -> > 0
    - 0인     -> == 0
    - 음수/마이너스 -> < 0
    """
    t = str(text or "").strip()
    compact = t.replace(" ", "")

    threshold = 0.0
    m = re.search(r"(-?\d+(?:\.\d+)?)", compact)
    if m:
        try:
            threshold = float(m.group(1))
        except Exception:
            threshold = 0.0

    if any(w in compact for w in ("음수", "마이너스")):
        return "<", 0.0, "미만"

    if any(w in compact for w in ("이상", "크거나같", ">=")):
        return ">=", threshold, "이상"

    if any(w in compact for w in ("초과", "보다큰", ">")):
        return ">", threshold, "초과"

    if any(w in compact for w in ("이하", "작거나같", "<=")):
        return "<=", threshold, "이하"

    if any(w in compact for w in ("미만", "보다작", "<")):
        return "<", threshold, "미만"

    if any(w in compact for w in ("같은", "같음", "동일", "0인", "=0", "==")):
        return "==", threshold, "같음"

    return "", threshold, ""

# 현재표 후속 질문에 대한 답변이 표일 때, SIMS_JSON 컨텍스트를 붙여서 푸시한다.
def _current_table_push_table(
    *,
    title: str,
    action: str,
    df: pd.DataFrame,
    query_summary: str,
    source_query: str,
    source_table_key: str = "",
    source_rows: int | None = None,
    display_limit: int | None = None,
) -> bool:
        
    from app.ui.chat_middleware import push_sims_result_to_chat

    source_table_key = str(source_table_key or "").strip()

    # 현재표 후속표 공통 보정:
    # 1) 제조사명/제조사 같은 별칭 컬럼 값 보존
    # 2) None/nan 표시값 정리
    if isinstance(df, pd.DataFrame) and not df.empty:
        df = _current_table_fill_alias_values(df)
        df = _current_table_clean_none_for_display(df)

    try:
        source_rows_int = int(source_rows) if source_rows is not None else 0
    except Exception:
        source_rows_int = 0

    if not isinstance(df, pd.DataFrame) or df.empty:
        message = f"해당 조회조건의 자료가 없습니다.\n\n조회조건: {query_summary}"
        payload = {
            "final": True,
            "type": "text",
            "title": title,
            "action": action,
            "params": {},
            "data": message,
            "message": message,
            "meta": {
                "nlq": True,
                "current_table_followup": True,
                "_force_push": True,
                "_nlq_nonce": str(uuid.uuid4()),
                "row_count": 0,
                "row_count_total": 0,
                "query_summary": query_summary,
                "summary_md": f"조회조건: {query_summary}",
                "source": "현재표 후속표",
                "source_query": source_query,
                "source_table_key": source_table_key,
                "source_rows": source_rows_int,
            },
        }
        push_sims_result_to_chat(payload, action)
        return True

    df_full = df.copy()

    # 현재표 후속표 공통 화면 표시 정책
    # - TOP N처럼 명시적으로 display_limit이 들어온 경우: 그 값 사용
    # - 일반 목록/분석표: 3,000건까지 화면 표시
    # - 3,000건 초과는 화면 표시 3,000건, 다운로드/분석 기준은 전체 df_full 유지
    try:
        total_rows = int(len(df_full))
    except Exception:
        total_rows = 0

    if display_limit is None:
        display_limit_int = min(3000, total_rows)
    else:
        try:
            display_limit_int = max(1, int(display_limit))
        except Exception:
            display_limit_int = min(3000, total_rows)

    df_display = df_full.head(display_limit_int).copy()

    payload = {
        "final": True,
        "type": "table",
        "title": title,
        "action": action,
        "params": {},
        "df": df_full,
        "df_display": df_display,
        "records": df_display.to_dict(orient="records"),
        "columns": list(df_display.columns),
        "message": f"{title} {len(df_full):,}건",
        "meta": {
            "nlq": True,
            "current_table_followup": True,
            "_force_push": True,
            "_nlq_nonce": str(uuid.uuid4()),
            "row_count": int(len(df_display)),
            "row_count_total": int(len(df_full)),
            "display_row_count": int(len(df_display)),
            "analysis_row_count": int(len(df_full)),
            "row_count_total_for_analysis": int(len(df_full)),
            "query_summary": query_summary,
            "summary_md": f"조회조건: {query_summary}",
            "source": "현재표 후속표",
            "source_query": source_query,
            "source_table_key": source_table_key,
            "source_rows": source_rows_int,
            "table_profile": "current_table_followup",
        },
    }

    push_sims_result_to_chat(payload, action)
    return True

# 현재표 후속 질문에 대한 답변이 텍스트일 때, SIMS_JSON 컨텍스트를 붙여서 푸시한다.
# 예: "이 결과 분석해줘", "이 결과 요약해줘", "이 결과에서 확인할 점 알려줘", "다음 조회 제안해줘"
def _current_table_push_notice(
    *,
    title: str,
    action: str,
    message: str,
    query_summary: str,
    source_query: str,
) -> bool:
    from app.ui.chat_middleware import push_sims_result_to_chat

    payload = {
        "final": True,
        "type": "text",
        "title": title,
        "action": action,
        "params": {},
        "data": message,
        "message": message,
        "content": message,
        "meta": {
            "nlq": True,
            "current_table_followup": True,
            "_force_push": True,
            "_nlq_nonce": str(uuid.uuid4()),
            "row_count": 0,
            "row_count_total": 0,
            "query_summary": query_summary,
            "summary_md": f"조회조건: {query_summary}",
            "source": "현재표 후속계산",
            "source_query": source_query,
            "render_as_text": True,
            "hide_meta_expander": True,
        },
    }

    push_sims_result_to_chat(payload, action)
    return True


def _current_table_group_top(
    df: pd.DataFrame,
    *,
    group_col: str,
    group_name: str,
    top_n: int,
) -> pd.DataFrame:
    tmp = df.copy()

    tmp[group_col] = (
        tmp[group_col]
        .fillna("")
        .astype(str)
        .str.strip()
        .replace("", "(없음)")
    )

    qty_col = _current_table_find_col(
        tmp,
        exact=("수량", "출고수량", "입고수량"),
        include_any=("수량",),
        exclude_any=("단가", "금액"),
    )

    supply_col = _current_table_find_col(
        tmp,
        exact=("공급가액", "공급금액", "매출공급가액", "입고공급가액", "출고공급가액"),
        include_any=("공급가액", "공급금액"),
    )

    tax_col = _current_table_find_col(
        tmp,
        exact=("세액", "부가세", "매출세액", "입고세액", "출고세액"),
        include_any=("세액", "부가세"),
    )

    amount_col = _current_table_find_col(
        tmp,
        exact=("합계금액", "총금액", "매출금액", "입고금액", "출고금액", "금액"),
        include_any=("합계금액", "총금액"),
    )

    numeric_cols: dict[str, str] = {}

    if qty_col:
        tmp["__qty__"] = _current_table_to_num(tmp[qty_col])
        numeric_cols["수량"] = "__qty__"

    if supply_col:
        tmp["__supply__"] = _current_table_to_num(tmp[supply_col])
        numeric_cols["공급가액"] = "__supply__"

    if tax_col:
        tmp["__tax__"] = _current_table_to_num(tmp[tax_col])
        numeric_cols["세액"] = "__tax__"

    if amount_col:
        tmp["__amount__"] = _current_table_to_num(tmp[amount_col])
        numeric_cols["합계금액"] = "__amount__"

    g = tmp.groupby(group_col, dropna=False)

    out = pd.DataFrame({
        group_name: g.size().index.astype(str),
        "건수": g.size().values,
    })

    for display_name, work_col in numeric_cols.items():
        out[display_name] = g[work_col].sum().values

    if "합계금액" not in out.columns and "공급가액" in out.columns and "세액" in out.columns:
        out["합계금액"] = out["공급가액"] + out["세액"]

    sort_col = ""
    for cand in ("합계금액", "공급가액", "수량", "건수"):
        if cand in out.columns:
            sort_col = cand
            break

    if sort_col:
        out = out.sort_values(sort_col, ascending=False)

    out = out.head(top_n).reset_index(drop=True)

    # 원본/이전 파생표에 순번·순위 계열이 섞여 들어온 경우 중복 표시를 방지한다.
    drop_rank_cols = [c for c in ("순번", "조회순번", "순위") if c in out.columns]
    if drop_rank_cols:
        out = out.drop(columns=drop_rank_cols)

    out.insert(0, "순번", range(1, len(out) + 1))

    return out

#def _current_table_to_num(s: pd.Series) -> pd.Series:
#    if pd.api.types.is_numeric_dtype(s):
#        return pd.to_numeric(s, errors="coerce").fillna(0)
#
#    return pd.to_numeric(
#        s.astype(str)
#        .str.replace(",", "", regex=False)
#        .str.replace("원", "", regex=False)
#        .str.replace("개", "", regex=False)
#        .str.strip(),
#        errors="coerce",
#    ).fillna(0)


def _current_table_amount_series(df: pd.DataFrame) -> tuple[pd.Series, str]:
    """
    매출금액 기준.
    - 합계금액이 있으면 합계금액 사용
    - 없으면 공급가액 + 세액
    - 그것도 안 되면 공급가액
    """
    total_col = _current_table_find_col(
        df,
        exact=("합계금액", "매출금액", "출고금액"),
        include_any=("합계금액", "매출금액", "출고금액"),
        exclude_any=("단가", "율"),
    )
    if total_col:
        return _current_table_to_num(df[total_col]), total_col

    supply_col = _current_table_find_col(
        df,
        exact=("공급가액",),
        include_any=("공급가액",),
        exclude_any=("확정", "단가"),
    )
    tax_col = _current_table_find_col(
        df,
        exact=("세액",),
        include_any=("세액",),
        exclude_any=("확정",),
    )

    if supply_col and tax_col:
        return _current_table_to_num(df[supply_col]) + _current_table_to_num(df[tax_col]), "공급가액+세액"

    if supply_col:
        return _current_table_to_num(df[supply_col]), supply_col

    return pd.Series([0] * len(df), index=df.index, dtype="float64"), ""


def _current_table_qty_series(df: pd.DataFrame) -> tuple[pd.Series, str]:
    qty_col = _current_table_find_col(
        df,
        exact=("수량", "출고수량", "매출수량"),
        include_any=("수량",),
        exclude_any=("할증", "재고", "이월", "현재", "부족"),
    )
    if qty_col:
        return _current_table_to_num(df[qty_col]), qty_col

    return pd.Series([0] * len(df), index=df.index, dtype="float64"), ""


def _current_table_date_series(df: pd.DataFrame) -> tuple[pd.Series, str]:
    date_col = _current_table_find_col(
        df,
        exact=("출고일자", "매출일자", "거래일자", "일자", "전표일자", "입고일자"),
        include_any=("일자", "날짜"),
        exclude_any=("등록", "수정", "보험", "유효", "마감"),
    )
    if not date_col:
        return pd.Series([pd.NaT] * len(df), index=df.index), ""

    raw = (
        df[date_col]
        .astype(str)
        .str.replace(r"\D", "", regex=True)
        .str[:8]
    )
    return pd.to_datetime(raw, format="%Y%m%d", errors="coerce"), date_col


def _current_table_push_text(
    *,
    title: str,
    action: str,
    message: str,
    query_summary: str,
    source_query: str,
    source_table_key: str = "",
    source_rows: int | None = None,
) -> bool:
    """
    현재표 후속계산 결과를 text로 채팅에 올린다.

    화면에는 표가 아니라 자연스러운 문장으로 보여준다.
    다만 직전 대형 원본표(예: 81,634건)의 다운로드 버튼이 매 rerun마다
    다시 만들어지는 문제를 막기 위해, 내부 last_table_key만 작은 hidden table로 교체한다.

    원본 후속분석 기준은 source_table_key로 계속 유지한다.
    """
    from app.ui.chat_middleware import push_sims_result_to_chat

    try:
        source_rows_int = int(source_rows) if source_rows is not None else 0
    except Exception:
        source_rows_int = 0

    source_table_key = str(source_table_key or "").strip()

    # 1) 속도용 hidden small table 생성
    #    화면에 직접 렌더하지는 않지만, __sims_last_table_key를 소형 결과로 바꿔
    #    대형 원본표 다운로드 재생성을 막는다.
    try:
        rows: list[dict] = []
        for line in str(message or "").splitlines():
            line = line.strip()
            if not line:
                continue

            if ":" in line:
                k, v = line.split(":", 1)
                rows.append({"항목": k.strip("- ").strip(), "값": v.strip()})
            else:
                rows.append({"항목": "내용", "값": line})

        if not rows:
            rows = [{"항목": "내용", "값": str(message or "").strip()}]

        df_hidden = pd.DataFrame(rows)

        table_key = f"sims_{uuid.uuid4().hex[:8]}"
        ss = st.session_state
        ss.setdefault("sims_tables", {})
        ss.setdefault("sims_export_tables", {})
        ss.setdefault("__sims_export_tables_by_key", {})

        ss["sims_tables"][table_key] = df_hidden
        ss["sims_export_tables"][table_key] = df_hidden
        ss["__sims_export_tables_by_key"][table_key] = df_hidden

        ss["__sims_last_table_key"] = table_key
        ss["__sims_last_table_action"] = str(action or title or "")

        # 후속 집계 기준은 원본 source table로 유지
        if source_table_key:
            ss["__sims_current_table_source_key"] = source_table_key

            # 중요:
            # 현재표 후속 텍스트 결과가 원본 조회 action을 덮어쓰면,
            # 다음 후속질문에서 "제품수불현황 조회" 같은 원본 업무명을 잃는다.
            # 따라서 현재표 후속결과 action으로 source_action을 덮어쓰지 않는다.
            prev_source_action = str(ss.get("__sims_current_table_source_action") or "").strip()
            new_action = str(action or title or "").strip()

            if not prev_source_action and not new_action.startswith("현재표"):
                ss["__sims_current_table_source_action"] = new_action

        log.info(
            "[chat.followup_text] hidden small table stashed table_key=%s rows=%s source_key=%s source_rows=%s action=%s",
            table_key,
            len(df_hidden),
            source_table_key,
            source_rows_int,
            action,
        )
    except Exception:
        log.exception("[chat.followup_text] hidden small table stash failed")

    # 2) 실제 화면 출력은 text payload
    payload = {
        "final": True,
        "type": "text",
        "title": title,
        "action": action,
        "params": {},
        "data": message,
        "message": message,
        "content": message,
        "meta": {
            "nlq": True,
            "current_table_followup": True,
            "_force_push": True,
            "_nlq_nonce": str(uuid.uuid4()),
            "row_count": 0,
            "row_count_total": 0,
            "query_summary": query_summary,
            "summary_md": f"조회조건: {query_summary}",
            "source": "현재표 후속계산",
            "source_query": source_query,
            "source_table_key": source_table_key,
            "source_rows": source_rows_int,
            "render_as_text": True,
            "hide_meta_expander": True,
        },
    }

    push_sims_result_to_chat(payload, action)
    return True


def _try_handle_current_table_dataframe_followup(
    text: str,
    *,
    room: dict,
    make_ts,
    next_seq,
) -> bool:
    """
    '현재표 ... 표로 만들어줘' 계열을 LLM 답변이 아니라 실제 pandas 표로 생성한다.
    """
    t = str(text or "").strip()
    compact = t.replace(" ", "")

    if not t:
        return False

    # 이번 1차는 명시적으로 현재표/현재 조회 결과를 말한 경우만 처리한다.
    if not (
        "현재표" in compact
        or "현재조회결과" in compact
        or "현재결과" in compact
    ):
        return False
    
    source_action_current = str(
        st.session_state.get("__sims_current_table_source_action")
        or st.session_state.get("__sims_last_table_action")
        or ""
    ).strip()

    is_analytics_kpi_source = source_action_current in _ANALYTICS_KPI_SOURCE_ACTIONS    

    # 현재표 후속질문 판정
    # 실제 처리 성격은 아래 action_dispatcher.py가 원본 action 기준으로 결정한다.
    wants_table = any(
        w in t
        for w in (
            "표로",
            "표 ",
            "상세표",
            "목록",
            "TOP",
            "top",
            "상위",
            "분석",
            "집계",
            "요약",
            "금액",
            "수량",
            "횟수",
            "회수",
            "건수",
            "불일치",
            "검증",
            "월별",
            "일자별",
            "날짜별",
            "요일",
            "최고",
            "가장",
            "많은",
            "1위",
        )
    )

    # 분석/KPI 현재표에서는 '조회', '감소', '추세' 같은 말만 있어도
    # LLM으로 보내지 말고 action_dispatcher → analytics_kpi handler로 보낸다.
    if is_analytics_kpi_source and not wants_table:
        wants_table = any(
            w in compact
            for w in (
                "조회",
                "목록",
                "상세",
                "품목",
                "제품",
                "매출추세",
                "추세판정",
                "추세",
                "감소",
                "증가",
                "안정",
                "자료부족",
                "반품주의",
                "TOP",
                "top",
                "상위",
            )
        )

    is_calc_followup = any(
        w in compact
        for w in (
            # 매출/매입/입출고
            "월별매출",
            "일자별매출",
            "월별입고",
            "월별매입",
            "일자별입고",
            "일자별매입",
            "입고금액이가장많은일자",
            "매출이가장많은일자",
            "수량이가장많은제품",
            "입고수량이가장많은제품",
            "출고수량이가장많은제품",
            "입고횟수가가장많은제품",
            "매출횟수가가장많은제품",

            # 제품수불/재고/월집계
            "입고수량합계",
            "출고수량합계",
            "월별입고수량",
            "월별출고수량",
            "월별입고출고수량",
            "제품별입고수량",
            "제품별출고수량",
            "제품별재고수량",
            "재고수량이가장많은제품",
            "제조사별재고수량",
            "발주처별재고수량",
            "입고처별재고수량",
            "매입처별재고수량",
            "거래처별재고수량",
            "재고위치별재고수량",

            # 거래명세서/세금계산서
            "월별계산서금액",
            "일자별계산서금액",
            "계산서금액이가장많은일자",
            "월별거래금액",
            "일자별거래금액",
            "거래금액이가장많은일자",
            "거래처별계산서금액",
            "거래처별거래금액",
            "세금계산서구분별계산서금액",
            "거래명세서구분별거래금액",

            # 검증
            "월별불일치",
            "거래처별불일치",
            "제품별불일치",
            "월별검증",
        )
    )

    if not wants_table and not is_calc_followup:
        return False

    df, table_key = _current_table_get_latest_df()
    if not isinstance(df, pd.DataFrame) or df.empty:
        return False

    # TOP N
    top_n = 20
    has_explicit_top = bool(re.search(r"(?:TOP|top|상위)\s*(\d{1,4})", t))

    m_top = re.search(r"(?:TOP|top|상위)\s*(\d{1,3})", t)
    if m_top:
        try:
            top_n = max(1, min(500, int(m_top.group(1))))
        except Exception:
            top_n = 20

#   @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@
#   현재표 후속질문 action/성격 우선 분기
    # 0) 현재표 action 우선 분기
    # 현재표 후속질문은 질문 단어보다 원본 action/성격을 먼저 본다.

    handled_by_action = handle_current_table_followup_by_action(
        df=df,
        query=t,
        top_n=top_n,
        table_key=table_key,
        source_action=source_action_current,
        helpers={
            "find_col": _current_table_find_col,
            "to_num": _current_table_to_num,
            "push_table": _current_table_push_table,
            "push_notice": _current_table_push_notice,
        },
        log=log,
    )

    if handled_by_action:
        return True

    # 1) 현재표 거래처명 대학약국 상세표 만들어줘
    m_detail = re.search(
        r"(?:현재\s*표|현재표|현재\s*조회\s*결과|현재조회결과)\s*"
        r"(거래처명|거래처|제품명|제품|재고위치명|재고위치)\s*[:=]?\s*"
        r"([^\s,]+)",
        t,
    )

    is_group_request = any(
        w in t
        for w in (
            "제품별",
            "품목별",
            "거래처별",
            "매입처별",
            "매출처별",
            "재고위치별",
            "영업사원별",
            "담당자별",
        )
    )

    if (
        m_detail
        and not is_group_request
        and any(w in t for w in ("상세", "상세표", "내역", "목록", "표로", "만들어"))
    ):
                
        label = m_detail.group(1)
        value = m_detail.group(2).strip()

        if "거래처" in label:
            col = _current_table_find_col(
                df,
                exact=("거래처명", "매출처명", "매입처명", "실납처명", "납품처명"),
                include_any=("거래처", "매출처", "매입처", "실납처", "납품처"),
                exclude_any=("코드", "번호", "분류", "구분"),
            )
            field_name = "거래처명"
        elif "제품" in label:
            col = _current_table_find_col(
                df,
                exact=("제품명", "품목명", "상품명"),
                include_any=("제품명", "품목명", "상품명"),
                exclude_any=("코드", "번호", "분류", "구분"),
            )
            field_name = "제품명"
        else:
            col = _current_table_find_col(
                df,
                exact=("재고위치명", "재고위치", "재고명"),
                include_any=("재고위치", "재고명"),
                exclude_any=("코드", "번호"),
            )
            field_name = "재고위치"

        if not col:
            log.warning("[chat.followup_table] detail col not found label=%s columns=%s", label, list(df.columns)[:30])
            return False

        mask = df[col].astype(str).str.contains(value, case=False, na=False, regex=False)
        out = df.loc[mask].copy()

        title = f"현재표 {field_name} {value} 상세표"
        query_summary = f"현재표 / {field_name} {value} / 상세표"

        log.info(
            "[chat.followup_table] detail table built title=%s source_rows=%s rows=%s table_key=%s",
            title,
            len(df),
            len(out),
            table_key,
        )

        return _current_table_push_table(
            title=title,
            action=title,
            df=out,
            query_summary=query_summary,
            source_query=t,
            source_table_key=table_key,
            source_rows=len(df),
        )

    # 1-1) 현재표에서 재고수량/출고수량/입고수량 조건 목록
    # 예:
    # - 현재표에서 재고수량이 0 이하인 제품 표로 보여줘
    # - 현재표에서 재고수량이 0 이상인 제품 표로 보여줘
    # - 현재표에서 출고수량이 0 초과인 제품 표로 보여줘
    # - 현재표에서 입고수량이 100 이상인 제품 표로 보여줘
    op, threshold, op_label = _current_table_numeric_filter_op(t)

    numeric_filter_hit = (
        bool(op)
        and any(w in t for w in ("수량", "재고수량", "출고수량", "입고수량"))
    )

    if numeric_filter_hit:
        if "재고수량" in t:
            qty_col = _current_table_find_col(
                df,
                exact=("재고수량", "현재재고수량", "기말재고수량", "실재고수량", "장부재고수량"),
                include_any=("재고수량", "현재재고수량"),
                exclude_any=("금액", "단가", "보험", "코드"),
            )
            qty_label = "재고수량"
        elif "출고수량" in t:
            qty_col = _current_table_find_col(
                df,
                exact=("출고수량",),
                include_any=("출고수량",),
                exclude_any=("금액", "단가", "코드"),
            )
            qty_label = "출고수량"
        elif "입고수량" in t:
            qty_col = _current_table_find_col(
                df,
                exact=("입고수량",),
                include_any=("입고수량",),
                exclude_any=("금액", "단가", "코드"),
            )
            qty_label = "입고수량"
        else:
            qty_col = _current_table_find_col(
                df,
                exact=("수량", "재고수량", "현재재고수량", "출고수량", "입고수량"),
                include_any=("수량",),
                exclude_any=("금액", "단가", "코드"),
            )
            qty_label = "수량"

        threshold_text = (
            f"{int(threshold)}"
            if abs(threshold - int(threshold)) < 1e-9
            else f"{threshold:g}"
        )

        title = f"현재표 {qty_label} {threshold_text} {op_label} 목록"
        query_summary = f"현재표 / {qty_label} {threshold_text} {op_label}"

        if not qty_col:
            cols_preview = ", ".join(str(c) for c in list(df.columns)[:25])
            msg = (
                f"현재표에는 '{qty_label}'로 판단할 수 있는 컬럼이 없습니다.\n\n"
                f"현재표 기준 행수: {len(df):,}건\n"
                f"현재표 주요 컬럼: {cols_preview}\n\n"
                "참고: 실재고월집계/장부재고월집계는 보통 재고수량 컬럼이 아니라 "
                "입고수량, 출고수량, 입고공급가액, 출고공급가액 중심의 월 집계표입니다.\n"
                "재고수량 기준 목록을 보려면 [제품재고현황 조회]를 먼저 실행한 뒤 다시 질문하세요.\n"
                "월집계 현재표에서는 예를 들어 '현재표에서 출고수량이 0 이하인 제품 표로 보여줘'처럼 조회할 수 있습니다."
            )

            return _current_table_push_notice(
                title=title,
                action=title,
                message=msg,
                query_summary=query_summary,
                source_query=t,
            )

        nums = _current_table_to_num(df[qty_col])

        if op == ">=":
            mask = nums >= threshold
        elif op == ">":
            mask = nums > threshold
        elif op == "<=":
            mask = nums <= threshold
        elif op == "<":
            mask = nums < threshold
        elif op == "==":
            mask = nums == threshold
        else:
            return False

        out = df.loc[mask].copy()

        # 조건 기준 컬럼은 화면에서도 숫자로 보이게 보정한다.
        # 제품재고현황은 0을 빈값처럼 표시하는 컬럼이 있을 수 있으므로,
        # 필터 기준 컬럼만큼은 0/음수/양수가 직접 보이게 한다.
        try:
            out[qty_col] = nums.loc[out.index].fillna(0).values
            out[qty_col] = pd.to_numeric(out[qty_col], errors="coerce").fillna(0)
        except Exception:
            pass

        # 원본표의 별칭 컬럼 값 보존 후 None/nan 표시 정리
        out = _current_table_fill_alias_values(out)
        out = _current_table_clean_none_for_display(out)

        # 제품 관련 컬럼을 앞쪽으로 보기 좋게 정렬
        front_candidates = [
            "제품코드",
            "제품명",
            "규격",
            "제조사명",
            "제품그룹명",
            "제품구분명",
            "제품분류명",
            "재고년월",
            "재고위치",
            "재고위치명",
            qty_col,
        ]

        front_cols = []
        for c in front_candidates:
            if c in out.columns and c not in front_cols:
                front_cols.append(c)

        rest_cols = [c for c in out.columns if c not in front_cols]
        if front_cols:
            out = out[front_cols + rest_cols]

        if "순번" not in out.columns:
            out.insert(0, "순번", range(1, len(out) + 1))

        log.info(
            "[chat.followup_table] numeric filter table built col=%s op=%s threshold=%s source_rows=%s rows=%s table_key=%s",
            qty_col,
            op,
            threshold,
            len(df),
            len(out),
            table_key,
        )

        return _current_table_push_table(
            title=title,
            action=title,
            df=out,
            query_summary=query_summary,
            source_query=t,
            source_table_key=table_key,
            source_rows=len(df),
            display_limit=500,
        )

def is_sims_related_question(text: str) -> bool:
    t = (text or "").strip()
    if not t:
        return False

    # 최신 SIMS 조회 결과에 대한 후속 분석/요약 질문도
    # LLM에는 SIMS_JSON을 붙여야 하므로 SIMS 관련 질문으로 본다.
    if is_sims_result_followup_question(t):
        return True

    if _SIMS_HINT_RE.search(t):
        return True

    # 최신 SIMS 조회 결과에 대한 후속 질문 보강
    # 예: "이 결과 분석해줘", "현재 조회 결과 요약", "방금 표에서 문제점 찾아줘"
    # 기존 정규식은 제품/거래처/재고 같은 명시 단어가 없으면 일반 질문으로 보아
    # SIMS_JSON을 붙이지 않는 경우가 있었다.
    try:
        has_sims_ctx = bool(
            st.session_state.get("__sims_ctx")
            or st.session_state.get("__sims_context")
            or st.session_state.get("__sims_context_obj")
        )
    except Exception:
        has_sims_ctx = False

    followup_words = (
        "이 결과", "현재 결과", "조회 결과", "현재 조회", "방금", "위 결과",
        "이 표", "현재 표", "현재 데이터", "컨텍스트", "sims 결과", "SIMS 결과",
    )
    analysis_words = (
        "분석", "요약", "정리", "해석", "설명", "인사이트", "특징",
        "문제점", "이상", "위험", "주의", "상위", "하위", "합계",
        "평균", "건수", "분포", "추세", "예상", "부족", "비교",
    )

    if has_sims_ctx and any(w in t for w in followup_words) and any(w in t for w in analysis_words):
        return True

    # 아주 짧은 후속 명령도 허용: "분석해줘", "요약해줘", "정리해줘"
    if has_sims_ctx and re.fullmatch(
        r"(분석|요약|정리|해석|설명)(해줘|해 주세요|해주시겠어요|해봐|해 봐)?[.!? ]*",
        t,
    ):
        return True

    # 거래처 마스터 라벨 단독 질의 보강
    master_label_words = (
        "대표자",
        "대표자명",
        "사업자",
        "사업자번호",
        "사업자등록번호",
        "영업사원",
        "영업사원명",
        "단가적용처",
        "단가적용처명",
        "재고적용처",
        "재고적용처명",
        "거래처명",
        "거래처코드",
    )

    master_verbs = (
        "조회",
        "검색",
        "목록",
        "찾아",
        "찾아줘",
        "보여줘",
        "알려줘",
        "있어",
        "있나",
        "확인",
    )

    if any(w in t for w in master_label_words) and any(v in t for v in master_verbs):
        return True

    return False



def _is_pending_product_pick_text(user_input: str) -> bool:
    pending = st.session_state.get("__io_pending_product_pick") or {}
    if not isinstance(pending, dict) or not pending.get("candidates"):
        return False

    t = (user_input or "").strip()
    if not t:
        return False

    # 예: 1번 / 2번으로 조회 / 3번 선택
    temp = t.replace("조회", "").replace("선택", "").replace("으로", "").strip()
    if temp.endswith("번"):
        num = temp[:-1].strip()
        if num.isdigit():
            return True

    # 예: 첫번째 제품 / 첫 번째 / 두번째
    if any(k in t for k in (
        "첫번째", "첫 번째",
        "두번째", "두 번째",
        "세번째", "세 번째",
        "네번째", "네 번째",
        "다섯번째", "다섯 번째",
    )):
        return True

    return False


MAX_MSG_CHARS_TO_MODEL = 20_000
def _clip_for_model(s: str, limit: int = MAX_MSG_CHARS_TO_MODEL) -> str:
    s = s or ""
    return s if len(s) <= limit else (s[:limit] + "\n\n...(길이 제한으로 이후 생략)")

#def call_chat_with_retry(*, messages, model=None, temperature=0.2, stream=False):
#    """
#    OpenAI 호환 API 호출을 타임아웃/재시도로 감싸는 헬퍼.
#    - stream=True  -> streaming generator 반환
#    - stream=False -> 완료 응답 객체 반환
#    """
#    # per-call timeout 부여 (with_options 없으면 원본 CLIENT 그대로 사용)
#    cli = getattr(CLIENT, "with_options", lambda **kw: CLIENT)(timeout=LLM_TIMEOUT_S)
#
#    last_err = None
#    total_tries = 1 + int(LLM_MAX_RETRY)
#    for attempt in range(total_tries):
#        try:
#            return cli.chat.completions.create(
#                model=model or st.session_state.get("selected_model") or "local-model",
#                messages=messages,
#                temperature=temperature,
#                stream=stream,
#                timeout=LLM_TIMEOUT_S,  # ← 이중 안전장치 (드라이버별 호환)
#           )
#        except Exception as e:
#            last_err = e
#            if attempt >= total_tries - 1:
#                raise RuntimeError(f"LLM 호출 실패(최종): {type(e).__name__}: {e}") from e
#            # 백오프(시퀀스 + 지터)
#            sleep_s = LLM_BACKOFF_SEQ[attempt] if attempt < len(LLM_BACKOFF_SEQ) else LLM_BACKOFF_SEQ[-1]
#            time.sleep(sleep_s + random.uniform(0, 0.3))

# per-call 오버라이드 + 회로 차단기 적용 래퍼
def call_chat_protected(
    *, 
    messages, 
    model: str | None = None, 
    temperature: float = 0.2, 
    stream: bool = False,
    # ⬇ per-call 오버라이드
    timeout_s: int | None = None, 
    max_retry: int | None = None, 
    backoff_seq: list[float] | None = None,
    allow_when_open: bool = False,   # 회로차단 open 시에도 시도할지 (half-open 트리거용)
):
    """
    OpenAI 호환 API 호출을 '타임아웃/재시도/회로차단'으로 감싼 보호 래퍼.
    - stream=True  -> streaming generator 반환
    - stream=False -> 완료 응답 객체 반환
    - timeout_s / max_retry / backoff_seq 로 호출별 오버라이드 가능
    """

    # ── 0) 회로차단 게이트 ──────────────────────────────────────────
    _cb_init()
    if _cb_is_open():
        # open 유지시간 끝났다면 half-open으로 1회만 시험 통과/실패 여부 확인
        if allow_when_open:
            st.session_state["__cb_state"] = "half_open"
        else:
            raise RuntimeError("LLM 호출이 잠시 중단되었습니다(회로차단). 잠시 후 다시 시도해 주세요.")

    # ── 1) per-call 파라미터 결합 ───────────────────────────────────
    local_timeout   = int(timeout_s if timeout_s is not None else LLM_TIMEOUT_S)
    local_max_retry = int(max_retry if max_retry is not None else LLM_MAX_RETRY)
    local_backoffs  = list(backoff_seq) if backoff_seq is not None else list(LLM_BACKOFF_SEQ)

    # ── 2) per-call 타임아웃을 클라이언트에 적용 ───────────────────
    cli = getattr(CLIENT, "with_options", lambda **kw: CLIENT)(timeout=local_timeout)


    # ── 회로차단 게이트 + 재시도 호출 ─────────────────────────────
    _cb_init()  # 기본 상태값 보장

    now = time.time()
    state = st.session_state["__cb_state"]

    # open 상태: 쿨다운 남아 있으면 막고, 만료되었으면 half_open으로 전환
    if state == "open":
        if now < st.session_state.get("__cb_open_until", 0.0):
            raise RuntimeError("회로차단: 열림 상태입니다. 잠시 후 다시 시도하세요.")
        st.session_state["__cb_state"] = "half_open"

    try:
        resp = call_chat_with_retry(
            messages=messages,
            model=model or (st.session_state.get("selected_model") or cfg_str("LMSTUDIO_MODEL", "local-model")),
            temperature=temperature,
            stream=stream,
            timeout_s=local_timeout,
            max_retry=local_max_retry,
            backoff_seq=local_backoffs,
        )
        _cb_on_success()   # 성공 → 카운터 리셋 & closed
        return resp

    except Exception as e:
        _cb_on_failure()   # 실패 → 실패 카운트 증가 / 필요 시 trip(open)
        raise RuntimeError(f"LLM 호출 실패(최종): {type(e).__name__}: {e}") from e
# =========================================================
# 공백제거 + (기본값: 행 수는 자르지 않음)
def _compact_records_for_model(
    records: list[dict],
    max_rows: int | None = None,
    max_str_len: int = 120,
    drop_empty: bool = True,
) -> list[dict]:
    """
    LLM에 보낼 records 정리:
    - 기본값: max_rows=None 이면 행 수 제한 없이 전체 사용
    - max_rows가 지정되면 그 개수까지만 사용
    - 문자열 값은 strip() 해서 불필요한 공백 제거
    - max_str_len: 너무 긴 문자열은 컷(주소/비고로 인한 토큰 폭발 방지)
    - drop_empty: None/"" 값은 키 자체를 제거하여 JSON 크기 감소    
    """
    if not isinstance(records, list):
        return []

    # max_rows가 지정되어 있으면 잘라 쓰고, 아니면 전체 사용
    if max_rows is not None and max_rows > 0:
        iter_rows = records[:max_rows]
    else:
        iter_rows = records

    trimmed: list[dict] = []
    for row in iter_rows:
        if not isinstance(row, dict):
            continue
        new_row: dict[str, Any] = {}
        for k, v in row.items():
            if isinstance(v, str):
                # CHAR(10) 같은 패딩 제거
                v = v.strip()
                # ✅ 너무 긴 문자열은 컷 (주소/비고/긴 메모 때문에 토큰 폭발 방지)
                if max_str_len and len(v) > max_str_len:
                    v = v[:max_str_len] + "…"
            # ✅ 빈값 제거(키 자체 제거) → JSON 크기 크게 감소
            if drop_empty and (v is None or v == "" or v == "None"):
                continue
            new_row[str(k)] = v
        trimmed.append(new_row)
    return trimmed

# ---------------------------------------------------------------------
# (D) 컨텍스트 메타 질문 즉답(SSOT=chat_bridge)
# ---------------------------------------------------------------------
#def _try_answer_ctx_meta_question(
#    user_text: str,
#    *,
#    room: dict,
#    make_ts,
#    next_seq,
#    logger,
#    max_age_sec: int = 900,
#) -> bool:
#    """
#    '컨텍스트/현재 데이터/행 수/컬럼' 같은 질문을 LLM 없이 즉답한다.
#    - SSOT는 get_sims_context_data/get_sims_context_text(=chat_bridge 계열)로 읽는다.
#    - 답변은 room.messages에 assistant 메시지로 저장한다.
#    """
#    import uuid
#    import re
#
#    txt = (user_text or "").strip()
#    if not txt:
#        return False
#
#    # 메타성 질문만 가로챔(너무 광범위하게 잡지 않기)
#    meta_pat = r"(컨텍스트|현재\s*데이터|지금\s*데이터|행\s*수|몇\s*건|총\s*몇|컬럼|열\s*목록|필드|스키마|meta|records|aggregations|집계)"
#    if not re.search(meta_pat, txt, flags=re.IGNORECASE):
#        return False
#
#    # SSOT: chat_bridge
#    try:
#        sims_data = get_sims_context_data(max_age_sec=max_age_sec)
#    except TypeError:
#        sims_data = get_sims_context_data()
#    except Exception:
#        logger.exception("[ctx.meta] get_sims_context_data failed")
#        sims_data = None
#
#    try:
#        try:
#            sims_ctx_text = get_sims_context_text(max_age_sec=max_age_sec)
#        except TypeError:
#            sims_ctx_text = get_sims_context_text()
#    except Exception:
#        sims_ctx_text = None
#
#    if not isinstance(sims_data, dict) or not sims_data:
#        # 텍스트 컨텍스트라도 있으면 안내
#        if isinstance(sims_ctx_text, str) and sims_ctx_text.strip():
#            msg = "현재 SIMS 컨텍스트는 '텍스트' 형태로만 존재합니다. (JSON data_container 없음)\n- 더 구체한 분석은 SIMS 조회(표 생성) 후 질문해 주세요."
#        else:
#            msg = "현재 SIMS 컨텍스트가 없습니다. 먼저 SIMS 조회를 실행해 표(결과)를 만든 뒤 질문해 주세요."
#
#        room.setdefault("messages", []).append({
#            "id": str(uuid.uuid4()),
#            "role": "assistant",
#            "content": msg,
#            "time": make_ts(),
#            "seq": next_seq(),
#        })
#        return True
#
#    # chat_middleware SSOT 호환: ctx_pack이면 data를 풀어서 사용
#    pack = sims_data
#    if isinstance(pack.get("data"), dict):
#        data = pack.get("data") or {}
#        action = pack.get("action")
#        ts = pack.get("ts")
#        params = pack.get("params")
#    else:
#        data = pack
#        action = None
#        ts = None
#        params = None
#
#    meta = data.get("meta") or {}
#    cols = data.get("columns") or []
#    records = data.get("records") or []
#    aggs = data.get("aggregations") or {}
#
#    # row_count_total 우선, 없으면 row_count, 그마저 없으면 records 길이
#    total = (
#        meta.get("row_count_total")
#        or meta.get("row_count")
#        or len(records)
#        or 0
#    )
#
#    # 액션/ts/params를 meta에 없으면 pack에서 보강
#    action_eff = meta.get("action") or action or "(unknown)"
#    ts_eff = meta.get("ts") or ts
#    params_eff = meta.get("params") or params
#
#    wants_cols = bool(re.search(r"(컬럼|열\s*목록|필드|스키마)", txt))
#    wants_aggs = bool(re.search(r"(집계|aggregations|요약)", txt, flags=re.IGNORECASE))
#
#    cols_preview = ", ".join([str(c) for c in cols[:30]])
#    if len(cols) > 30:
#        cols_preview += f" ... (+{len(cols)-30}개)"
#
#    lines = []
#    lines.append("현재 SIMS 컨텍스트 요약")
#    lines.append(f"- action: {action_eff}")
#    lines.append(f"- rows: {total}")
#    lines.append(f"- cols: {len(cols)}")
#    if ts_eff:
#        lines.append(f"- ts: {ts_eff}")
#    if params_eff:
#        lines.append(f"- params: {params_eff}")
#
#    if wants_cols:
#        lines.append(f"- columns: {cols_preview if cols else '(없음)'}")
#    if wants_aggs:
#        # aggs는 너무 길 수 있으니 키만 요약
#        if isinstance(aggs, dict) and aggs:
#            keys = ", ".join(list(aggs.keys())[:20])
#            tail = "" if len(aggs) <= 20 else f" ... (+{len(aggs)-20}개)"
#            lines.append(f"- aggregations keys: {keys}{tail}")
#        else:
#            lines.append("- aggregations: (없음)")
#
#    msg = "\n".join(lines)
#
#    room.setdefault("messages", []).append({
#        "id": str(uuid.uuid4()),
#        "role": "assistant",
#        "content": msg,
#        "time": make_ts(),
#        "seq": next_seq(),
#    })
#    return True


# 시스템 프롬프트 포함 메시지 빌더
# - SIMS 관련 질문이면 SIMS 전용 시스템 프롬프트, 아니면 일반 시스템 프롬프트 사용
# - SIMS 관련 질문이면서 SIMS 컨텍스트가 있으면 SIMS_JSON 블록으로 주입, 없으면 텍스트 컨텍스트 주입
# - 규칙 문구는 최소화하고, JSON 데이터 자체를 보고 추론하도록 유도
# - SIMS 관련 질문 판단은 is_sims_related_question() 사용
def build_messages_with_system(
    history_msgs: list[dict],
    system_prompt: str | None = None,
    user_text: str | None = None,
) -> list[dict]:
    """
    대화 히스토리 + (선택) 시스템 프롬프트 + (선택) 최신 SIMS 컨텍스트를 결합해 messages를 만든다.
    - SIMS JSON이 있으면 [SIMS_JSON] 블록으로, 없으면 텍스트 컨텍스트를 주입.
    - 규칙 문구는 최소화하고, JSON 데이터 자체를 보고 추론하도록 유도한다.
    """
    # ✅ 질문이 SIMS 관련이면 SIMS 시스템프롬프트, 아니면 일반 시스템프롬프트
    attach_sims = is_sims_related_question(user_text or "")
    base_system = system_prompt or (BASE_SYSTEM_PROMPT if attach_sims else GENERAL_SYSTEM_PROMPT)

    msgs: list[dict] = [{"role": "system", "content": base_system}]

    # 1) 최신 SIMS 컨텍스트 조회
    sims_block: str | None = None

    sims_context_kind = "none"

    try:
        sims_data = get_sims_context_data(max_age_sec=900)
        # ✅ chat_bridge/chat_middleware 호환: max_age_sec 인자 유무 모두 지원
        try:
            sims_ctx = get_sims_context_text(max_age_sec=900)
        except TypeError:
            sims_ctx = get_sims_context_text()
        try:
            log.debug(
                "[chat] build_messages_with_system: sims_data_type=%s, has_ctx=%s",
                type(sims_data).__name__,
                bool(sims_ctx and sims_ctx.strip()),
            )
        except Exception:
            pass
    except Exception:
        # ❗ 여기서 예외가 나면 SIMS 쪽 전부 죽어버리니, 로그 남기고 None 처리
        log.exception("[chat] get_sims_context_xxx failed")
        sims_data = None
        sims_ctx = None

    # 2) ✅ SIMS 관련 질문일 때만 SIMS JSON/CONTEXT를 주입
    analysis_ctx = st.session_state.get("__sims_analysis_ctx")

    def _wants_current_table_source_ctx(q: str) -> bool:
        """
        '현재표 ...별/집계/분석'은 최신 파생표가 아니라
        원본 현재표(source DF) 기준으로 LLM이 보게 한다.

        단, '방금 표', '위 표', '이 표'는 화면에 바로 보이는 최신 파생표를 뜻할 수 있으므로 제외한다.
        """
        s = str(q or "").strip()
        compact = s.replace(" ", "")

        if not any(w in compact for w in ("현재표", "현재조회결과", "현재결과")):
            return False

        return any(
            w in compact
            for w in (
                "제품별",
                "품목별",
                "거래처별",
                "재고위치별",
                "영업사원별",
                "담당자별",

                # 금액/수량/매출 계열
                "매출금액",
                "매출액",
                "매출",
                "출고수량",
                "수량",
                "공급가액",
                "세액",
                "합계금액",

                # 일자/요일/기간 분석 계열
                "일자",
                "날짜",
                "요일",
                "일별",
                "일자별",
                "월별",
                "연도별",
                "기간별",

                # 순위/최대/최소 계열
                "가장많은",
                "가장큰",
                "최대",
                "최고",
                "많은",
                "큰",
                "TOP",
                "top",
                "상위",

                # 일반 분석 계열
                "집계",
                "분석",
                "요약",
            )
        )
    try:
        source_analysis_ctx = st.session_state.get("__sims_current_table_source_analysis_ctx")
        latest_is_current_followup = bool(
            isinstance(analysis_ctx, dict)
            and analysis_ctx.get("current_table_followup")
        )

        if (
            attach_sims
            and _wants_current_table_source_ctx(user_text or "")
            and latest_is_current_followup
            and isinstance(source_analysis_ctx, dict)
            and source_analysis_ctx.get("kind") == "SIMS_ANALYSIS_CONTEXT_V1"
        ):
            log.debug(
                "[chat] use source analysis ctx for current-table LLM followup "
                "latest_action=%s latest_rows=%s source_action=%s source_rows=%s",
                analysis_ctx.get("action"),
                analysis_ctx.get("row_count"),
                source_analysis_ctx.get("action"),
                source_analysis_ctx.get("row_count"),
            )
            analysis_ctx = source_analysis_ctx
    except Exception:
        log.exception("[chat] source analysis ctx selection failed")

    if attach_sims and isinstance(analysis_ctx, dict) and analysis_ctx.get("kind") == "SIMS_ANALYSIS_CONTEXT_V1":
        try:
            try:
                _ctx_len = int(st.session_state.get("__model_ctx_len") or 16384)
            except Exception:
                _ctx_len = 16384

            if _ctx_len <= 12288:
                SIMS_JSON_CHAR_LIMIT = 9000
            elif _ctx_len <= 16384:
                SIMS_JSON_CHAR_LIMIT = 12000
            elif _ctx_len <= 32768:
                SIMS_JSON_CHAR_LIMIT = 18000
            else:
                SIMS_JSON_CHAR_LIMIT = 24000

            def _shrink_analysis_ctx_for_limit(ctx: dict, limit: int) -> dict:
                """분석 컨텍스트는 JSON 유효성을 유지하면서 대표 목록만 단계적으로 축소한다."""
                base = dict(ctx)

                def _with_limits(risk_n: int, grade_n: int) -> dict:
                    out = dict(base)
                    if isinstance(out.get("risk_products_top"), list):
                        out["risk_products_top"] = out["risk_products_top"][:risk_n]
                    gs = out.get("grade_samples")
                    if isinstance(gs, dict):
                        out["grade_samples"] = {
                            str(k): (v[:grade_n] if isinstance(v, list) else v)
                            for k, v in gs.items()
                        }
                    return out

                for risk_n, grade_n in ((50, 8), (40, 6), (30, 5), (25, 4), (20, 3), (15, 2), (10, 1)):
                    cand = _with_limits(risk_n, grade_n)
                    raw = json.dumps(cand, ensure_ascii=False, default=str)
                    if len(raw) <= limit:
                        return cand
                return _with_limits(8, 1)

            llm_sims_data = _shrink_analysis_ctx_for_limit(analysis_ctx, SIMS_JSON_CHAR_LIMIT)
            sims_json = json.dumps(llm_sims_data, ensure_ascii=False, default=str)
            ctx_txt = _clip_for_model(sims_json, limit=SIMS_JSON_CHAR_LIMIT)

            log_cols = int(llm_sims_data.get("column_count") or len(llm_sims_data.get("columns") or []))
            log_rows = int(llm_sims_data.get("row_count") or 0)
            log_risk = len(llm_sims_data.get("risk_products_top") or [])

            meta_hint = (
                f"[SIMS_META] kind=analysis action={llm_sims_data.get('action')} "
                f"row_count={log_rows} cols={log_cols} risk_top={log_risk}"
            )

            sims_block = (
                meta_hint + "\n"
                "다음 SIMS_ANALYSIS_CONTEXT_V1 JSON만 보고 질문에 답하세요.\n"
                "\n"
                "- 이 JSON은 최신 SIMS 조회 결과 전체를 Python으로 집계한 분석 컨텍스트입니다.\n"
                "- summary와 *_counts는 전체 결과 기준입니다. 절대 '샘플 데이터'라고 말하지 마세요.\n"
                "- 일자/날짜/요일/일별/월별 매출 질문은 sales_time_profile의 daily_sales_top, monthly_sales, weekday_sales를 우선 사용하세요.\n"
                "- 제품별/거래처별/영업사원별/수량 TOP 질문은 sales_group_profile을 우선 사용하세요.\n"
                "- risk_products_top과 grade_samples는 품목 설명용 대표 목록입니다.\n"
                "- 이전 표, 이전 조회 결과, 이전 답변은 분석 근거로 사용하지 마세요.\n"
                "- 사용자가 특정 항목/목록/값을 물으면 그 질문에만 직접 답하세요. 이때 핵심 요약, 주의/확인할 점, 다음 조회 제안은 생략하세요.\n"
                "- 사용자가 요약/분석을 요청한 경우에만 ①핵심 요약 ②주요 수치 ③주의/확인할 점 ④다음 조회 제안 순서로 짧게 정리하세요.\n"
                "- 사용자가 '표로'라고 요청하면 주요 수치나 목록은 표로 답하세요.\n"
                "- 내부 key 이름(low_stock_records, zero_stock_records, sample_records, risk_products_top, grade_samples 등)은 답변에 노출하지 마세요.\n"


                "- stock_mode 또는 stock_label 값이 '실수불'이면 반드시 '실수불 기준'이라고 그대로 표현하고, '실재고 기준'으로 바꾸지 마세요.\n"
                "- 답변 본문에 '샘플', '샘플 데이터', 'sample_rows', '대표 목록'이라는 표현을 쓰지 말고, 필요한 경우 '전체 조회 결과 기준'이라고만 표현하세요.\n"
                f"[SIMS_JSON]\n{ctx_txt}\n[/SIMS_JSON]"
            )

            log.debug(
                "[chat] build_messages_with_system → SIMS_JSON attached "
                "(len=%s, kind=analysis, cols=%s, rows=%s, risk_top=%s)",
                len(ctx_txt),
                log_cols,
                log_rows,
                log_risk,
            )
        except Exception:
            log.exception("[chat] build_messages_with_system → analysis ctx attach failed")
            sims_block = None

    elif attach_sims and isinstance(sims_data, dict) and sims_data:
        # ✅ chat_middleware SSOT 호환: ctx_pack이면 data를 풀어서 사용
        # ctx_pack = {"text":..., "data":{columns/records/meta}, "action":..., "params":..., "ts":...}
        if isinstance(sims_data.get("data"), dict):
            pack = sims_data
            sims_data = pack["data"]  # 여기서부터는 columns/records/meta가 있는 "data_container"로 처리
            # meta에 action/ts/params도 같이 넣어두면 LLM이 맥락 이해에 도움됨(없으면 자동 생성)
            try:
                meta0 = sims_data.get("meta", {}) or {}
                if isinstance(meta0, dict):
                    meta0 = dict(meta0)
                    meta0.setdefault("action", pack.get("action"))
                    meta0.setdefault("ts", pack.get("ts"))
                    meta0.setdefault("params", pack.get("params"))
                    sims_data["meta"] = meta0
            except Exception:
                pass
        try:
            cols = sims_data.get("columns", [])
            records = sims_data.get("records", [])
            # 원본 meta 를 건드리지 않도록 복사
            meta_orig = sims_data.get("meta", {}) or {}
            meta = dict(meta_orig)
            aggs = sims_data.get("aggregations", {}) or {}

            # 전체 행 수 계산 & meta 기본값 보정
            total_rows = int(
                meta.get("row_count_total")
                or meta.get("row_count")
                or len(records)
                or 0
            )
            meta.setdefault("row_count_total", total_rows)
            meta.setdefault("row_count", total_rows)

            # ✅ 액션명 기준으로, 얼마나 많이 records 를 보낼지 결정
            action_name = (meta.get("action") or "").strip()

            def _pick_llm_columns(action_name: str, all_cols: list[str]) -> list[str]:
                """
                LLM 분석용 컬럼 축소.

                전체 표에는 90개 이상의 컬럼이 있어도,
                LLM 분석에는 핵심 컬럼만 넘긴다.
                원본 표 표시에는 영향 없음.
                """
                colset = set(map(str, all_cols))
                action = str(action_name or "")

                if "재고부족" in action:
                    preferred = [
                        "순번",
                        "제품코드",
                        "제품명",
                        "규격",
                        "제조사명",
                        "제품그룹명",
                        "제품구분명",
                        "제품분류명",
                        "재고기준",
                        "현재재고수량",
                        "현재재고금액",
                        "최근3개월평균수량",
                        "최근6개월평균수량",
                        "월평균출고수량",
                        "예상기준월수량",
                        "재고커버월수",
                        "1개월부족수량",
                        "2개월부족수량",
                        "3개월부족수량",
                        "부족등급",
                    ]
                    return [c for c in preferred if c in colset]

                if "매출 예상" in action or "매출예상" in action:
                    preferred = [
                        "순번",
                        "제품코드",
                        "제품명",
                        "제조사명",
                        "제품그룹명",
                        "제품구분명",
                        "총매출액",
                        "월평균매출",
                        "최근3개월평균매출",
                        "최근6개월평균매출",
                        "최근3개월증감률",
                        "예상매출",
                        "예상등급",
                    ]
                    return [c for c in preferred if c in colset]

                if "매출 추세" in action or "추세" in action:
                    preferred = [
                        "순번",
                        "제품코드",
                        "제품명",
                        "제조사명",
                        "제품그룹명",
                        "제품구분명",
                        "매출년월",
                        "매출액",
                        "총매출액",
                        "월평균매출",
                        "최근3개월평균매출",
                        "최근6개월평균매출",
                        "추세판정",
                    ]
                    return [c for c in preferred if c in colset]

                return []

            # ── 중요 포인트 ──────────────────────────────
            # 사용자 목록(378행) 같이 적당히 작은 데이터는
            #  → LLM 에게 "전 행"을 보여줘서 검색이 가능하게 함.
            # 그 외 큰 테이블(입출고 등)은 80행 정도만 샘플로 유지.
            # ────────────────────────────────────────────

            # ✅ 툴콜 붙이기 전 임시 운영:
            max_rows = min(len(records), 300)

            # ✅ 컨텍스트 길이 예산
            #  - 16K(16384)에서는 45,000 chars면 오버플로우가 자주 남.
            #  - 로그 기준 len≈43,530 에서 overflow 발생했으므로, 16K는 18k~22k 수준이 안전.
            #  - ctx_len을 얻을 수 있으면 그 값에 맞춰 동적으로 제한.
            try:
                _ctx_len = int(st.session_state.get("__model_ctx_len") or 16384)
            except Exception:
                _ctx_len = 16384

            # Gemma 3 27B를 12K context로 쓰는 경우,
            # JSON 20K chars도 n_keep 초과가 발생한다.
            # SIMS_JSON은 summary/meta/counts/핵심 컬럼 중심으로 줄인다.
            if _ctx_len <= 12288:
                SIMS_JSON_CHAR_LIMIT = 7000
            elif _ctx_len <= 16384:
                SIMS_JSON_CHAR_LIMIT = 9000
            elif _ctx_len <= 32768:
                SIMS_JSON_CHAR_LIMIT = 16000
            else:
                SIMS_JSON_CHAR_LIMIT = 24000

            # ✅ 마지막으로 만든 sample_rows를 기록( nonlocal 대신 dict 사용 )
            _last = {"rows": 0}

            llm_cols = _pick_llm_columns(action_name, cols)

            def _build_json_with_rows(n: int) -> str:
                raw_trimmed = _compact_records_for_model(
                    records, max_rows=n, max_str_len=80, drop_empty=True
                )

                # 분석/KPI 결과는 LLM에 핵심 컬럼만 전달한다.
                # 원본 표 렌더링에는 영향이 없다.
                if llm_cols:
                    trimmed = []
                    for r in raw_trimmed:
                        rr = {}
                        for c in llm_cols:
                            if c in r and r.get(c) not in ("", None):
                                rr[c] = r.get(c)
                        trimmed.append(rr)
                else:
                    trimmed = raw_trimmed

                _last["rows"] = len(trimmed)

                meta2 = dict(meta)
                meta2["sample_rows"] = len(trimmed)

                if llm_cols:
                    meta2["llm_columns_limited"] = "Y"
                    meta2["original_column_count"] = len(cols)
                    meta2["llm_column_count"] = len(llm_cols)

                container: Dict[str, Any] = {}
                if meta2:
                    container["meta"] = meta2
                if aggs:
                    container["aggregations"] = aggs

                container["columns"] = llm_cols or cols
                container["records"] = trimmed

                return json.dumps(container, ensure_ascii=False, default=str)

            sims_json = _build_json_with_rows(max_rows)

            # ✅ 그래도 크면 행 수를 줄여서 맞춘다 (300 → 200 → 150 → 120 → 100 → 80 → 60 ...)
            if len(sims_json) > SIMS_JSON_CHAR_LIMIT:
                for n in (200, 150, 120, 100, 80, 60, 40, 30, 20, 10, 5, 3):
                    sims_json = _build_json_with_rows(min(n, len(records)))
                    if len(sims_json) <= SIMS_JSON_CHAR_LIMIT:
                        break

            # (여기서 meta.sample_rows는 _build_json_with_rows에서 이미 세팅됨)

        except Exception:
            # 준비 과정에서 문제가 나면, 마지막 수단으로 원본 sims_data 를 그대로 직렬화
            log.exception("[chat] build_messages_with_system → SIMS JSON prepare failed")
            try:
                sims_json = json.dumps(sims_data, ensure_ascii=False, default=str)
            except Exception:
                log.exception("[chat] json.dumps(sims_data) failed (fallback)")
                sims_json = "{}"

        # ✅ 길이 제한 적용 (meta/aggregations를 앞에 두었기 때문에 잘려도 핵심 정보는 살아남음)
        #ctx_txt = _clip_for_model(sims_json, limit=MAX_MSG_CHARS_TO_MODEL)
        #    → 사용자 300~400행 수준에서는 통째로 보내도 여유가 있으므로
        #      일단 SIMS_JSON은 자르지 않고 그대로 보낸다.
        # ✅ 최후 안전장치: 모델 컨텍스트 초과를 막기 위해 마지막에 clip 적용
        # (meta.sample_rows가 들어가 있으므로, 일부 잘리더라도 "샘플"임을 모델이 인지 가능)
        ctx_txt = _clip_for_model(sims_json, limit=SIMS_JSON_CHAR_LIMIT)

        
#        sims_block = (
#            "다음 SIMS ERP JSON을 보고 질문에 사실만 정확히 요약해 답변하세요.\n"
#            "- meta/columns/records 값을 기준으로 판단하세요.\n"
#            "- JSON에 존재하지 않는 필드는 지어내지 마세요.\n"
#            "- records는 일부 샘플일 수 있으므로, 전체 개수는 meta.row_count_total을 우선 사용하세요.\n"
#            "- 필요한 경우 records를 직접 분석하여 집계하세요.\n"
#            "- 답변은 간결한 표나 문장 형태로 짧게 정리하세요.\n"
#            f"[SIMS_JSON]\n{ctx_txt}\n[/SIMS_JSON]"
#        )
        # ✅ LLM이 빨리 핵심을 잡도록 meta 요약을 JSON 바깥에 한 줄로 제공
        meta_hint = f"[SIMS_META] action={meta.get('action')} row_count_total={meta.get('row_count_total')} sample_rows={_last.get('rows')} cols={len(cols) if isinstance(cols, list) else ''}"

        sims_block = (
            meta_hint + "\n" +
            "다음 SIMS ERP JSON만 보고 질문에 답하세요.\n"
            "\n"
            "- 전체 개수는 meta.row_count_total → 없으면 meta.row_count 를 사용하세요.\n"
            "- 질문이 '총 몇 개/몇 명/개수'면 meta.row_count_total(또는 meta.row_count)을 **먼저** 답하세요.\n"
            "- 질문이 \"영업사원 몇명\"(또는 유사) 이면 aggregations.sales_reps.distinct_count 를 먼저 사용하세요.\n"
            "- \"거래처종류/등급/그룹\" 관련 질문은 aggregations.vendor_kind / vendor_rank / vendor_group 을 우선 사용하세요.\n"
            "- '~상세내역', '~목록', '~모두 보여줘' 같은 질문은 조건에 맞는 records 를 표로 보여주고,\n"
            "  필요하면 한두 문장으로만 짧게 설명하세요.\n"
            f"[SIMS_JSON]\n{ctx_txt}\n[/SIMS_JSON]"
        )

        try:
            log.debug(
                "[chat] build_messages_with_system → SIMS_JSON attached (len=%s, cols=%s, rows=%s)",
                len(ctx_txt),
                len(cols) if isinstance(cols, list) else None,
                _last.get("rows"),
            )
        except Exception:
            pass

        # 2.9) LLM 분석 전용 컨텍스트 우선 사용
        # - 위의 기존 columns/records 샘플 JSON은 하위호환용으로 일단 만들어지지만,
        # - __sims_analysis_ctx가 있으면 최종 sims_block은 분석용 컨텍스트로 덮어쓴다.
        # - 이렇게 해야 LLM이 sample_rows=10 같은 예전 샘플 표가 아니라
        #   최신 전체 결과 기준 분석 컨텍스트를 보게 된다.
        try:
            analysis_ctx = st.session_state.get("__sims_analysis_ctx")

            if (
                isinstance(analysis_ctx, dict)
                and analysis_ctx.get("kind") == "SIMS_ANALYSIS_CONTEXT_V1"
            ):
                sims_context_kind = "analysis"

                # LLM 전달용으로 너무 큰 대표목록만 살짝 축소한다.
                # 전체 판단은 summary/counts가 기준이고, 품목 설명은 대표 목록을 사용한다.
                analysis_for_llm = dict(analysis_ctx)

                # 현재표 후속분석용 핵심 집계는 LLM JSON 앞쪽에 명시적으로 유지한다.
                sales_time_profile = analysis_for_llm.get("sales_time_profile") or {}
                sales_group_profile = analysis_for_llm.get("sales_group_profile") or {}

                if isinstance(sales_time_profile, dict):
                    monthly_sales = sales_time_profile.get("monthly_sales") or []
                    if isinstance(monthly_sales, list):
                        # 2020~2026이면 월 84개 정도라 그대로 가능하나, 안전하게 120개 제한
                        sales_time_profile["monthly_sales"] = monthly_sales[:120]

                    daily_top = sales_time_profile.get("daily_sales_top") or []
                    if isinstance(daily_top, list):
                        sales_time_profile["daily_sales_top"] = daily_top[:20]

                    weekday_sales = sales_time_profile.get("weekday_sales") or []
                    if isinstance(weekday_sales, list):
                        sales_time_profile["weekday_sales"] = weekday_sales[:7]

                    weekday_top = sales_time_profile.get("weekday_sales_top") or []
                    if isinstance(weekday_top, list):
                        sales_time_profile["weekday_sales_top"] = weekday_top[:7]

                    analysis_for_llm["sales_time_profile"] = sales_time_profile

                if isinstance(sales_group_profile, dict):
                    for k in (
                        "product_sales_top",
                        "product_quantity_top",
                        "vendor_sales_top",
                        "vendor_quantity_top",
                        "staff_sales",
                    ):
                        v = sales_group_profile.get(k) or []
                        if isinstance(v, list):
                            sales_group_profile[k] = v[:20]

                    analysis_for_llm["sales_group_profile"] = sales_group_profile

                try:
                    log.debug(
                        "[chat] analysis profile for llm time_keys=%s group_keys=%s daily_top=%s monthly=%s product_qty_top=%s",
                        list(sales_time_profile.keys()) if isinstance(sales_time_profile, dict) else None,
                        list(sales_group_profile.keys()) if isinstance(sales_group_profile, dict) else None,
                        len((sales_time_profile or {}).get("daily_sales_top") or []) if isinstance(sales_time_profile, dict) else 0,
                        len((sales_time_profile or {}).get("monthly_sales") or []) if isinstance(sales_time_profile, dict) else 0,
                        len((sales_group_profile or {}).get("product_quantity_top") or []) if isinstance(sales_group_profile, dict) else 0,
                    )
                except Exception:
                    pass

                risk_top = analysis_for_llm.get("risk_products_top") or []
                if isinstance(risk_top, list):
                    analysis_for_llm["risk_products_top"] = risk_top[:30]

                grade_samples = analysis_for_llm.get("grade_samples") or {}
                if isinstance(grade_samples, dict):
                    analysis_for_llm["grade_samples"] = {
                        str(k): (v[:5] if isinstance(v, list) else v)
                        for k, v in grade_samples.items()
                    }

                sims_json = json.dumps(analysis_for_llm, ensure_ascii=False, default=str)

                try:
                    _ctx_len2 = int(st.session_state.get("__model_ctx_len") or 16384)
                except Exception:
                    _ctx_len2 = 16384

                if _ctx_len2 <= 12288:
                    analysis_json_limit = 12000
                elif _ctx_len2 <= 16384:
                    analysis_json_limit = 14000
                elif _ctx_len2 <= 32768:
                    analysis_json_limit = 20000
                else:
                    analysis_json_limit = 28000

                ctx_txt = _clip_for_model(sims_json, limit=analysis_json_limit)

                sims_block = (
                    "[SIMS_META] "
                    f"kind=analysis "
                    f"action={analysis_for_llm.get('action')} "
                    f"rows={analysis_for_llm.get('row_count')} "
                    f"cols={analysis_for_llm.get('column_count')} "
                    f"included_detail_products={len(analysis_for_llm.get('risk_products_top') or [])}\n"
                    "다음 SIMS ERP JSON은 샘플 표가 아니라, 최신 SIMS 조회 결과 전체 원본 DataFrame을 "
                    "Python에서 집계한 LLM 분석용 컨텍스트입니다.\n"
                    "\n"
                    "- summary, shortage_grade_counts, forecast_grade_counts, trend_judge_counts는 전체 결과 기준입니다.\n"
                    "- 일자/날짜/요일/일별/월별 매출 질문은 sales_time_profile을 우선 근거로 답하세요.\n"
                    "- 제품별/거래처별/영업사원별/수량 TOP 질문은 sales_group_profile을 우선 근거로 답하세요.\n"
                    "- whole_table_profile은 전체 원본 DataFrame 기준으로 만든 등급별/제조사별/제품군별/TOP 품목 집계입니다. 전체 판단은 이 값을 우선 사용하세요.\n"
                    "- risk_products_top과 grade_samples는 내부 분석용 상세 품목 목록입니다. 답변에는 내부 키 이름을 노출하지 마세요.\n"
                    "- 절대로 '샘플 데이터', '전체를 대표하지 않을 수 있음', 'sample_rows=10'이라고 말하지 마세요.\n"
                    "- 이전 대화, 이전 표, 이전 답변은 분석 근거로 사용하지 마세요.\n"
                    "- 전체 수치 판단은 summary와 *_counts를 우선 사용하세요.\n"
                    "- stock_mode 또는 stock_label 값이 '실수불'이면 반드시 '실수불 기준'이라고 그대로 표현하고, '실재고 기준'으로 바꾸지 마세요.\n"
                    "- 답변 본문에 '샘플', '샘플 데이터', 'sample_rows', '대표 목록'이라는 표현을 쓰지 말고, 필요한 경우 '전체 조회 결과 기준'이라고만 표현하세요.\n"
                    f"[SIMS_JSON]\n{ctx_txt}\n[/SIMS_JSON]"
                )

                log.debug(
                    "[chat] build_messages_with_system → SIMS_JSON attached "
                    "(len=%s, kind=analysis, cols=%s, rows=%s, risk_top=%s)",
                    len(ctx_txt),
                    analysis_for_llm.get("column_count"),
                    analysis_for_llm.get("row_count"),
                    len(analysis_for_llm.get("risk_products_top") or []),
                )

        except Exception:
            log.exception("[chat] build analysis SIMS_JSON failed")


    # 3) JSON이 없으면 텍스트 컨텍스트라도 사용
    elif attach_sims and isinstance(sims_ctx, str) and sims_ctx.strip():
        ctx_txt = _clip_for_model(sims_ctx.strip())
        sims_block = f"[SIMS CONTEXT]\n{ctx_txt}\n[/SIMS CONTEXT]"
        log.debug(
            "[chat] build_messages_with_system → attach SIMS CONTEXT (len=%s)",
            len(ctx_txt),
        )

    else:
        log.debug(
            "[chat] build_messages_with_system → NO SIMS BLOCK "
            "(attach_sims=%s, has_sims_data=%s, has_sims_ctx=%s)",
            attach_sims,
            isinstance(sims_data, dict) and bool(sims_data),
            isinstance(sims_ctx, str) and bool(sims_ctx.strip()),
        )

    # 4) system 메시지에 SIMS 블록 병합
    if sims_block:
        if msgs and msgs[0]["role"] == "system":
            if sims_block not in msgs[0]["content"]:
                msgs[0]["content"] = (msgs[0]["content"].rstrip() + "\n\n" + sims_block).strip()
        else:
            msgs.append({"role": "system", "content": sims_block})

    # 4.5) SIMS_ANALYSIS_CONTEXT_V1 모드에서는 과거 대화 이력을 절대 붙이지 않는다.
    # 조건을 sims_context_kind 하나에만 의존하지 않고,
    # sims_block 내용과 session_state의 분석 컨텍스트까지 함께 확인한다.
    try:
        _analysis_ctx = st.session_state.get("__sims_analysis_ctx")
        analysis_mode = bool(
            attach_sims
            and (
                sims_context_kind == "analysis"
                or (
                    isinstance(_analysis_ctx, dict)
                    and _analysis_ctx.get("kind") == "SIMS_ANALYSIS_CONTEXT_V1"
                )
                or ("SIMS_ANALYSIS_CONTEXT_V1" in str(sims_block or ""))
                or ("kind=analysis" in str(sims_block or ""))
            )
        )
    except Exception:
        analysis_mode = False

    if analysis_mode:
        current_question = (user_text or "").strip()
        if not current_question:
            current_question = "현재 조회 결과를 핵심 요약, 주요 수치, 주의할 점, 다음 조회 제안 순서로 분석해줘"

        analysis_rule = (
            "\n\n[SIMS_ANALYSIS_RULE]\n"
            "- 현재 SIMS_JSON은 SIMS_ANALYSIS_CONTEXT_V1이다.\n"
            "- 이 컨텍스트는 최신 SIMS 조회 결과 전체 원본 DataFrame을 Python에서 집계한 분석 컨텍스트다.\n"
            "- summary와 shortage_grade_counts, forecast_grade_counts, trend_judge_counts는 전체 결과 기준이다.\n"
            "- 일자/날짜/요일/일별/월별 매출 질문은 sales_time_profile을 우선 근거로 답한다.\n"
            "- 제품별/거래처별/영업사원별/수량 TOP 질문은 sales_group_profile을 우선 근거로 답한다.\n"
            "- risk_products_top은 품목 설명용 대표 위험 목록이며, 전체 판단은 summary와 *_counts를 우선 사용한다.\n"
            "- risk_products_top의 건수가 적어도 '샘플 데이터', '전체를 대표하지 않을 수 있음'이라고 말하지 마라.\n"
            "- sample_rows=10, sample_rows=300 같은 표현을 답변에 쓰지 마라.\n"
            "- stock_mode가 real 또는 실재고이면 Rddbc220이라고 말하지 마라.\n"
            "- 이전 대화, 이전 표, 이전 답변은 분석 근거로 사용하지 마라.\n"
            "- source_table 또는 source_table_label이 명시되어 있지 않으면 테이블명(Rddbc210/Rddbc220)을 추정해서 말하지 마라.\n"
            "- risk_products_top이라는 내부 키 이름을 답변에 그대로 쓰지 마라.\n"
            "- risk_products_top의 전달 건수가 8건/30건이어도, 전체 분석이 8건만 기준이라고 말하지 마라.\n"
            "- 필요하면 '우선 확인 대상 품목' 또는 '위험 품목 상세 목록'이라고 표현하라.\n"
            "- 재고 기준과 기준 테이블은 SIMS_JSON.summary.stock_basis_label, SIMS_JSON.summary.source_table_label 값만 사용하라.\n"
            "- low_stock_records, sample_records, risk_products_top, grade_samples 같은 내부 key 이름을 답변에 노출하지 마세요.\n"
            "- 답변 본문에 '샘플', '샘플 데이터', 'sample_rows', '대표 목록'이라는 표현을 쓰지 말고, 필요한 경우 '전체 조회 결과 기준' 또는 '참고 목록'이라고만 표현하세요.\n"
            "- source_table_label이 없으면 Rddbc210/Rddbc220 같은 테이블명을 추정해서 말하지 말고, '실재고 기준' 또는 '장부재고 기준'까지만 말하라.\n"
            "- stock_mode가 real 또는 실재고이면 Rddbc220이라고 말하지 마라.\n"
        )

        system_text = msgs[0]["content"] if msgs and msgs[0].get("role") == "system" else ""
        if "[SIMS_ANALYSIS_RULE]" not in system_text:
            system_text = system_text.rstrip() + analysis_rule

        msgs = [
            {"role": "system", "content": system_text},
            {"role": "user", "content": current_question},
        ]

        log.debug(
            "[chat] build_messages_with_system analysis mode → skip history, msgs=%s, kind=%s",
            len(msgs),
            sims_context_kind,
        )
        return msgs

    # 5) 히스토리 user/assistant 메시지 이어 붙이기


    # ✅ 2번 정책과 세트: SIMS_JSON이 커지므로 히스토리는 예산 내에서만 "뒤에서부터" 붙인다.
    #    (토큰 기준이 이상적이지만, 여기서는 안전한 문자 예산으로 근사)
    HISTORY_CHAR_BUDGET = 1200  # SIMS_JSON 우선. 이전 대화는 짧게 유지

    used = 0
    tail: list[dict] = []

    # ✅ 일반 질문일 때는 SIMS 관련 히스토리를 제외해 "SIMS에 없다" 답변을 방지
    sims_noise = re.compile(r"(SIMS|ERP|\[SIMS_|컨텍스트|거래처|사용자목록|부서별|Rddbc0\d+)", re.IGNORECASE)

    for m in reversed(history_msgs):
        if m.get("role") not in ("user", "assistant"):
            continue
        txt = (m.get("content", "") or "")
        if not attach_sims and sims_noise.search(txt):
            continue
        txt = _clip_for_model(txt)  # 각 메시지 자체가 너무 길면 컷
        if used + len(txt) > HISTORY_CHAR_BUDGET:
            break
        tail.append({"role": m["role"], "content": txt})
        used += len(txt)
    msgs.extend(reversed(tail))

    return msgs

@st.cache_data(ttl=60)
def get_models() -> List[str]:
    try:
        models = CLIENT.models.list()
        return [m.id for m in models.data]
    except Exception:
        return ["local-model"]

# =========================
# 저장/로드
# =========================
def _normalize_messages(messages: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    기존 메시지 호환 정리.

    중요:
    - user_id/company_id/message_type 같은 신규 메타를 버리지 않는다.
    """
    norm = []

    for m in messages or []:
        if not isinstance(m, dict):
            continue

        item = dict(m)
        item["role"] = item.get("role", "user")
        item["content"] = item.get("content", "")
        item["time"] = (
            item.get("time")
            or item.get("timestamp")
            or datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        )

        norm.append(item)

    return norm

def _current_chat_meta() -> dict[str, Any]:
    user = get_current_user()
    company = get_selected_company()

    meta: dict[str, Any] = {}

    if user:
        meta.update(
            {
                "user_id": int(user.user_id),
                "owner_login_id": str(user.login_id or ""),
                "owner_user_type": str(user.user_type or ""),
                "owner_user_grade": str(user.user_grade or ""),
            }
        )

    if isinstance(company, dict):
        meta.update(
            {
                "company_id": company.get("company_id"),
                "company_name": company.get("company_name"),
                "db_name": company.get("db_name"),
            }
        )

    return meta


def _default_room_name() -> str:
    now_text = datetime.now().strftime("%Y-%m-%d %H:%M")

    user = get_current_user()
    company = get_selected_company()

    user_type = str(getattr(user, "user_type", "") or "")

    if user_type.startswith("WHOLESALE") and isinstance(company, dict):
        company_name = str(company.get("company_name") or "").strip()
        if company_name:
            if len(company_name) > 20:
                company_name = company_name[:20] + "…"
            return f"{now_text} {company_name} 업무 대화"

    return f"{now_text} 업무 대화"

def _room_title_text_from_message(content: str, *, limit: int = 32) -> str:
    """
    첫 사용자 메시지에서 채팅방 제목으로 쓸 짧은 문구를 만든다.
    LLM 호출 없이 즉시 처리한다.
    """
    text = str(content or "").strip()

    if not text:
        return ""

    # 코드블록/마크다운 일부 정리
    text = text.replace("```", " ")
    text = re.sub(r"[\r\n\t]+", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    text = text.strip(" \"'“”‘’`")

    # 너무 흔한 접두어 제거
    text = re.sub(r"^(질문|문의|요청)\s*[:：]\s*", "", text).strip()

    if not text:
        return ""

    if len(text) > limit:
        text = text[:limit].rstrip() + "…"

    return text


def _auto_room_name_from_first_message(room: dict[str, Any]) -> str:
    """
    첫 user 메시지를 기준으로 자동 채팅방 이름을 만든다.
    예: 2026-06-28 19:29 저녁시간을 즐기자
    """
    if not isinstance(room, dict):
        return ""

    first_user_text = ""
    first_user_time = ""

    for msg in room.get("messages") or []:
        if not isinstance(msg, dict):
            continue

        if msg.get("role") != "user":
            continue

        content = str(msg.get("content") or "").strip()
        if content:
            first_user_text = content
            first_user_time = str(msg.get("time") or "").strip()
            break

    title_text = _room_title_text_from_message(first_user_text)

    if not title_text:
        return ""

    # 제목 시간은 방 생성 시간이 아니라 첫 사용자 메시지 시간을 우선 사용한다.
    base_time = (
        first_user_time
        or str(room.get("created_at") or "")
        or make_ts()
    )
    base_time = str(base_time).replace("T", " ")[:16]

    return f"{base_time} {title_text}"


def _ensure_auto_room_name(room: dict[str, Any]) -> None:
    """
    자동 생성 방은 첫 메시지가 들어온 뒤 방 이름을 첫 메시지 기준으로 바꾼다.
    단, 사용자가 직접 이름을 바꾼 방은 건드리지 않는다.

    주의:
    - 이 함수는 채팅 입력 처리 중 호출될 수 있다.
    - Streamlit 위젯 key(__room_rename_buf)는 이미 생성된 뒤 수정하면 오류가 날 수 있으므로
      여기서는 변경 플래그만 남기고, 다음 rerun에서 사이드바 렌더 전에 입력칸이 갱신되게 한다.
    """
    if not isinstance(room, dict):
        return

    if room.get("name_auto") is not True:
        return

    new_name = _auto_room_name_from_first_message(room)

    if not new_name:
        return

    if room.get("name") != new_name:
        old_name = room.get("name")

        room["name"] = new_name
        room["title_source"] = "first_user_message"

        try:
            if str(st.session_state.get("current_room") or "") == str(room.get("id") or ""):
                st.session_state["__chat_room_title_changed_room_id"] = str(room.get("id") or "")
                st.session_state["__chat_room_title_changed_name"] = new_name
        except Exception:
            pass

        try:
            log.info(
                "[chat.room.title] auto renamed %s old_name=%s new_name=%s",
                _chat_log_kv(room),
                _safe_log_value(old_name),
                _safe_log_value(new_name),
            )
        except Exception:
            pass

def _make_chat_room(*, name: str | None = None, auto_created: bool = False) -> dict[str, Any]:
    meta = _current_chat_meta()

    room = {
        "id": str(uuid.uuid4()),
        "name": name or _default_room_name(),
        "created_at": make_ts(),
        "updated_at": make_ts(),
        "auto_created": bool(auto_created),
        "name_auto": name is None,
        "title_source": "default",        
        "messages": [],
        "gen_messages": [],
        "sims_messages": [],
    }

    room.update(meta)
    return room


def _room_has_any_messages(room: dict[str, Any] | None) -> bool:
    """저장 대상 메시지/SIMS history가 하나라도 있는 채팅방인지 판정한다."""
    if not isinstance(room, dict):
        return False

    # SIMS 패널 결과는 messages가 아니라 history에 저장될 수 있다.
    for key in ("messages", "gen_messages", "sims_messages", "history"):
        value = room.get(key)
        if isinstance(value, list) and len(value) > 0:
            return True

    return False

def _is_empty_auto_room(room: dict[str, Any] | None) -> bool:
    """로그인 직후 입력 대기용 임시방인지 판정한다."""
    return bool(
        isinstance(room, dict)
        and room.get("auto_created") is True
        and not _room_has_any_messages(room)
    )


def _select_pending_new_room() -> dict[str, Any]:
    """
    로그인 직후 기본 입력 대상은 기존 방이 아니라 새 대화 대기방이다.

    정책:
    - 기존 채팅방은 목록에만 보여준다.
    - 사용자가 기존 방을 직접 선택하지 않고 입력하면 이 임시방이 정식 방으로 저장된다.
    - 임시방은 메시지가 생기기 전까지 JSON 파일에 저장하지 않는다.
    """
    ss = st.session_state
    ss.setdefault("chat_rooms", [])

    # 이미 만들어 둔 빈 임시방이 있으면 재사용한다.
    for room in ss.chat_rooms:
        if _is_empty_auto_room(room):
            ss.current_room = room.get("id")
            return room

    room = _make_chat_room(auto_created=True)
    ss.chat_rooms.append(room)
    ss.current_room = room["id"]

    try:
        log.info("[chat.room] pending new room %s", _chat_log_kv(room))
    except Exception:
        pass

    return room


def _get_current_room_or_pending() -> dict[str, Any]:
    """현재 선택 방을 가져오되, 없거나 깨졌으면 새 대화 대기방을 반환한다."""
    ss = st.session_state
    ss.setdefault("chat_rooms", [])

    current_id = str(ss.get("current_room") or "").strip()
    if current_id:
        for room in ss.chat_rooms:
            if str(room.get("id") or "") == current_id:
                return room

    return _select_pending_new_room()

def _drop_empty_auto_rooms(*, keep_room_id: str = "") -> int:
    """
    사용자가 기존 채팅방을 선택했을 때,
    저장 전 새 대화 대기방이 목록에 계속 남아 어색하게 보이는 것을 방지한다.
    """
    ss = st.session_state
    rooms = ss.get("chat_rooms") or []

    kept = []
    removed = 0
    keep_room_id = str(keep_room_id or "")

    for room in rooms:
        rid = str(room.get("id") or "") if isinstance(room, dict) else ""

        if rid == keep_room_id:
            kept.append(room)
            continue

        if _is_empty_auto_room(room):
            removed += 1
            continue

        kept.append(room)

    if removed:
        ss["chat_rooms"] = kept

    return removed


def _ensure_sims_panel_room_title(room: dict[str, Any], action: str) -> None:
    """
    첫 이벤트가 일반 채팅이 아니라 SIMS 패널 조회인 경우,
    빈 새 대화 대기방을 정식 방으로 전환하고 제목을 SIMS action 기준으로 확정한다.
    """
    if not isinstance(room, dict):
        return

    action_text = _room_title_text_from_message(str(action or "SIMS 조회"), limit=36) or "SIMS 조회"
    new_name = f"{make_ts()[:16]} {action_text}"

    if room.get("auto_created") is True and not _room_has_any_messages(room):
        old_name = room.get("name")

        room["name"] = new_name
        room["auto_created"] = False
        room["name_auto"] = True
        room["title_source"] = "sims_panel_action"
        room["updated_at"] = make_ts()

        try:
            st.session_state["__chat_room_title_changed_room_id"] = str(room.get("id") or "")
        except Exception:
            pass

        try:
            log.info(
                "[chat.room.title] sims panel renamed %s old_name=%s new_name=%s",
                _chat_log_kv(room),
                _safe_log_value(old_name),
                _safe_log_value(new_name),
            )
        except Exception:
            pass


def _push_panel_result_to_current_chat(
    room: dict[str, Any],
    *,
    selected_for_render: dict[str, Any] | None = None,
    run_seq: int | None = None,
) -> bool:
    """
    SIMS 패널에서 조회한 최종 결과를 현재 채팅방 history에 1회 저장한다.

    정책:
    - 패널에는 표를 직접 렌더링하지 않는다.
    - 채팅방에 조회조건 + 요약 + 표 + 다운로드 기준으로 표시한다.
    - df는 전체 기준, df_display/records는 화면 표시 제한 기준.
    """
    ss = st.session_state

    if not ss.get("__sims_was_final"):
        return False

    selected_for_render = selected_for_render or {}

    stored_payload = ss.get("__sims_last_final_payload_for_chat")
    if not isinstance(stored_payload, dict):
        stored_payload = ss.get("__sims_panel_last_final_payload")

    if isinstance(stored_payload, dict) and not _sims_payload_matches_current_company(stored_payload):
        payload_company_id, payload_db_name = _sims_payload_company_sig(stored_payload)
        current_company_id, current_db_name = _sims_current_company_sig()

        ss.pop("__sims_last_final_payload_for_chat", None)
        ss.pop("__sims_last_final_payload_for_chat_action", None)
        ss.pop("__sims_panel_last_final_payload", None)
        ss.pop("__sims_panel_last_final_action", None)
        ss.pop("__sims_panel_source_promoted_sig", None)

        log.info(
            "[chat.panel.push] skip stale company payload payload_company_id=%s payload_db=%s current_company_id=%s current_db=%s run_seq=%s",
            payload_company_id,
            payload_db_name,
            current_company_id,
            current_db_name,
            run_seq or ss.get("__sims_run_seq"),
        )
        return False

    action = str(
        (stored_payload or {}).get("action")
        or ss.get("__sims_last_table_action")
        or ss.get("__sims_current_table_source_action")
        or selected_for_render.get("action")
        or "SIMS 조회"
    ).strip()

    table_key = str(
        ((stored_payload or {}).get("meta") or {}).get("table_key")
        or ss.get("__sims_last_table_key")
        or ss.get("__sims_current_table_source_key")
        or ""
    ).strip()

    df_full = None
    df_display = None

    if isinstance(stored_payload, dict):
        cand_full = stored_payload.get("df")
        cand_display = stored_payload.get("df_display")

        if isinstance(cand_full, pd.DataFrame) and not cand_full.empty:
            df_full = cand_full

        if isinstance(cand_display, pd.DataFrame) and not cand_display.empty:
            df_display = cand_display

    # fallback: table store에서 복구
    if not isinstance(df_full, pd.DataFrame) or df_full.empty:
        try:
            if table_key:
                for store_name in ("__sims_export_tables_by_key", "sims_export_tables"):
                    store = ss.get(store_name)
                    if isinstance(store, dict):
                        cand = store.get(table_key)
                        if isinstance(cand, pd.DataFrame) and not cand.empty:
                            df_full = cand
                            break

                store = ss.get("sims_tables")
                if isinstance(store, dict):
                    cand = store.get(table_key)
                    if isinstance(cand, pd.DataFrame) and not cand.empty:
                        df_display = cand
        except Exception:
            log.exception("[chat.panel.push] table store lookup failed")

    if not isinstance(df_full, pd.DataFrame) or df_full.empty:
        try:
            df_full, table_key2 = _current_table_get_latest_df()
            if not table_key:
                table_key = table_key2
        except Exception:
            log.exception("[chat.panel.push] current table lookup failed")
            df_full = None

    if not isinstance(df_full, pd.DataFrame) or df_full.empty:
        # 0건 조회/안내성 결과는 df가 없어도 text payload로 채팅창에 올린다.
        if isinstance(stored_payload, dict) and str(stored_payload.get("type") or "").strip().lower() in {"text", "message"}:
            try:
                from app.ui.chat_middleware import push_sims_result_to_chat

                message = str(
                    stored_payload.get("data")
                    or (stored_payload.get("meta") or {}).get("message")
                    or "해당 조회조건의 자료가 없습니다."
                ).strip()

                if not message:
                    message = "해당 조회조건의 자료가 없습니다."

                payload = dict(stored_payload)
                meta = dict(payload.get("meta") or {})

                meta.update(
                    {
                        "panel_push": True,
                        "_force_push": True,
                        "_panel_run_seq": run_seq or ss.get("__sims_run_seq"),
                        "row_count": 0,
                        "row_count_total": 0,
                        "display_row_count": 0,
                        "download_row_count": 0,
                        "column_count": 0,
                        "empty_result": True,
                        "source": "SIMS 패널",
                        "hide_meta_expander": str(os.getenv("SSAI_DEBUG_META", "false")).strip().lower()
                        not in {"1", "true", "yes", "y", "on"},
                    }
                )

                if not str(meta.get("query_summary") or "").strip():
                    qs = str(meta.get("condition") or "").strip()
                    if not qs:
                        qs = str(action or "").strip()
                    meta["query_summary"] = qs

                payload.update(
                    {
                        "final": True,
                        "type": "text",
                        "title": payload.get("title") or f"📋 조회 결과 — {action}",
                        "action": action,
                        "data": message,
                        "meta": meta,
                    }
                )

                sig_src = f"{run_seq or ss.get('__sims_run_seq') or 0}::text::{action}::{message}"
                sig = hashlib.sha256(sig_src.encode("utf-8")).hexdigest()[:16]

                if ss.get("__sims_panel_chat_push_sig") == sig:
                    return False

                _ensure_sims_panel_room_title(room, action)
                push_sims_result_to_chat(payload, action=action)

                ss["__sims_panel_chat_push_sig"] = sig

                log.info(
                    "[chat.panel.push] text saved action=%s run_seq=%s message=%s",
                    action,
                    run_seq,
                    message[:80],
                )
                return True

            except Exception:
                log.exception("[chat.panel.push] text payload push failed")
                return False

        log.warning(
            "[chat.panel.push] skip: no df action=%s table_key=%s run_seq=%s",
            action,
            table_key,
            run_seq,
        )
        return False



    # 표시용 df 제한
    try:
        display_cap = int(os.getenv("SIMS_CHAT_DISPLAY_MAX_ROWS", "2000"))
    except Exception:
        display_cap = 2000

    if not isinstance(df_display, pd.DataFrame) or df_display.empty:
        df_display = df_full.head(min(display_cap, len(df_full))).copy()
    elif display_cap > 0 and len(df_display) > display_cap:
        df_display = df_display.head(display_cap).copy()

    sig = f"{run_seq or ss.get('__sims_run_seq') or 0}::{table_key}::{action}::{len(df_full)}"
    if ss.get("__sims_panel_chat_push_sig") == sig:
        return False

    _ensure_sims_panel_room_title(room, action)

    if isinstance(stored_payload, dict):
        payload = dict(stored_payload)
        meta = dict(payload.get("meta") or {})
    else:
        payload = {}
        meta = {}

    meta.update({
        "panel_push": True,
        "_force_push": True,
        "_panel_run_seq": run_seq or ss.get("__sims_run_seq"),
        "_panel_table_key": table_key,
        "table_key": table_key,
        "row_count": int(len(df_display)),
        "row_count_total": int(len(df_full)),
        "display_row_count": int(len(df_display)),
        "download_row_count": int(len(df_full)),
        "column_count": int(len(df_display.columns)),
        "source": "SIMS 패널",
        "hide_meta_expander": str(os.getenv("SSAI_DEBUG_META", "false")).strip().lower()
        not in {"1", "true", "yes", "y", "on"},
    })

    # 조회조건/요약 보존
    if not str(meta.get("query_summary") or "").strip():
        qs = str(meta.get("condition") or "").strip()
        if not qs:
            qs = str(action or "").strip()
        meta["query_summary"] = qs

    if not str(meta.get("summary_md") or "").strip():
        meta["summary_md"] = (
            f"조회조건: {meta.get('query_summary') or action}\n\n"
            f"조회결과: 전체 {len(df_full):,}건 / 화면 표시 {len(df_display):,}건"
        )

    payload.update({
        "final": True,
        "type": "table",
        "title": payload.get("title") or action,
        "action": action,
        "params": payload.get("params") or {},
        "df": df_full,
        "df_display": df_display,
        "data": df_display,
        "records": df_display.to_dict(orient="records"),
        "columns": list(df_display.columns),
        "message": payload.get("message") or f"{action} {len(df_full):,}건",
        "meta": meta,
    })

    from app.ui.chat_middleware import push_sims_result_to_chat

    push_sims_result_to_chat(payload, action)

    _sync_room_meta(room, materialize=True)
    ss["__sims_panel_chat_push_sig"] = sig

    # 같은 payload가 다음 rerun에서 다시 push되지 않도록 정리
    ss.pop("__sims_last_final_payload_for_chat", None)
    ss.pop("__sims_last_final_payload_for_chat_action", None)

    save_chat_rooms()

    log.info(
        "[chat.panel.push] saved %s action=%s table_key=%s rows=%s display_rows=%s run_seq=%s",
        _chat_log_kv(room),
        action,
        table_key,
        len(df_full),
        len(df_display),
        run_seq or ss.get("__sims_run_seq"),
    )

    return True


def _room_meta_from_last_message(room: dict[str, Any]) -> dict[str, Any]:
    """
    기존 저장된 방을 보정할 때 마지막 메시지의 사용자/회사 메타를 가져온다.
    마지막 메시지에 메타가 없으면 뒤에서부터 메타가 있는 메시지를 찾는다.
    """
    if not isinstance(room, dict):
        return {}

    keys = (
        "user_id",
        "owner_login_id",
        "owner_user_type",
        "owner_user_grade",
        "company_id",
        "company_name",
        "db_name",
    )

    for msg in reversed(room.get("messages") or []):
        if not isinstance(msg, dict):
            continue

        meta: dict[str, Any] = {}

        for key in keys:
            value = msg.get(key)
            if value not in (None, ""):
                meta[key] = value

        if meta:
            return meta

    return {}


def _sync_room_meta(
    room: dict[str, Any],
    *,
    materialize: bool = False,
) -> None:
    """
    현재 로그인 사용자/선택 회사 정보를 채팅방 최상위 메타에 동기화한다.

    materialize=True:
    - 임시방(auto_created=True)이 실제 메시지를 가진 정식 채팅방으로 전환됨.
    """
    if not isinstance(room, dict):
        return

    room.update(_current_chat_meta())

    if materialize:
        room["auto_created"] = False
        _ensure_auto_room_name(room)

    room["updated_at"] = make_ts()

def _message_meta(message_type: str = "chat") -> dict[str, Any]:
    meta = _current_chat_meta()
    meta["message_type"] = message_type
    return meta


def _consume_company_change_notice(room: dict[str, Any]) -> None:
    """
    ssai_login.py에서 큐잉한 회사 변경 안내를 현재 채팅방에 표시/저장한다.
    """
    notice = st.session_state.pop("__ssai_company_change_notice", None)

    if not isinstance(notice, dict):
        return

    # 회사 변경 시점에는 이전 회사 기준의 현재표/컨텍스트/다운로드 캐시를 먼저 비운다.
    _clear_sims_runtime_for_company_change("company_change_notice")

    old_name = str(notice.get("old_company_name") or "이전 회사").strip()
    new_name = str(notice.get("new_company_name") or "변경 회사").strip()
    new_db = str(notice.get("new_db_name") or "").strip()

    message = (
        f"사용 DB가 **{old_name}**에서 **{new_name}**으로 변경되었습니다.\n\n"
        "이전 DB 기준의 현재표/조회결과/후속분석 컨텍스트는 초기화되었습니다.\n"
        f"이후 조회와 분석은 **{new_name}** 기준으로 실행됩니다."
    )

    if new_db:
        message += f"\n\nDB: `{new_db}`"

    # 1) 화면에 즉시 표시
    st.info(message)

    # 2) 같은 회사 변경 안내가 이미 현재 채팅방 마지막 메시지에 있으면 중복 저장 방지
    messages = room.setdefault("messages", [])

    if messages:
        last = messages[-1]
        if (
            isinstance(last, dict)
            and last.get("message_type") == "company_change"
            and str(last.get("content") or "").strip() == message.strip()
        ):
            log.info(
                "[company.change.notice] duplicated skip %s old=%s new=%s db=%s",
                _chat_log_kv(room),
                old_name,
                new_name,
                new_db,
            )
            return

    # 3) 채팅방에도 assistant 메시지로 저장
    item = {
        "id": str(uuid.uuid4()),
        "seq": _next_seq(),
        "role": "assistant",
        "content": message,
        "time": make_ts(),
        **_message_meta("company_change"),
    }

    messages.append(item)

    _sync_room_meta(room, materialize=True)

    save_chat_rooms()

    log.info(
        "[company.change.notice] saved %s old=%s new=%s db=%s",
        _chat_log_kv(room),
        old_name,
        new_name,
        new_db,
    )


# =================================
# 일명(액션명 슬러그+타임스탬프)
# =================================
def migrate_rooms(data: Union[dict, list, None]) -> List[Dict[str, Any]]:
    if data is None:
        return []

    current_meta = _current_chat_meta()

    def _normalize_room(room: dict, fallback_id: str | None = None, fallback_name: str | None = None) -> dict:
        rid = room.get("id") or fallback_id or str(uuid.uuid4())
        name = room.get("name") or fallback_name or "업무 대화"

        out = dict(room)
        out["id"] = rid
        out["name"] = name
        out["messages"] = _normalize_messages(room.get("messages", []))

        if "gen_messages" in room:
            out["gen_messages"] = _normalize_messages(room.get("gen_messages", []))
        else:
            out["gen_messages"] = list(out["messages"])

        if "sims_messages" in room:
            out["sims_messages"] = _normalize_messages(room.get("sims_messages", []))
        else:
            out["sims_messages"] = []

        # user_id 기준 파일로 분리되므로, 없는 메타는 현재 로그인 사용자 기준으로 보강
        for k, v in current_meta.items():
            out.setdefault(k, v)

        out.setdefault("created_at", room.get("created_at") or make_ts())
        out.setdefault("updated_at", room.get("updated_at") or make_ts())

        return out

    if isinstance(data, list):
        out = []
        for i, room in enumerate(data):
            if isinstance(room, dict):
                out.append(_normalize_room(room, fallback_name=f"대화 {i + 1}"))
        return out

    if isinstance(data, dict):
        out = []
        for rid, room in data.items():
            if isinstance(room, dict):
                out.append(_normalize_room(room, fallback_id=str(rid), fallback_name=str(rid)))
        out.sort(key=lambda r: str(r.get("name") or ""))
        return out

    return []


def _effective_chat_file() -> str:
    """
    채팅방 저장 파일.

    정책:
    - 로그인 전에는 .env CHAT_FILE 사용
    - 로그인 후에는 user_id 기준 파일로 분리
    - 1호기/2호기는 .env의 CHAT_FILE parent가 다르므로 공유하지 않음
    """
    try:
        user = get_current_user()
    except Exception:
        user = None

    if not user:
        return str(CHAT_FILE)

    base = Path(str(CHAT_FILE))
    return str(base.parent / f"user_{int(user.user_id)}_chat_rooms.json")



def _safe_log_value(value: Any, limit: int = 120) -> str:
    """
    운영 로그용 값 정리.
    - 비밀번호/질문 전문/답변 전문은 이 함수에 넘기지 않는다.
    - 너무 긴 값은 잘라서 로그 폭주를 막는다.
    """
    try:
        s = str(value if value is not None else "").replace("\n", " ").replace("\r", " ").strip()
    except Exception:
        s = ""
    if len(s) > limit:
        return s[:limit].rstrip() + "..."
    return s


def _chat_log_context(room: dict[str, Any] | None = None, **extra: Any) -> dict[str, Any]:
    """
    여러 사용자가 동시에 사용할 때 로그만 보고도
    사용자/회사/DB/채팅방/저장파일을 식별할 수 있게 하는 공통 컨텍스트.
    """
    try:
        user = get_current_user()
    except Exception:
        user = None

    try:
        company = get_selected_company()
    except Exception:
        company = None

    ctx: dict[str, Any] = {
        "user_id": getattr(user, "user_id", None),
        "login_id": getattr(user, "login_id", None),
        "user_type": getattr(user, "user_type", None),
        "user_grade": getattr(user, "user_grade", None),
        "company_id": None,
        "company_name": "",
        "db_name": "",
        "room_id": "",
        "chat_file": "",
    }

    if isinstance(company, dict):
        ctx.update(
            {
                "company_id": company.get("company_id"),
                "company_name": company.get("company_name"),
                "db_name": company.get("db_name"),
            }
        )

    if isinstance(room, dict):
        ctx["room_id"] = room.get("id") or ""

    try:
        ctx["chat_file"] = _effective_chat_file()
    except Exception:
        ctx["chat_file"] = ""

    ctx.update(extra)
    return ctx


def _chat_log_kv(room: dict[str, Any] | None = None, **extra: Any) -> str:
    ctx = _chat_log_context(room, **extra)
    order = [
        "user_id",
        "login_id",
        "user_type",
        "user_grade",
        "company_id",
        "company_name",
        "db_name",
        "room_id",
        "chat_file",
    ]
    order += [k for k in ctx.keys() if k not in order]

    return " ".join(
        f"{key}={_safe_log_value(ctx.get(key))}"
        for key in order
        if ctx.get(key) not in (None, "")
    )


def load_chat_rooms() -> List[Dict[str, Any]]:
    chat_file = _effective_chat_file()

    if not os.path.exists(chat_file):
        log.debug("[chat.load] missing %s", _chat_log_kv(chat_file=chat_file))
        return []

    try:
        with open(chat_file, "r", encoding="utf-8") as f:
            raw = json.load(f)
    except Exception:
        log.exception("[chat.load] failed %s", _chat_log_kv(chat_file=chat_file))
        return []

    rooms = migrate_rooms(raw)
    log.info("[chat.load] %s rooms=%s", _chat_log_kv(chat_file=chat_file), len(rooms))
    return rooms

def _reset_chat_session_when_user_changed() -> None:
    """
    같은 브라우저 세션에서 로그인 사용자가 바뀌면
    이전 사용자의 chat_rooms/current_room을 버리고 다시 로드하게 한다.
    """
    try:
        user = get_current_user()
    except Exception:
        user = None

    owner_id = str(getattr(user, "user_id", "") or "").strip()
    prev_owner_id = str(st.session_state.get("__chat_owner_user_id") or "").strip()

    if not owner_id:
        return

    if prev_owner_id == owner_id:
        return

    log.info(
        "[chat.owner] session owner changed prev_user_id=%s new_user_id=%s %s",
        prev_owner_id or "-",
        owner_id,
        _chat_log_kv(user_id=owner_id),
    )

    for key in [
        "chat_rooms",
        "current_room",
        "__room_pick_radio",
        "__room_prev_selected",
        "__room_rename_buf",
        "__room_prev_name",
        "__chat_room_title_changed_room_id",
        "__chat_room_title_changed_name",
        "__room_filter",
        "__room_page",
        "__seq",
        "__queue_ai",
        "__an_busy",
        "__an_job",
        "__an_cancel",
        "__deferred_current_table_followup",
        "__sims_auto_user_input",
    ]:
        st.session_state.pop(key, None)

    st.session_state["__chat_owner_user_id"] = owner_id


# =========================
# 저장 (원자적)



def _json_sanitize(obj, _depth: int = 0):
    """
    JSON 직렬화 불가 객체(DataFrame, Streamlit DeltaGenerator, set 등)를
    안전한 형태로 변환한다.
    - chat_rooms 저장/백업(zip) 시 TypeError 방지 목적
    """
    if _depth > 6:
        return str(obj)

    # pandas DataFrame
    try:
        import pandas as _pd
        if isinstance(obj, _pd.DataFrame):
            try:
                return {
                    "__type__": "dataframe",
                    "columns": list(obj.columns),
                    "records": obj.to_dict(orient="records"),
                }
            except Exception:
                return {"__type__": "dataframe", "shape": [int(obj.shape[0]), int(obj.shape[1])]}
    except Exception:
        pass

    # dict / list / tuple
    if isinstance(obj, dict):
        return {str(k): _json_sanitize(v, _depth + 1) for k, v in obj.items()}
    if isinstance(obj, (list, tuple)):
        return [_json_sanitize(v, _depth + 1) for v in obj]

    # set
    if isinstance(obj, set):
        return [_json_sanitize(v, _depth + 1) for v in sorted(obj, key=lambda x: str(x))]

    # datetime/date
    try:
        import datetime as _dt
        if isinstance(obj, (_dt.datetime, _dt.date)):
            return obj.isoformat()
    except Exception:
        pass

    # numpy scalars
    try:
        import numpy as _np
        if isinstance(obj, (_np.generic,)):
            return obj.item()
    except Exception:
        pass

    # streamlit DeltaGenerator 등: 문자열로 대체
    tname = type(obj).__name__
    if "DeltaGenerator" in tname:
        return {"__type__": "DeltaGenerator"}

    # 기본: json이 처리 가능한 타입이면 그대로, 아니면 str()
    try:
        import json as _json
        _json.dumps(obj)
        return obj
    except Exception:
        return str(obj)



def _json_default(o):
    """json.dump/json.dumps default handler (DataFrame 등 비직렬 객체 방어)."""
    try:
        import pandas as pd  # lazy
        if isinstance(o, pd.DataFrame):
            head = o.head(500)
            return {
                "__type__": "DataFrame",
                "rows": int(len(o)),
                "cols": [str(c) for c in o.columns],
                "data": head.to_dict(orient="records"),
            }
    except Exception:
        pass
    try:
        return str(o)
    except Exception:
        return repr(o)

def save_chat_rooms():
    """원자적 저장: 임시파일에 쓴 뒤 os.replace로 교체"""
    # ✅ chat_rooms 안에 DataFrame/Streamlit 객체(DeltaGenerator 등)가 섞이면
    # json.dump에서 TypeError가 발생한다.
    # 저장 시에는 '직렬화 가능한 형태'로만 변환해서 파일에 기록한다.
    def _json_safe(o):
        # pandas DataFrame → records/columns로 축약 저장
        try:
            import pandas as _pd
            if isinstance(o, _pd.DataFrame):
                return {
                    "__type__": "DataFrame",
                    "columns": list(o.columns),
                    "records": o.to_dict(orient="records"),
                }
        except Exception:
            pass

        # dict/list 재귀 처리 + table payload의 df/df_display/data 키 제거
        if isinstance(o, dict):
            out = {}
            for k, v in o.items():
                if k in ("df", "df_display", "data"):
                    # DF는 위에서 변환했거나, records/columns로 이미 들어있을 수 있으니 제거
                    # (그대로 남기면 직렬화 에러의 원인이 됨)
                    if "records" not in o and "columns" not in o:
                        try:
                            import pandas as _pd
                            if isinstance(v, _pd.DataFrame):
                                out["columns"] = list(v.columns)
                                out["records"] = v.to_dict(orient="records")
                        except Exception:
                            pass
                    continue
                out[k] = _json_safe(v)
            return out
        if isinstance(o, list):
            return [_json_safe(x) for x in o]
        if isinstance(o, tuple):
            return [_json_safe(x) for x in o]

        # 기본 JSON 타입은 그대로
        try:
            json.dumps(o, ensure_ascii=False)
            return o
        except TypeError:
            # Streamlit DeltaGenerator 등 → 문자열로 안전 처리
            return str(o)

    chat_file = _effective_chat_file()
    tmp_dir = os.path.dirname(chat_file) or "."
    os.makedirs(tmp_dir, exist_ok=True)
    fd, tmp_path = tempfile.mkstemp(prefix="chat_rooms_", suffix=".json", dir=tmp_dir)

    try:
        # 로그인 직후 입력 대기용 빈 auto_created 방은 파일에 저장하지 않는다.
        # 사용자가 입력하거나 회사 변경 안내/SIMS 결과가 들어와 materialize되면 저장 대상이 된다.
        rooms_to_save = [
            room
            for room in (st.session_state.get("chat_rooms") or [])
            if not _is_empty_auto_room(room)
        ]

        with os.fdopen(fd, "w", encoding="utf-8") as f:
            json.dump(_json_sanitize(rooms_to_save), f, ensure_ascii=False, indent=2)
        os.replace(tmp_path, chat_file)

        try:
            log.info(
                "[chat.save] %s rooms=%s saved_rooms=%s",
                _chat_log_kv(chat_file=chat_file),
                len(st.session_state.get("chat_rooms") or []),
                len(rooms_to_save),
            )
        except Exception:
            pass
    except Exception:
        # 실패 시 임시파일 정리
        try:
            os.remove(tmp_path)
        except Exception:
            pass
        raise
# =========================


def migrate_rooms_seq_if_needed(rooms: List[Dict[str, Any]], *, chat_file: str, logger) -> int:
    """
    rooms.json 안에 seq 없는 옛 메시지를 안전하게 마이그레이션.
    - 룸 단위로 messages 순서를 보존하면서 seq를 1..N으로 재부여(필요한 룸만).
    - 변경이 있으면 rooms.json을 백업(.bak_YYYYMMDD_HHMMSS) 후 저장.
    - 반환: 전체 rooms에서의 최대 seq
    """
    import shutil
    from datetime import datetime

    if not isinstance(rooms, list) or not rooms:
        return 0

    changed = False
    max_seq_all = 0

    for r in rooms:
        msgs = r.get("messages") or []
        if not isinstance(msgs, list) or not msgs:
            continue

        # 룸의 seq 상태 점검
        need_renumber = False
        prev = -1
        for m in msgs:
            s = m.get("seq")
            if not isinstance(s, int):
                need_renumber = True
                break
            if s <= prev:
                need_renumber = True
                break
            prev = s

        # 필요 시: 현재 저장 순서를 그대로 유지한 채 1..N 재부여
        if need_renumber:
            for i, m in enumerate(msgs, start=1):
                m["seq"] = i
            changed = True
            logger.info("[chat.migrate] room=%s seq renumbered (n=%s)", r.get("id"), len(msgs))

        # 전체 최대 seq 계산
        local_max = 0
        for m in msgs:
            s = m.get("seq")
            if isinstance(s, int) and s > local_max:
                local_max = s
        if local_max > max_seq_all:
            max_seq_all = local_max

    # 파일 저장(변경이 있을 때만) + 백업
    if changed:
        try:
            if chat_file and os.path.exists(chat_file):
                ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                bak = f"{chat_file}.bak_{ts}"
                shutil.copy2(chat_file, bak)
                logger.info("[chat.migrate] backup created: %s", bak)
        except Exception:
            logger.exception("[chat.migrate] backup failed")

        try:
            # 세션에 반영 후 기존 원자 저장 루틴 사용
            st.session_state.chat_rooms = rooms
            save_chat_rooms()
            logger.info("[chat.migrate] rooms saved (seq migrated)")
        except Exception:
            logger.exception("[chat.migrate] save failed")

    return max_seq_all

# =========================
# 파일 처리
# =========================
def process_file(file, preview: bool = True) -> str:
    """업로드된 파일을 분석해 텍스트/표 미리보기 반환 (견고/다국어/대용량 대응, 캐시 포함)"""
    if file is None:
        return "파일이 업로드되지 않았습니다."

    # 0) 크기 제한 가드
    try:
        file.seek(0, os.SEEK_END)
        size = file.tell()
        file.seek(0)
        if size > MAX_FILE_SIZE_MB * 1024 * 1024:
            return f"[용량 초과] 파일이 {size/1024/1024:.1f}MB 입니다. 최대 {MAX_FILE_SIZE_MB}MB까지 허용됩니다."
    except Exception:
        pass

    # 1) 캐시 키
    try:
        file_hash = _sha256_of_filelike(file)
    except Exception:
        file_hash = None
    ocr_conf = _ocr_conf_tuple()
    cache_key = (file_hash, bool(preview), ocr_conf)

    if file_hash and cache_key in st.session_state.extraction_cache:
        cached = st.session_state.extraction_cache[cache_key]
        if cached is not None:
            return cached

    def _cache_and_return(txt: str) -> str:
        if file_hash:
            st.session_state.extraction_cache[cache_key] = txt
        return txt

    # 2) 메타
    name = getattr(file, "name", "uploaded")
    ext = os.path.splitext(name)[1].lower()
    try:
        pos = file.tell()
        head = file.read(4096)
        file.seek(pos)
    except Exception:
        head = b""
    mime = _sniff_mime(head, getattr(file, "type", "") or "")

    # 3) 타입별
    # PDF
    if ext == ".pdf" or "pdf" in mime:
        try:
            import PyPDF2
            file.seek(0)
            reader = PyPDF2.PdfReader(file)
            parts = []
            for i, page in enumerate(reader.pages):
                try:
                    txt = page.extract_text() or ""
                    if txt.strip():
                        parts.append(txt)
                except Exception:
                    parts.append(f"[{i+1}쪽 텍스트 추출 실패, 이후 계속 진행]")
            text = "\n".join(parts).strip()
            if not text:
                return _cache_and_return("PDF에서 텍스트를 추출하지 못했습니다. (스캔/이미지형 PDF 가능성)")
            return _cache_and_return(_out(text, preview))
        except Exception as e:
            return _cache_and_return(f"[PDF 읽기 오류] {e}")

    # CSV
    if ext == ".csv" or "csv" in mime or (mime.startswith("text/") and ext == ".csv"):
        try:
            import csv
            file.seek(0)
            if _CHARDET_AVAILABLE:
                raw = file.read()
                enc = chardet.detect(raw).get("encoding") or "utf-8"
                text4 = raw.decode(enc, errors="ignore")
                try:
                    dialect = csv.Sniffer().sniff(text4[:4000])
                    sep = dialect.delimiter
                except Exception:
                    sep = None
                if sep is None:
                    df = pd.read_csv(io.BytesIO(raw), encoding=enc, low_memory=False, engine="python", nrows=200, on_bad_lines="skip")
                else:
                    df = pd.read_csv(io.BytesIO(raw), encoding=enc, low_memory=False, sep=sep, nrows=200, on_bad_lines="skip")
            else:
                df = pd.read_csv(file, low_memory=False, nrows=200, on_bad_lines="skip")
            preview_txt = df.head(20).to_string()
            return _cache_and_return("데이터 미리보기(상위 20행):\n" + (_truncate(preview_txt) if preview else preview_txt))
        except Exception as e:
            return _cache_and_return(f"[CSV 읽기 오류] {e}")

    # ---- Excel (.xlsx/.xlsm, 첫 시트 nrows=200) ----
    if ext in (".xlsx", ".xlsm") or "spreadsheet" in mime:
        try:
            file.seek(0)
            try:
                xl = pd.ExcelFile(file, engine="openpyxl")
            except Exception:
                file.seek(0)
                xl = pd.ExcelFile(file)  # 엔진 자동 선택 폴백

            first_name = xl.sheet_names[0]
            first_df = xl.parse(first_name, nrows=200)
            header = f"[시트 {len(xl.sheet_names)}개 중 첫 시트: {first_name}]\n"
            preview_txt = first_df.head(20).to_string()
            return _cache_and_return(header + "데이터 미리보기(상위 20행):\n" + (_truncate(preview_txt) if preview else preview_txt))
        except ImportError:
            return _cache_and_return("`.xlsx`를 읽으려면 `openpyxl`이 필요합니다. 설치: `pip install openpyxl`")
        except Exception as e:
            return _cache_and_return(f"[XLSX 읽기 오류] {e}")

    # Excel (.xls)
    if ext == ".xls":
        try:
            file.seek(0)
            df = pd.read_excel(file, engine="xlrd")
            preview_txt = df.head(20).to_string()
            return _cache_and_return("데이터 미리보기(상위 20행):\n" + (_truncate(preview_txt) if preview else preview_txt))
        except ImportError:
            return _cache_and_return("`.xls`는 `xlrd==1.2.0`이 필요합니다. 설치: `pip install xlrd==1.2.0`")
        except Exception as e:
            return _cache_and_return(f"[XLS 읽기 오류] {e}")

    # TXT
    if ext == ".txt" or mime.startswith("text/"):
        try:
            file.seek(0)
            raw = file.read()

            # 바이너리 의심 탐지 (널바이트/제어문자 비율)
            if isinstance(raw, bytes):
                suspicious = raw.count(b"\x00") > 0 or sum(b < 9 and b not in (10,13) for b in raw[:2048]) > 50
                if suspicious:
                    return _cache_and_return("이 파일은 텍스트가 아니라 바이너리로 보입니다. (.txt 확장자더라도 내용이 이진 데이터일 수 있습니다.)")

            if isinstance(raw, bytes):
                enc = chardet.detect(raw).get("encoding") if _CHARDET_AVAILABLE else "utf-8"
                if not enc:
                    enc = "utf-8"
                text = raw.decode(enc, errors="ignore")
            else:
                text = raw
            return _cache_and_return(_out(text, preview))
        except Exception as e:
            return _cache_and_return(f"[TXT 읽기 오류] {e}")

    # DOCX (표 포함)
    if ext == ".docx" or "officedocument.wordprocessingml.document" in mime:
        try:
            import docx
            file.seek(0)
            d = docx.Document(file)
            paras = [p.text for p in d.paragraphs]
            for t in d.tables:
                for row in t.rows:
                    paras.append(" | ".join(cell.text.strip() for cell in row.cells))
            text = "\n".join([p for p in paras if p is not None]).strip()
            return _cache_and_return(_out(text, preview) if text else "DOCX에서 텍스트를 추출하지 못했습니다.")
        except Exception as e:
            return _cache_and_return(f"[DOCX 읽기 오류] {e}")

    # ---- 이미지 (요약 + OCR) ----
    if ext in {".png", ".jpg", ".jpeg"} or mime.startswith("image/"):
        if not _PIL_AVAILABLE:
            return _cache_and_return("이미지 요약/OCR은 Pillow 설치 시 제공됩니다. 설치: `pip install pillow`")
        try:
            file.seek(0)
            img = Image.open(file)

            # EXIF 회전 보정
            try:
                img = ImageOps.exif_transpose(img)
            except Exception:
                pass

            w, h = img.size
            info = [f"이미지 크기: {w}x{h}px", f"모드: {img.mode}"]

            # EXIF 일부 표시
            exif = {}
            try:
                exif_data = img._getexif() or {}
                from PIL import ExifTags
                tag_map = {ExifTags.TAGS.get(k, k): v for k, v in exif_data.items()}
                for key in ("DateTime", "Model", "Make", "Software"):
                    if key in tag_map:
                        exif[key] = tag_map[key]
            except Exception:
                pass
            if exif:
                info.append("EXIF: " + ", ".join(f"{k}={v}" for k, v in exif.items()))

            # 자동 OCR
            if _TESS_AVAILABLE and st.session_state.get("ocr_auto", False):
                ocr_text = ocr_image_pil(
                    img=img,
                    langs=st.session_state.get("ocr_langs", ["kor", "eng"]),
                    psm=int(st.session_state.get("ocr_psm", 3)),
                    oem=int(st.session_state.get("ocr_oem", 3)),
                    upscale=bool(st.session_state.get("ocr_upscale", True)),
                    binarize=bool(st.session_state.get("ocr_binarize", True)),
                    denoise=bool(st.session_state.get("ocr_denoise", False)),
                )
                if ocr_text:
                    info.append("\n[OCR 결과]\n" + (_truncate(ocr_text) if preview else ocr_text))
                else:
                    info.append("\n[OCR 결과] 추출된 텍스트가 없습니다.")

            return _cache_and_return("\n".join(info))

        except Image.DecompressionBombError:
            # PIL 안전장치에 걸릴 때
            return _cache_and_return(
                "이미지 해상도가 너무 커서 안전상 처리가 중단되었습니다. "
                "(DecompressionBombError) — 해상도를 낮춰 다시 시도해 주세요."
            )
        except Exception as e:
            return _cache_and_return(f"[이미지 분석 오류] {e}")
    # ---- 기타(미지원 타입) ----
    return _cache_and_return(f"지원하지 않거나 분석 대상이 아닙니다: {name} ({mime})")
# =========================
def process_file_with_cache(file, preview: bool = True) -> str:
    try:
        file_hash = _sha256_of_filelike(file)
    except Exception:
        file_hash = uuid.uuid4().hex
    ocr_conf = _ocr_conf_tuple()
    cache_key = (file_hash, bool(preview), ocr_conf)

    cached = st.session_state.extraction_cache.get(cache_key)
    if cached:
        return cached

    out = process_file(file, preview=preview)
    st.session_state.extraction_cache[cache_key] = out
    # 상한선 500개 정도로 관리
    MAX_CACHE_ITEMS = 500
    cache = st.session_state.extraction_cache
    if len(cache) > MAX_CACHE_ITEMS:
        # FIFO 식 단순 제거
        try:
            first_key = next(iter(cache.keys()))
            cache.pop(first_key, None)
        except Exception:
            pass
    return out

# =========================
# 스트리밍 응답
# 캐시 관리 버튼 + 캐시 상한
# =========================
def stream_and_append_assistant(
    messages_for_ai: List[Dict[str, str]],
    room: Dict[str, Any],
    *,
    temperature: float = 0.7,
    max_retries: int = 2,   # 스트리밍 중 예외 시 '전체 재시도' 횟수
    backoff: float = 1.5,
    history_channel: str = "gen_messages",
) -> str:
    model_id = st.session_state.get("selected_model") or "local-model"
    collected: List[str] = []
    final_text = ""
    finish_reason = None
    assistant_time = ""

    import time as _time, random as _random

    with st.chat_message("assistant"):
        container = st.empty()
        caption_slot = st.empty()
        last_err = None

        for attempt in range(max_retries + 1):
            collected.clear()
            last_render = 0.0
            render_interval = 0.04  # 40ms

            try:
                # ✅ 요청 보호막 + 회로차단 포함 (스트림 모드)
                response_stream = call_chat_protected(
                    messages=messages_for_ai,
                    model=model_id,
                    temperature=temperature,
                    stream=True,
                    # 필요 시 per-call 오버라이드 사용:
                    # timeout_s=LLM_TIMEOUT_S,
                    # max_retry=LLM_MAX_RETRY,
                    # backoff_seq=LLM_BACKOFF_SEQ,
                    allow_when_open=True,  # 회로차단 open이더라도 half-open 1회 허용
                )

                for chunk in response_stream:
                    delta = getattr(chunk.choices[0].delta, "content", None) or ""
                    if delta:
                        collected.append(delta)

                    fr = getattr(chunk.choices[0], "finish_reason", None)
                    if fr:
                        finish_reason = fr

                    now_ts = _time.perf_counter()
                    if (now_ts - last_render) >= render_interval and collected:
                        container.markdown("".join(collected) + "▌")
                        last_render = now_ts

                final_text = "".join(collected).strip()

                if finish_reason in ("length",) and final_text:
                    final_text += "\n\n_※ 응답이 길이 제한으로 중단된 것 같습니다._"

                assistant_time = make_ts()
                container.markdown(final_text or "_(빈 응답)_")
                caption_slot.caption(assistant_time)

                last_err = None
                break  # 성공

            except Exception as e:
                last_err = e
                if attempt < max_retries:
                    sleep_s = (backoff ** attempt) + _random.uniform(0, 0.3)
                    container.markdown(f"_연결 이슈로 재시도 중... ({attempt+1}/{max_retries})_")
                    _time.sleep(sleep_s)
                else:
                    msg = f"{type(e).__name__}: {str(e)}"
                    if len(msg) > 400:
                        msg = msg[:400] + " …"
                    final_text = f"⚠️ 모델 스트리밍 중 오류: {msg}"
                    assistant_time = make_ts()
                    container.markdown(final_text)
                    caption_slot.caption(assistant_time)


    # ✅ 기존 정렬/앵커/중복제거 로직과 100% 호환되게 저장
    room.setdefault("messages", []).append({
        "id": str(uuid.uuid4()),
        "seq": _next_seq(),
        "role": "assistant",
        "content": final_text,
        "time": assistant_time or make_ts(),
        **_message_meta("chat"),
    })
    _sync_room_meta(room, materialize=True)

    log.info(
        "[chat.assistant] %s answer_len=%s history_channel=%s",
        _chat_log_kv(room),
        len(final_text or ""),
        history_channel,
    )

    # ✅ 2-채널 히스토리에도 저장(LLM 히스토리 분리용)
    room.setdefault(history_channel, []).append({
        "id": str(uuid.uuid4()),
        "seq": _next_seq(),
        "role": "assistant",
        "content": final_text,
        "time": assistant_time or make_ts(),
        **_message_meta("chat"),
    })

    save_chat_rooms()

    return final_text

# =========================
# 세션 초기화
# =========================
_reset_chat_session_when_user_changed()

if "chat_rooms" not in st.session_state:
    st.session_state.chat_rooms = load_chat_rooms()

    # ✅ 2-채널 히스토리 마이그레이션(기존 rooms.json 호환)
    # - messages는 UI 렌더용(기존 유지)
    # - gen_messages는 일반 대화 히스토리(LLM 일반 질문에 사용)
    # - sims_messages는 SIMS 대화 히스토리(LLM/SIMS 질문에 사용)
    try:
        touched = False
        for r in (st.session_state.chat_rooms or []):
            if "gen_messages" not in r:
                # 초기에는 기존 messages를 gen_messages로 복사(안전한 기본값)
                r["gen_messages"] = list(r.get("messages") or [])
                touched = True
            if "sims_messages" not in r:
                r["sims_messages"] = []
                touched = True
        if touched:
            save_chat_rooms()
            log.info("[chat.migrate] added gen_messages/sims_messages to rooms %s", _chat_log_kv())
    except Exception:
        log.exception("[chat.migrate] failed to add history channels")

    # ✅ 기존 저장 방 보정
    # - 메시지가 이미 있는 방은 더 이상 임시방이 아니므로 auto_created=False
    # - 방 최상위 company_id/db_name이 예전 회사로 남아 있으면 마지막 메시지 기준으로 보정
    try:
        touched = False

        for r in (st.session_state.chat_rooms or []):
            if not isinstance(r, dict):
                continue

            msgs = r.get("messages") or []

            if msgs:
                if r.get("auto_created") is not False:
                    r["auto_created"] = False
                    touched = True

                last_meta = _room_meta_from_last_message(r)
                for k, v in last_meta.items():
                    if r.get(k) != v:
                        r[k] = v
                        touched = True

            else:
                if "auto_created" not in r:
                    r["auto_created"] = False
                    touched = True

        if touched:
            save_chat_rooms()
            log.info("[chat.migrate] materialized rooms and synced room meta from last message %s", _chat_log_kv())

    except Exception:
        log.exception("[chat.migrate] failed to materialize room meta")


    # ✅ seq 없는 옛 메시지 마이그레이션(필요 시에만 백업+저장)
    try:
        max_seq = migrate_rooms_seq_if_needed(
            st.session_state.chat_rooms,
            chat_file=CHAT_FILE,
            logger=log,
        )
        # ✅ __seq 카운터도 동기화(새 메시지가 맨 위로 올라가는 현상 방지)
        cur = st.session_state.get("__seq", 0)
        if not isinstance(cur, int):
            cur = 0
        st.session_state["__seq"] = max(cur, int(max_seq or 0))
        log.debug("[chat] synced __seq=%s (migrated max_seq=%s)", st.session_state["__seq"], max_seq)
    except Exception:
        log.exception("[chat] seq migration failed")

    # ✅ 메시지 seq 카운터 동기화(중요!)
    # 기존 rooms.json에 저장된 seq 최대값보다 __seq가 작으면,
    # 새로 입력한 메시지가 정렬에서 맨 위로 올라가는 현상이 생김.
    try:
        max_seq = 0
        for r in (st.session_state.chat_rooms or []):
            for m in (r.get("messages") or []):
                s = m.get("seq")
                if isinstance(s, (int, float)):
                    if int(s) > max_seq:
                        max_seq = int(s)
        cur = st.session_state.get("__seq", 0)
        if not isinstance(cur, (int, float)):
            cur = 0
        st.session_state["__seq"] = max(int(cur), max_seq)
        log.debug("[chat] synced __seq=%s (max_seq=%s)", st.session_state["__seq"], max_seq)
    except Exception:
        log.exception("[chat] failed to sync __seq from chat_rooms")

# 로그인 직후 기본 입력 대상은 기존 방이 아니라 "새 대화 대기방"이다.
# 기존 방은 목록에만 표시하고, 사용자가 직접 선택해야 해당 방에 이어 쓴다.
st.session_state.setdefault("chat_rooms", [])
if "current_room" not in st.session_state or not st.session_state.current_room:
    _select_pending_new_room()


if "selected_model" not in st.session_state:
    st.session_state.selected_model = get_models()[0]

current_room = _get_current_room_or_pending()

_render_login_greeting_banner()
_consume_company_change_notice(current_room)

#=========================
# SIMS 사이드바 옵션 fragment
# =========================
# SIMS 패널 열기/카테고리/작업 선택 시 전체 앱 rerun 방지 + 기존 채팅 표 재렌더링 방지
# 실제 [SIMS 실행] 버튼을 누를 때만 전체 앱 rerun
# SIMS 옵션은 st.session_state["__sims_selected"]에 저장되고,
# [SIMS 실행] 클릭 시 "__sims_selected_snapshot"으로 복사되어 메인 화면에서 참조된다.
@st.fragment
def _render_sims_sidebar_fragment() -> None:
    """
    SIMS 사이드바 옵션 전용 fragment.

    목적:
    - SIMS 패널 열기/카테고리/작업 선택 시 전체 앱 rerun 방지
    - 기존 채팅 표, 특히 974 x 97 styled table 재렌더링 방지
    - 실제 [SIMS 실행] 버튼을 누를 때만 전체 앱 rerun
    """
    st.markdown("---")
    st.markdown("### 🧩 SIMS 모드")

    # 조회 결과를 채팅창으로 보낸 뒤에도 SIMS 패널은 열린 상태를 유지한다.
    # 이전 v2 자동 닫기 플래그가 남아 있어도 토글을 끄지 않고 플래그만 소비한다.
    st.session_state.pop("__sims_close_after_push", None)

    # 조회 결과 push 직후 다음 rerun에서 토글이 OFF로 보이는 현상을 보정한다.
    # 주의: __sims_open은 toggle 위젯 key라서, 반드시 st.toggle 생성 전에만 보정한다.
    if st.session_state.pop("__sims_keep_open_after_push", False):
        st.session_state["__sims_open"] = True
        st.session_state["__sims_panel_active"] = True



    st.session_state.setdefault("__sims_open", False)
    st.session_state.setdefault("__sims_q", "")
    st.session_state.setdefault("__sims_run", False)

    sims_panel_open = st.toggle("SIMS 패널 열기", key="__sims_open")

    if not sims_panel_open:
        # fragment 안의 토글만 rerun되면 메인 영역의 기존 패널이 화면에 남을 수 있다.
        # 기존 패널이 활성 상태였으면 app 전체 rerun을 1회 요청해 즉시 닫힘을 반영한다.
        need_app_rerun = bool(
            st.session_state.get("__sims_panel_active")
            or st.session_state.get("__sims_force_open")
            or st.session_state.get("__sims_open_ui")
            or st.session_state.get("__sims_run_flag")
            or st.session_state.get("__sims_inner_submit")
        )

        # 사용자가 토글을 끄면 프로그램 강제 열림 상태보다 사용자 닫기가 우선이다.
        # 단, 현재표/다운로드/LLM용 내부 컨텍스트는 유지한다.
        st.session_state["__sims_force_open"] = False
        st.session_state["__sims_open_ui"] = False
        st.session_state["__sims_panel_active"] = False
        st.session_state["__sims_run_flag"] = False
        st.session_state["__sims_inner_submit"] = False

        if need_app_rerun:
            log.info("[ui.fragment] SIMS panel close requested → app rerun")
            try:
                st.rerun(scope="app")
            except TypeError:
                st.rerun()

        return

    with st.expander("SIMS 옵션", expanded=True):
        st.text_input("쿼리/키워드", key="__sims_q", label_visibility="collapsed")

        sims_mode_selector(key="__sims_mode")

        selected = render_sims_sidebar_controls()
        if not selected:
            try:
                from app.ui.sims_entry import render_sims_sidebar_controls as _render_sims_sidebar_controls
                selected = _render_sims_sidebar_controls(parent=st)
            except Exception:
                selected = {}

        if selected:
            new_selected = dict(selected)
            # 중요:
            # 카테고리/작업 선택만으로는 기존 SIMS 결과/채팅 표를 지우지 않는다.
            # 실제 실행은 [SIMS 실행] 버튼이 담당한다.
            st.session_state["__sims_selected"] = new_selected
            st.session_state["__sims_selected_snapshot"] = dict(new_selected)
            st.caption("선택됨. ‘SIMS 작업 열기’를 눌러 선택한 화면/폼을 여세요.")
        else:
            st.session_state.pop("__sims_selected", None)
            st.session_state["__sims_selected_snapshot"] = {}

        clicked = st.button(
            "SIMS 작업 열기",
            use_container_width=True,
            type="primary",
            key="__sims_run_btn",
        )

        if clicked:
            st.session_state["__sims_force_open"] = True
            st.session_state["__sims_panel_active"] = True
            st.session_state["__sims_run_flag"] = True
            st.session_state["__sims_inner_submit"] = False

            st.session_state.setdefault("__sims_form_id", 0)
            st.session_state.setdefault("__sims_run_seq", 0)

            st.session_state["__sims_form_id"] += 1
            st.session_state["__sims_run_seq"] += 1
            st.session_state["__sims_selected_snapshot"] = dict(
                st.session_state.get("__sims_selected") or {}
            )

            log.info(
                "[ui.fragment] SIMS 작업 열기 클릭됨 — open=%s, form_id=%s, run_seq=%s, selected=%r",
                st.session_state.get("__sims_open"),
                st.session_state["__sims_form_id"],
                st.session_state["__sims_run_seq"],
                st.session_state.get("__sims_selected"),
            )
            st.toast("🧩 SIMS 작업 화면을 엽니다.")

            # 실제 조회/메인 결과 렌더링은 전체 앱 rerun이 필요하다.
            st.rerun()

        def _sims_reset() -> None:
            st.session_state["__sims_reset_requested"] = True
            log.info("[sims.reset] requested=True (from sidebar fragment)")
            st.toast("SIMS 옵션이 초기화되었습니다.", icon="🧹")

            # 옵션 초기화는 메인 상태까지 정리해야 하므로 전체 앱 rerun.
            st.rerun()

        st.button("옵션 초기화", use_container_width=True, on_click=_sims_reset)

# =========================
# 사이드바 (통합본)
# =========================
with st.sidebar:
    # 상단 배너
    st.markdown("""
    <div style="
        background-color:#1E3A8A;
        color:white;
        padding:12px;
        border-radius:8px;
        text-align:center;
        font-size:18px;
        font-weight:bold;
        margin-bottom:10px;">
        🤖 SSAI LM Studio Chatbot
    </div>
    """, unsafe_allow_html=True)

    # 1) 모델 선택
    st.markdown("## 🧠 모델 선택")
    try:
        models = get_models() or ["local-model"]
    except Exception:
        models = ["local-model"]
    current = st.session_state.get("selected_model")
    default_idx = models.index(current) if current in models else 0
    st.session_state.selected_model = st.selectbox("모델", options=models, index=default_idx)

    # 2) 채팅방 관리
    # =========================
    # 💬 사이드바: 2) 채팅방 (핵심 고정 + 추가 옵션 접기)
    # =========================
    st.markdown("---")
    st.markdown("## 💬 채팅방")

    ss = st.session_state

    # ── 안전 기본값 ───────────────────────────────────────────
    # 로그인 직후 채팅방이 없으면 임시 자동 대화방 1개만 만든다.
    # 여기서는 저장하지 않는다.
    # 첫 채팅 / SIMS 결과 / 회사 변경 안내가 들어갈 때 save_chat_rooms()로 실제 저장된다.
    ss.setdefault("chat_rooms", [])

    if "current_room" not in ss or not ss.current_room:
        _select_pending_new_room()
    else:
        _get_current_room_or_pending()

    # 추가 옵션 상태(필터/페이지) — 목록은 이 상태를 기준으로 보여줌
    ss.setdefault("__room_filter", "")
    ss.setdefault("__room_page", 0)
    PAGE_SIZE = 20

    # ── 필터/페이지 계산(옵션 패널을 안 열어도 상태는 반영됨) ──────────
    q = (ss.__room_filter or "").strip().lower()
    rooms_filtered = [r for r in ss.chat_rooms if q in str(r.get("name") or "").lower()]

    # 로그인 직후 새 대화 대기방이 첫 화면에 보이도록 현재 방이 포함된 페이지를 우선 표시한다.
    current_id = str(ss.get("current_room") or "").strip()
    current_idx = next(
        (idx for idx, room in enumerate(rooms_filtered) if str(room.get("id") or "") == current_id),
        -1,
    )

    total = len(rooms_filtered)
    max_page = max(0, (total - 1) // PAGE_SIZE)

    if current_idx >= 0:
        ss.__room_page = max(0, min(current_idx // PAGE_SIZE, max_page))
    else:
        ss.__room_page = max(0, min(ss.__room_page, max_page))

    start = ss.__room_page * PAGE_SIZE
    end = start + PAGE_SIZE
    view = rooms_filtered[start:end]

    # 필터 때문에 현재 방이 view에 없으면 첫 줄에 현재 방을 임시로 보여준다.
    # 라디오 위젯이 첫 번째 기존 방을 자동 선택해 current_room을 바꾸는 것을 막기 위함이다.
    if current_id and current_id not in [str(r.get("id") or "") for r in view]:
        current_room_for_list = next(
            (r for r in ss.chat_rooms if str(r.get("id") or "") == current_id),
            None,
        )
        if isinstance(current_room_for_list, dict):
            view = [current_room_for_list] + view[: max(0, PAGE_SIZE - 1)]

    # 최조 대화 등록 하기 전
    id_to_name = {}

    for r in ss.chat_rooms:
        rid = r.get("id")
        if not rid:
            continue

        if _is_empty_auto_room(r):
            id_to_name[rid] = "➕ 새 대화 입력 대기 (저장 전)"
        else:
            id_to_name[rid] = r.get("name") or "업무 대화"


    # ── (항상 보임) 핵심: 새 대화 + 목록 + 이름 변경 ────────────────
    def _new_room():
        # 사용자가 직접 누른 경우는 임시방이 아니라 정식 새 채팅방이다.
        new_room = _make_chat_room(auto_created=False)

        ss.chat_rooms.append(new_room)
        ss.current_room = new_room["id"]
        ss["__room_prev_selected"] = new_room["id"]
        ss["__room_rename_buf"] = new_room.get("name") or ""
        ss["__room_filter"] = ""

        log.info("[chat.room] new room created %s", _chat_log_kv(new_room))
        save_chat_rooms()
        st.rerun()

    st.button("➕ 새 대화 시작", use_container_width=True, key="__room_new_btn", on_click=_new_room)

    # 목록(라디오)
    options_ids = [r["id"] for r in view]
    picked_index = options_ids.index(ss.current_room) if (options_ids and ss.current_room in options_ids) else 0
    # Streamlit radio는 같은 key를 쓰면 이전 선택값을 계속 기억한다.
    # 로그인 직후 기본값이 기존 방으로 되돌아가지 않도록 현재 방 ID 기반 key를 사용한다.
    room_radio_key = f"__room_pick_radio_{str(ss.current_room or '')[:12]}"
    picked = st.radio(
        "채팅방 목록",
        options=options_ids,
        format_func=lambda rid: id_to_name.get(rid, rid),
        index=picked_index if options_ids else 0,
        label_visibility="collapsed",
        key=room_radio_key,
    )
    if picked and picked != ss.current_room:
        old_room_id = ss.current_room
        ss.current_room = picked

        removed_empty_pending = _drop_empty_auto_rooms(keep_room_id=picked)

        log.info(
            "[chat.room] selected %s old_room_id=%s new_room_id=%s removed_empty_pending=%s",
            _chat_log_kv(_get_current_room_or_pending()),
            old_room_id,
            picked,
            removed_empty_pending,
        )

        save_chat_rooms()
        st.rerun()

    # 이름 변경(경고 방지: 선택 바뀌면 버퍼 초기화)
    cur_name = id_to_name.get(ss.current_room, "")

    ss.setdefault("__room_prev_selected", ss.current_room)
    ss.setdefault("__room_prev_name", cur_name)
    ss.setdefault("__room_rename_buf", cur_name)

    # 선택 방이 바뀌었거나, 같은 방의 실제 이름이 자동 변경된 경우 입력창도 갱신
    if (
        ss.__room_prev_selected != ss.current_room
        or ss.__room_prev_name != cur_name
    ):
        ss["__room_prev_selected"] = ss.current_room
        ss["__room_prev_name"] = cur_name
        ss["__room_rename_buf"] = cur_name

    st.markdown("**✏️ 채팅방 이름 변경**")
    st.text_input("새 이름", key="__room_rename_buf", label_visibility="collapsed")

    def _apply_rename():
        new_name = (ss.__room_rename_buf or "").strip()
        if new_name:
            for r in ss.chat_rooms:
                if r["id"] == ss.current_room:
                    r["name"] = new_name
                    r["name_auto"] = False
                    r["title_source"] = "manual"
                    break                
            save_chat_rooms()
        st.rerun()

    st.button("이름 적용", use_container_width=True, key="__room_rename_apply", on_click=_apply_rename)

    # 간단한 현황(현재 페이지 범위/총개수)
    if total:
        st.caption(f"{start+1}–{min(end, total)} / 총 {total}개")

    # ── (접는) 추가 옵션: 필터/페이지/삭제/내보내기 ────────────────
    with st.expander("채팅방 관리", expanded=False):
        # 필터
        st.text_input("이름 검색", key="__room_filter", placeholder="채팅방 이름 필터")

        # 페이지네이션
        c_prev, c_page, c_next = st.columns([1, 2, 1])
        with c_prev:
            st.button("◀ ", use_container_width=True,
                    disabled=ss.__room_page <= 0,
                    key="__room_prev",
                    on_click=lambda: ss.update(__room_page=max(0, ss.__room_page-1)) or st.rerun())
        with c_page:
            st.caption(f"페이지 {ss.__room_page+1} / {max_page+1}")
        with c_next:
            st.button(" ▶", use_container_width=True,
                    disabled=ss.__room_page >= max_page,
                    key="__room_next",
                    on_click=lambda: ss.update(__room_page=min(max_page, ss.__room_page+1)) or st.rerun())

        st.write("")

        # 삭제(확인 단계)
        ss.setdefault("__room_delete_ask", False)
        def _ask_delete(): ss.update(__room_delete_ask=True)
        st.button("🗑️ 현재 채팅방 삭제", use_container_width=True, key="__room_delete_btn", on_click=_ask_delete)

        if ss.__room_delete_ask:
            st.warning("정말로 이 채팅방을 삭제할까요? (되돌릴 수 없음)", icon="⚠️")
            c_ok, c_cancel = st.columns(2)

            def _do_delete():
                target = ss.current_room
                ss.chat_rooms = [r for r in ss.chat_rooms if r["id"] != target]
                if not ss.chat_rooms:
                    ss.chat_rooms.append(_make_chat_room(auto_created=True))                    
                ss.current_room = ss.chat_rooms[0]["id"]
                ss["__room_delete_ask"] = False
                save_chat_rooms()
                st.rerun()

            with c_ok:
                st.button("삭제 확인", type="primary", use_container_width=True,
                        key="__room_delete_confirm", on_click=_do_delete)
            with c_cancel:
                st.button("취소", use_container_width=True,
                        key="__room_delete_cancel",
                        on_click=lambda: ss.update(__room_delete_ask=False))

        st.markdown("---")
        st.markdown("**내보내기**")

        # 전체 ZIP
        buf_zip = io.BytesIO()
        with zipfile.ZipFile(buf_zip, "w", zipfile.ZIP_DEFLATED) as zipf:
            for room in ss.chat_rooms:
                safe = re.sub(r"[^a-zA-Z0-9가-힣]+", "_", room["name"])
                zipf.writestr(f"{safe}.json", json.dumps(_json_sanitize(room), ensure_ascii=False, indent=2))
        buf_zip.seek(0)
        st.download_button("💾 모든 대화 ZIP", data=buf_zip, file_name="chat_rooms.zip",
                        mime="application/zip", use_container_width=True, key="__dl_all_rooms_zip")

        # 현재 JSON
        cur_room = next((r for r in ss.chat_rooms if r["id"] == ss.current_room), ss.chat_rooms[0])
        cur_json = json.dumps(cur_room, ensure_ascii=False, indent=2, default=_json_default)
        cur_json = json.dumps(cur_room, ensure_ascii=False, indent=2).encode("utf-8")
        st.download_button("⬇️ 현재 대화(JSON)", data=cur_json,
                        file_name=f"{re.sub(r'[^a-zA-Z0-9가-힣]+','_', cur_room['name'])}.json",
                        mime="application/json", use_container_width=True, key="__dl_current_json")
        # --- 현재 대화 Markdown / HTML 내보내기 ---
        def _room_to_markdown(room: dict) -> str:
            lines = [f"# 대화방: {room.get('name','')}", ""]
            for m in room.get("messages", []):
                role = m.get("role","assistant")
                t    = m.get("time","")
                content = (m.get("content") or "").strip()
                lines.append(f"**{role}** ({t})\n\n{content}\n")
            return "\n".join(lines)

        def _room_to_html(room: dict) -> str:
            # 간단 변환(정확한 렌더가 필요하면 markdown 라이브러리 사용 권장)
            md = _room_to_markdown(room)
            html = md.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            html = html.replace("\n", "<br/>").replace("**", "")
            return f"<html><body>{html}</body></html>"

        md_bytes   = _room_to_markdown(cur_room).encode("utf-8")
        html_bytes = _room_to_html(cur_room).encode("utf-8")

        st.download_button(
            "📝 현재 대화 (Markdown)",
            data=md_bytes,
            file_name=f"{re.sub(r'[^a-zA-Z0-9가-힣]+','_', cur_room['name'])}.md",
            mime="text/markdown",
            use_container_width=True,
            key="__dl_current_md",
        )
        st.download_button(
            "🧾 현재 대화 (HTML)",
            data=html_bytes,
            file_name=f"{re.sub(r'[^a-zA-Z0-9가-힣]+','_', cur_room['name'])}.html",
            mime="text/html",
            use_container_width=True,
            key="__dl_current_html",
        )
    # =========================
    # 🔎 사이드바의 3) "채팅 검색" 섹션 (옵션만) ===
    # =========================
    st.markdown("---")
    st.markdown("## 🔎 채팅 검색")
    
    #  (A) 기본값 방어: 위젯 렌더 전에만 setdefault
    st.session_state.setdefault("__search_open", False)
    st.session_state.setdefault("__search_q", "")
    st.session_state.setdefault("__search_roles", ["user", "assistant"])
    st.session_state.setdefault("__search_regex", False)
    st.session_state.setdefault("__search_case", False)
    st.session_state.setdefault("__search_limit", 50)

    # (B) 초기화 훅 (위젯 렌더 전에만 값 변경)
    if st.session_state.pop("__search_reset", False):
        st.session_state["__search_q"] = ""
        st.session_state["__search_roles"] = ["user", "assistant"]
        st.session_state["__search_regex"] = False
        st.session_state["__search_case"] = False
        st.session_state["__search_limit"] = 50

    # (C) 토글: 반환값 재대입 금지! value=도 금지!
    st.toggle("검색 패널 열기", key="__search_open")

    # (D) 옵션 UI: value= 전부 제거, key만 유지
    if st.session_state.get("__search_open"):
        with st.expander("검색 옵션", expanded=False):
            # (선택) Enter 즉시 반영 보장용 no-op 콜백
            def _search_touch(): 
                pass

            st.text_input("키워드",
                        key="__search_q",
                        label_visibility="collapsed",
                        on_change=_search_touch,
            )  # <- 선택

            st.checkbox("대/소문자", key="__search_case")
            st.checkbox("정규식",   key="__search_regex")

            st.multiselect("대상 역할",
                        options=["user", "assistant"],
#                        default=st.session_state["__search_roles"],
                        key="__search_roles",
            )

            st.slider("표시 개수", min_value=5, max_value=200, step=5,
                    key="__search_limit",
            )

            if st.session_state["__search_regex"]:
                st.caption(r"예: \bERROR\b  /  ^Hello")

            # 초기화 버튼 (on_click에서 플래그만 세우고 rerun)
            def _queue_search_reset():
                st.session_state["__search_reset"] = True
                st.rerun()

            st.button("검색 초기화", key="__search_clear",
                    use_container_width=True, on_click=_queue_search_reset)

    # =========================
    # 4) SIMS 모드
    # =========================
    _render_sims_sidebar_fragment()

    # =========================
    # 🖼️ 5) OCR 옵션 (정리판: 토글+익스펜더)
    # =========================
    st.markdown("---")
    st.markdown("## 🖼️ OCR 옵션")

    # (A) 기본값(위젯 렌더 전 setdefault)
    st.session_state.setdefault("__ocr_open", False)
    st.session_state.setdefault("__ocr_auto", True)
    st.session_state.setdefault("__ocr_langs", ["kor", "eng"])
    st.session_state.setdefault("__ocr_psm", 3)           # 0~13
    st.session_state.setdefault("__ocr_oem", 3)           # 0:legacy, 1:LSTM, 2:both, 3:auto
    st.session_state.setdefault("__ocr_upscale", True)
    st.session_state.setdefault("__ocr_binarize", True)
    st.session_state.setdefault("__ocr_denoise", False)
    st.session_state.setdefault("__ocr_tesseract_cmd", "")

    # (B) 유틸: 초기화/프리셋
    def _ocr_reset():
        st.session_state["__ocr_auto"] = True
        st.session_state["__ocr_langs"] = ["kor", "eng"]
        st.session_state["__ocr_psm"] = 3
        st.session_state["__ocr_oem"] = 3
        st.session_state["__ocr_upscale"] = True
        st.session_state["__ocr_binarize"] = True
        st.session_state["__ocr_denoise"] = False
        st.session_state["__ocr_tesseract_cmd"] = ""
        st.rerun()

    def _ocr_preset(name: str):
        if name == "kor+eng":
            st.session_state["__ocr_langs"] = ["kor", "eng"]
            st.session_state["__ocr_psm"] = 3
            st.session_state["__ocr_oem"] = 3
        elif name == "eng":
            st.session_state["__ocr_langs"] = ["eng"]
            st.session_state["__ocr_psm"] = 3
            st.session_state["__ocr_oem"] = 3
        elif name == "jpn_receipt":
            st.session_state["__ocr_langs"] = ["jpn"]
            st.session_state["__ocr_psm"] = 4
            st.session_state["__ocr_oem"] = 3
        st.rerun()

    # (C) 패널 토글
    st.toggle("OCR 옵션 열기", key="__ocr_open")

    # (D) 옵션 UI (토글 ON일 때만)
    if st.session_state["__ocr_open"]:

        with st.expander("기본", expanded=True):
            st.checkbox("이미지 자동 OCR", key="__ocr_auto")
            st.multiselect(
                "언어", 
                options=["kor", "eng", "jpn", "chi_sim"],
                key="__ocr_langs",  # ← default/value 없이 key만!
                help="여러 언어를 선택하면 'kor+eng'처럼 함께 인식합니다."
            )

            st.markdown("**프리셋**")
            st.button("🇰🇷+🇺🇸 일반",   use_container_width=True, on_click=_ocr_preset, args=("kor+eng",))
            st.button("🇺🇸 영어전용",   use_container_width=True, on_click=_ocr_preset, args=("eng",))
            st.button("🇯🇵 영수증",     use_container_width=True, on_click=_ocr_preset, args=("jpn_receipt",))


        with st.expander("고급", expanded=False):
            st.number_input("PSM (페이지 세그먼트)", min_value=0, max_value=13, key="__ocr_psm")
            st.selectbox("OEM (엔진)", options=[0,1,2,3], key="__ocr_oem")

            st.checkbox("고해상도 스케일업", key="__ocr_upscale")
            st.checkbox("그레이스케일+이진화", key="__ocr_binarize")
            st.checkbox("노이즈 제거(미디안)", key="__ocr_denoise")

            st.text_input(
                "Tesseract 경로(옵션, Windows용)",
                key="__ocr_tesseract_cmd",
                placeholder=r"C:\Program Files\Tesseract-OCR\tesseract.exe"
            )

            # 경로 적용 (가능할 때만)
            try:
                if (globals().get("_TESS_AVAILABLE") and st.session_state["__ocr_tesseract_cmd"].strip()):
                    import pytesseract
                    pytesseract.pytesseract.tesseract_cmd = st.session_state["__ocr_tesseract_cmd"].strip()
            except Exception:
                pass
            st.markdown("**동작**")
            st.button("초기화", use_container_width=True, on_click=_ocr_reset)
            st.button("추천(🇰🇷+🇺🇸)", use_container_width=True, on_click=_ocr_preset, args=("kor+eng",))
            st.caption("변경 즉시 적용됩니다.")

        # 상태/요약
        _PIL = bool(globals().get("_PIL_AVAILABLE"))
        _TESS = bool(globals().get("_TESS_AVAILABLE"))
        _CHAR = bool(globals().get("_CHARDET_AVAILABLE"))
        _MAG  = bool(globals().get("_MAGIC_AVAILABLE"))
        st.caption(f"🔧 deps — PIL:{_PIL} · Tesseract:{_TESS} · chardet:{_CHAR} · magic:{_MAG}")

    #  사용자에 따른 감추기
    if _can_show_admin_diagnostics_sidebar():

        # 6) 환경 요약/헬스체크/진단
        def _mask(v: str | None, show: int = 3) -> str:
            if not v: return "(미설정)"
            v = str(v); return v if len(v) <= show else v[:show] + "…" + f"({len(v)} chars)"

        st.markdown("---")
        st.markdown("## ⚙️ 환경 요약")
        st.write({
            "LMSTUDIO_BASE_URL": os.getenv("LMSTUDIO_BASE_URL"),
            "MSSQL_SERVER":      os.getenv("MSSQL_SERVER"),
            "MSSQL_DATABASE":    os.getenv("MSSQL_DATABASE"),
            "LMSTUDIO_API_KEY":  _mask(os.getenv("LMSTUDIO_API_KEY", "")),
        })

        st.markdown("---")
        st.markdown("## 🩺 Health Check")
        if st.button("Check LLM & DB", use_container_width=True, key="btn_health_check"):
            # LLM
            try:
                from app.services.llm_health import check_llm
                llm = check_llm()
                if isinstance(llm, dict) and llm.get("ok"):
                    st.success(f"LLM OK ({llm.get('count', 0)} models)")
                    with st.expander("Models (Top 20)", expanded=False):
                        st.write(llm.get("models", []))
                        st.caption(f"Base URL: {llm.get('base_url')}")
                else:
                    st.error(f"LLM 오류: {getattr(llm, 'get', lambda *_: None)('error') or llm}")
            except Exception as e:
                st.error(f"LLM 연결 오류: {e}")

            # DB
            try:
                from app.db.mssql_client import health_check, list_tables, search_columns
                info = health_check()

                # ✅ health_check() 반환 형식 보정(dict / bool 모두 지원)
                if isinstance(info, dict):
                    ok = info.get("ok", False)
                    if ok:
                        st.success(f"DB OK: {info.get('database')}")
                        with st.expander("DB Version", expanded=False):
                            st.code(info.get("version", ""), language="text")
                        with st.expander("샘플 테이블", expanded=False):
                            st.dataframe(list_tables(20))
                        with st.expander("컬럼 탐색(샘플: 'Code')", expanded=False):
                            st.dataframe(search_columns("Code", 50).head(50))
                    else:
                        st.error(f"DB 오류: {info.get('error')}")
                elif isinstance(info, bool):
                    # 예전 스타일: True/False 만 리턴하는 health_check()
                    if info:
                        st.success("DB OK (health_check() → True)")
                        with st.expander("샘플 테이블", expanded=False):
                            st.dataframe(list_tables(20))
                    else:
                        st.error("DB 오류: health_check() → False")
                else:
                    st.error(f"DB health_check() 반환 형식을 알 수 없음: {info!r}")

            except Exception as e:
                st.error(f"DB 체크 예외: {e}")

        st.markdown("---")
        # 7) 디버그
        # =========================
        # 7) 디버그 / 환경 진단 패널
        # =========================
        DEBUG = st.toggle("🪲 Debug 모드", value=False, key="__debug_toggle")

        if DEBUG:
            st.markdown("## 🪲 Debug / 환경 진단")

            # ---- 1) LLM / 세션 상태 ----
            with st.expander("1️⃣ LLM / 세션 상태", expanded=True):
                chat_rooms = st.session_state.get("chat_rooms", [])
                selected_model = st.session_state.get("selected_model")
                current_room_id = st.session_state.get("current_room_id")
                messages_count = sum(len(r.get("messages", [])) for r in chat_rooms)

                st.json(
                    {
                        "selected_model": selected_model,
                        "current_room_id": current_room_id,
                        "chat_rooms": len(chat_rooms),
                        "messages_total": messages_count,
                    },
                    expanded=False,
                )

            # ---- 2) DB 진단 ----
            with st.expander("2️⃣ MSSQL / DB 진단", expanded=True):
                db_diag = _diagnose_db()
                status = db_diag.get("status")
                reason = db_diag.get("reason") or "원인 미상"

                if db_diag.get("ok"):
                    st.success(f"DB: ✅ OK — {reason}")
                elif status == "WARN":
                    st.warning(f"DB: ⚠ 설정 필요 — {reason}")
                else:
                    st.error(f"DB: ❌ FAIL — {reason}")

                st.json(
                    {
                        "status": status,
                        "reason": reason,
                        "info": db_diag.get("info"),
                        "detail": db_diag.get("detail"),
                    },
                    expanded=False,
                )

            # ---- 3) 경로 / 설정 값 정리 ----
            with st.expander("3️⃣ 경로 / 설정 값", expanded=False):
                # 상단에서 정의된 상수들을 안전하게 읽어옴
                def _safe(name, default=None):
                    try:
                        return globals().get(name, default)
                    except Exception:
                        return default

                app_env = os.getenv("APP_ENV") or "development"
                log_level = os.getenv("LOG_LEVEL") or "DEBUG"

                CHAT_FILE_val = _safe("CHAT_FILE")
                UPLOAD_DIR_val = _safe("UPLOAD_DIR")
                MAX_FILE_SIZE_MB_val = _safe("MAX_FILE_SIZE_MB")
                MAX_PREVIEW_CHARS_val = _safe("MAX_PREVIEW_CHARS")

                base_dir = Path(__file__).resolve().parent.parent  # app/ 상위 디렉토리 기준

                st.json(
                    {
                        "env": {
                            "APP_ENV": app_env,
                            "LOG_LEVEL": log_level,
                            "TZ": os.getenv("TZ") or "Asia/Seoul",
                        },
                        "paths": {
                            "BASE_DIR": str(base_dir),
                            "CHAT_FILE": str(CHAT_FILE_val) if CHAT_FILE_val else None,
                            "UPLOAD_DIR": str(UPLOAD_DIR_val) if UPLOAD_DIR_val else None,
                        },
                        "limits": {
                            "MAX_FILE_SIZE_MB": MAX_FILE_SIZE_MB_val,
                            "MAX_PREVIEW_CHARS": MAX_PREVIEW_CHARS_val,
                            "SIMS_MAX_ROWS_VIEW": os.getenv("SIMS_MAX_ROWS_VIEW"),
                            "SIMS_MAX_ROWS_CHAT": os.getenv("SIMS_MAX_ROWS_CHAT"),
                        },
                    },
                    expanded=False,
                )

            # ---- 4) 최근 SIMS 컨텍스트 요약 ----
            with st.expander("4️⃣ 최근 SIMS 컨텍스트 요약", expanded=False):
                ctx_info = _extract_recent_sims_context()
                if not ctx_info:
                    st.info("SIMS 컨텍스트 관련 session_state 키를 찾지 못했습니다.")
                else:
                    key = ctx_info["key"]
                    value = ctx_info["value"]
                    st.write(f"🔑 사용 중인 컨텍스트 키: `{key}`")

                    # 1단계: 메타 구조 JSON 으로 요약
                    summary = {
                        "type": None,
                        "title": None,
                        "action": None,
                        "columns": None,
                        "rows": None,
                    }

                    if isinstance(value, dict):
                        summary["type"] = value.get("type")
                        summary["title"] = value.get("title")
                        summary["action"] = value.get("action")
                        cols = None
                        rows = None

                        # SIMS 결과 형식을 table 중심으로 가정
                        data = value.get("data") or value.get("rows") or value.get("df")
                        if hasattr(data, "to_dict"):  # DataFrame 등
                            try:
                                df = data if isinstance(data, pd.DataFrame) else pd.DataFrame(data)
                                cols = list(df.columns)
                                rows = len(df)
                                df_preview = df.head(3)
                            except Exception:
                                df_preview = None
                        else:
                            # list[dict] 형태일 수 있음
                            try:
                                df = pd.DataFrame(data) if data is not None else None
                                if df is not None:
                                    cols = list(df.columns)
                                    rows = len(df)
                                    df_preview = df.head(3)
                                else:
                                    df_preview = None
                            except Exception:
                                df_preview = None

                        summary["columns"] = cols
                        summary["rows"] = rows

                        st.json(summary, expanded=False)
                        if "df_preview" in locals() and df_preview is not None:
                            st.write("미리보기 (상위 3행):")
                            st.dataframe(df_preview, use_container_width=True)
                        else:
                            st.caption("표 형태로 변환할 수 없어서 미리보기를 생략합니다.")
                    else:
                        # dict 이 아니면 그냥 repr만 짧게
                        st.write("컨텍스트 값 타입:", type(value).__name__)
                        preview = textwrap.shorten(repr(value), width=300, placeholder=" ...")
                        st.code(preview)
# @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@

def _display_room_name(room: dict[str, Any] | None) -> str:
    if _is_empty_auto_room(room):
        return "➕ 새 대화 입력 대기 (저장 전)"

    if isinstance(room, dict):
        return str(room.get("name") or "업무 대화")

    return "➕ 새 대화 입력 대기"

# =========================
# 메인 영역
# =========================
st.markdown(
    f"""
    <h1 style="margin-bottom:4px;">🤖 SSAI LM Studio Chatbot</h1>
    <p style="color:#6b7280; margin-top:0;">현재 대화방: <strong>{_display_room_name(current_room)}</strong></p>
    <hr style="margin-top:10px;">
    """,
    unsafe_allow_html=True,
)

# 디버그(원하면 주석처리)
#st.caption(
#    "DEBUG "
#    f"search_open={st.session_state.get('__search_open')}, "
#    f"q={st.session_state.get('__search_q')!r}, "
#    f"roles={st.session_state.get('__search_roles')}, "
#    f"regex={st.session_state.get('__search_regex')}, "
#    f"case={st.session_state.get('__search_case')}, "
#    f"limit={st.session_state.get('__search_limit')}"
#)

# ✅ rooms(list) 모델을 쓰는 동안에는 single(messages) 모델 무시/정리 (중복 방지)
if isinstance(st.session_state.get("chat_rooms"), list):
    st.session_state.pop("messages", None)

# =========================
# 1) 입력 처리
# =========================
# 메시지 입력 UI는 채팅 결과 바로 아래, 파일 첨부 바로 위에 별도 inline form으로 렌더한다.
# 여기서는 이전 run에서 inline form이 넘긴 값을 먼저 처리한다.
typed_user_input = None
auto_user_input = st.session_state.pop("__sims_auto_user_input", None)

user_input = auto_user_input or typed_user_input

if user_input and user_input.strip():
    user_input = user_input.strip()

    st.session_state.pop("__sims_flash", None)
    # ✅ 키보드 보정(입력 버블에도 반영)
    try:
        from app.sims.nlq.nlq_router import keyboard_fix
        fixed = keyboard_fix(user_input)
        if fixed and fixed != user_input:
            log.debug("[chat] keyboard-fix(user_input): %r -> %r", user_input[:60], fixed[:60])
            user_input = fixed
    except Exception:
        pass

    st.session_state["__did_user_input"] = True

    is_pending_product_pick = _is_pending_product_pick_text(user_input)
    is_sims_result_followup = is_sims_result_followup_question(user_input)
    is_sims_input = is_sims_related_question(user_input) or is_sims_result_followup or is_pending_product_pick

    # ------------------------------------------------------------
    # 패널 B단계 성능 패치 신호
    # ------------------------------------------------------------
    # 일반 채팅/일반 SIMS 입력 rerun에서는 기존 view 재호출을 줄인다.
    # 단, "현재표 ..." 후속질문은 조회조건 화면이 사라져 보이면 안 되므로
    # skip_view를 걸지 않는다.
    try:
        if (
            st.session_state.get("__sims_panel_active")
            and st.session_state.get("__sims_selected")
            and not is_sims_result_followup
        ):
            st.session_state["__sims_panel_skip_view_once"] = True
            st.session_state["__sims_panel_skip_view_reason"] = "chat_input"
        else:
            st.session_state.pop("__sims_panel_skip_view_once", None)
            st.session_state.pop("__sims_panel_skip_view_reason", None)
    except Exception:
        pass


    # ✅ 1-현재 방에도 저장(원본 그대로, 히스토리 분리 전)
    current_room.setdefault("messages", []).append({
        "id": str(uuid.uuid4()),
        "role": "user",
        "content": user_input,
        "time": make_ts(),
        "seq": _next_seq(),
        **_message_meta("chat"),
    })

    _sync_room_meta(current_room, materialize=True)

    log.info(
        "[chat.input] %s input_len=%s auto_input=%s is_sims_input=%s is_followup=%s",
        _chat_log_kv(current_room),
        len(user_input),
        bool(auto_user_input),
        bool(is_sims_input),
        bool(is_sims_result_followup),
    )

    # ✅ 2-채널 히스토리에도 저장(LLM 히스토리 분리용)
    # SIMS 패널 ON + SIMS 관련 질문이면 sims_messages, 아니면 gen_messages
    ch = "sims_messages" if (st.session_state.get("__sims_open", False) and is_sims_input) else "gen_messages"

    current_room.setdefault(ch, []).append({
        "id": str(uuid.uuid4()),
        "role": "user",
        "content": user_input,
        "time": make_ts(),
        "seq": _next_seq(),
        **_message_meta("chat"),
    })

#   @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@
    # ✅ (D) 컨텍스트 메타 질문은 NLQ/LLM 전에 즉답 처리 (삽입 위치 고정)
#    if _try_answer_ctx_meta_question(
#        user_input,
#        room=current_room,
#        make_ts=make_ts,
#        next_seq=_next_seq,
#        logger=log,
#        max_age_sec=900,
#    ):
#        save_chat_rooms()
#        st.session_state["__queue_ai"] = False
#        st.rerun()
#   @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@

    # ✅ 제품 후보표가 열려 있을 때 '취소'도 SIMS/NLQ 입력으로 인정
    pending_pick_for_cancel = st.session_state.get("__io_pending_product_pick")
    has_pending_product_pick_for_cancel = (
        isinstance(pending_pick_for_cancel, dict)
        and isinstance(pending_pick_for_cancel.get("candidates"), list)
        and bool(pending_pick_for_cancel.get("candidates"))
    )

    cancel_norm = str(user_input or "").strip().replace(" ", "")
    is_pending_product_cancel = cancel_norm in {
        "취소",
        "후보취소",
        "선택취소",
        "제품선택취소",
        "제품후보취소",
        "그만",
    }

    # ✅ NLQ 자동 처리(가능하면 DB조회 표 반환 후 LLM 스킵)
    st.session_state.pop("__sims_last_push_sig", None)

    handled = False

    deferred_current_followup = False    

    compact_current = re.sub(r"\s+", "", str(user_input or ""))

    # ✅ 분석/KPI 명시 조회는 최근 SIMS 결과가 있어도 LLM 후속질문으로 보내지 않는다.
    # 예:
    # - 품목별추세요약표 조회
    # - 품목별 매출추세요약표 조회
    # - 품목별 매출 추세 요약표 조회
    # - 품목별 매출 추세 분석 조회
    # - 품목별 매출 예상 조회
    # - 품목별 재고부족현황 조회
    compact_analytics = compact_current

    is_explicit_analytics_nlq = (
        "조회" in compact_analytics
        and any(
            key in compact_analytics
            for key in (
                # 매출추세 요약표 / 요약
                "품목별추세요약표",
                "품목별추세요약",
                "품목별매출추세요약표",
                "품목별매출추세요약",
                "매출추세요약표",
                "매출추세요약",
                "추세요약표",
                "추세요약",

                # 매출추세 분석
                "품목별매출추세분석",
                "매출추세분석",

                # 매출예상
                "품목별매출예상",
                "매출예상",
                "예상매출",

                # 재고부족
                "품목별재고부족현황",
                "재고부족현황",
            )
        )
    )

    is_current_table_forced_followup = (
        "현재표" in compact_current
        or "현재조회결과" in compact_current
        or "현재조회자료" in compact_current
    )

    is_implicit_analytics_current_followup = (
        not is_explicit_analytics_nlq
        and _looks_like_implicit_analytics_current_followup(user_input)
    )

    if is_implicit_analytics_current_followup:
        log.info(
            "[chat.followup_table] implicit analytics current-table candidate: %r",
            str(user_input or "")[:80],
        )

    if (
        is_sims_input
        or is_pending_product_pick
        or is_pending_product_cancel
        or is_current_table_forced_followup
        or is_implicit_analytics_current_followup
    ):
        # 1) 현재표 후속 "표 생성" 요청은 LLM으로 보내지 말고 실제 pandas 표를 만든다.
        # 예:
        # - 현재표 거래처명 대학약국 상세표 만들어줘
        # - 현재표 제품별 매출 TOP 20 표로 만들어줘
        # - 현재표 거래처별 매출 TOP 20 표로 만들어줘
        try:
            handled = _try_handle_current_table_dataframe_followup(
                user_input,
                room=current_room,
                make_ts=make_ts,
                next_seq=_next_seq,
            )
        except Exception:
            log.exception("[chat.followup_table] handler failed")
            handled = False

        if handled:
            log.debug("[chat.followup_table] handled → skip LLM/NLQ, content=%r", user_input[:80])

        elif is_current_table_forced_followup:
            # 조회 화면에서는 현재표 stash가 panel render 이후에 완료될 수 있다.
            # 따라서 여기서 바로 LLM fallback으로 보내지 말고,
            # panel render 이후 한 번 더 현재표 후속분석을 재시도한다.
            st.session_state["__deferred_current_table_followup"] = {
                "user_input": user_input,
                "ts": time.time(),
                "retry": 0,
            }
            deferred_current_followup = True
            log.debug(
                "[chat.followup_table] defer current-table followup until after panel render: %r",
                user_input[:80],
            )            

        elif (
            is_implicit_analytics_current_followup
            and not is_pending_product_pick
            and not is_pending_product_cancel
        ):
            current_followup_query = _normalize_implicit_analytics_current_followup(user_input)

            log.info(
                "[chat.followup_table] implicit analytics current-table followup: %r -> %r",
                user_input[:80],
                current_followup_query[:80],
            )

            try:
                handled = _try_handle_current_table_dataframe_followup(
                    current_followup_query,
                    room=current_room,
                    make_ts=make_ts,
                    next_seq=_next_seq,
                )
            except Exception:
                log.exception("[chat.followup_table] implicit analytics followup handler failed")
                handled = False

            if not handled:
                st.session_state["__deferred_current_table_followup"] = {
                    "user_input": current_followup_query,
                    "ts": time.time(),
                    "retry": 0,
                }
                deferred_current_followup = True
                log.debug(
                    "[chat.followup_table] implicit analytics followup deferred until after panel render: %r",
                    current_followup_query[:80],
                )

        elif (
            is_sims_result_followup
            and not is_explicit_analytics_nlq
            and not is_pending_product_pick
            and not is_pending_product_cancel
        ):
            # 최신 SIMS 결과에 대한 분석/요약/설명 요청은
            # 새 DB 조회가 아니므로 NLQ router로 보내지 않는다.
            # 단, "품목별추세요약표 조회"처럼 분석/KPI 조회가 명확한 문장은
            # NLQ router로 보내야 한다.
            log.debug("[chat] skip NLQ (SIMS result follow-up → LLM): %r", user_input[:80])
            handled = False
            
        else:
            if is_pending_product_pick:
                log.debug("[chat] pending product pick accepted as sims-related: %r", user_input[:80])

            if is_pending_product_cancel:
                log.debug("[chat] pending product cancel accepted as sims-related: %r", user_input[:80])

            nlq_input = _normalize_doc_common_nlq_text(user_input)
            nlq_input = _normalize_doc_vendor_nlq_text(nlq_input)

            if nlq_input != user_input:
                log.info("[chat] normalize doc nlq: %r -> %r", user_input, nlq_input)

            handled = try_handle_nlq(
                nlq_input,
                room=current_room,
                session_state=st.session_state,
                make_ts=make_ts,
                next_seq=_next_seq,
                logger=log,
            )

    else:
        log.debug("[chat] skip NLQ (not sims-related): %r", user_input[:80])

    save_chat_rooms()

    if handled:
        st.session_state["__queue_ai"] = False

        # 현재표 후속표/채팅 NLQ 처리 후에도 SIMS 패널 토글은 열린 상태로 유지한다.
        # __sims_open은 toggle 위젯 key이므로 여기서 직접 True 대입하지 않고,
        # 사이드바 fragment가 다음 rerun에서 toggle 생성 전에 보정하도록 keep flag만 세운다.
        try:
            if st.session_state.get("__sims_open") or st.session_state.get("__sims_panel_active"):
                st.session_state["__sims_panel_active"] = True
                st.session_state["__sims_force_open"] = False
                st.session_state["__sims_keep_open_after_push"] = True
                st.session_state.pop("__sims_close_after_push", None)
        except Exception:
            pass

        log.debug("[chat] nlq handled → skip LLM, content=%r", user_input[:80])
        st.rerun()


    elif deferred_current_followup:
        # 현재표 질문은 panel render 이후 재시도한다.
        # 여기서 LLM을 켜면, 현재표가 준비되기 전에 일반 답변으로 빠진다.
        st.session_state["__queue_ai"] = False
        st.session_state.pop("__sims_last_push_sig", None)
        log.debug(
            "[chat.followup_table] deferred current-table followup queued, content=%r",
            user_input[:80],
        )

        # 첫 메시지로 방 제목이 바뀐 경우, 이미 그려진 상단 제목/사이드바 입력칸을 갱신한다.
        changed_room_id = str(st.session_state.pop("__chat_room_title_changed_room_id", "") or "").strip()
        if changed_room_id and changed_room_id == str(current_room.get("id") or ""):
            st.rerun()

    else:
        st.session_state["__queue_ai"] = True
        st.session_state.pop("__sims_last_push_sig", None)
        log.debug("[chat] user_input accepted, __queue_ai=True, content=%r", user_input[:80])

        # 첫 메시지로 채팅방 제목이 자동 변경된 경우,
        # 상단 현재 대화방/좌측 이름변경 입력칸은 이미 렌더링된 뒤라
        # 한 번 rerun해서 새 제목을 화면에 반영한다.
        changed_room_id = str(st.session_state.pop("__chat_room_title_changed_room_id", "") or "").strip()
        if changed_room_id and changed_room_id == str(current_room.get("id") or ""):
            st.rerun()


# =========================
# 0) 전역 기본값/상수
# =========================
DEFAULT_MODE = "Panel (A)"

# =========================
# 2) 채팅 렌더 + 파일/SIMS
# =========================
with st.container():


    # ===== SIMS 버튼 프리패스 (결과 렌더보다 위에서 상태 선반영) =====
    # ✅ 토글 전용 키 — setdefault만. 절대 대입(=True/False)하지 말 것.
    st.session_state.setdefault("__sims_open", False)
    st.session_state.setdefault("__sims_panel_active", False)
    st.session_state.setdefault("__sims_rendered", False)
    st.session_state.setdefault("__sims_was_final", False)
    st.session_state.setdefault("__sims_selected_snapshot", {})
    st.session_state.setdefault("__sims_run_flag", False)

    # 🔁 옵션 초기화 버튼이 눌렸을 때 바로 리셋 (단일 경로)
    if st.session_state.pop("__sims_reset_requested", False):
        _ns = (
            st.session_state.get("__sims_widget_ns")
            or st.session_state.get("__sims_form_id")
            or st.session_state.get("__sims_run_seq")
        )
        _ns = str(_ns) if _ns is not None else None
        log.info("[sims.reset] apply start ns=%s", _ns)

        # 1) SIMS 핵심 상태키 정리
        for k in [
            # 입력/선택
            "__sims_q", "__sims_cat", "__sims_action", "__sims_top",
            "__sims_selected", "__sims_selected_snapshot", "__sims_selected_action",
            "__sims_hub_action", "__sims_last_action",
            # 실행/렌더 플래그
            "__sims_panel_active", "__sims_run_flag", "__sims_run_seq", "__sims_form_id",
            "__sims_rendered", "__sims_was_final", "__sims_force_open", "__sims_widget_ns",
            # 컨텍스트/결과/부가 상태
            "__sims_context", "__sims_context_note", "__sims_result", "__sims_last_push_sig",
            "__sims_flash", "__sims_flash_close", "__sims_flash_csv", "__sims_flash_xlsx",
        ]:
            st.session_state.pop(k, None)

        # 2) ✅ 뷰별 옵션 위젯키 정리(현재 ns만)
        prefixes = ("__vendors_", "__users_", "__codes_", "__rddbc")
        for kk in list(st.session_state.keys()):
            if kk.startswith(prefixes):
                if (_ns is None) or kk.endswith(f"__{_ns}"):
                    st.session_state.pop(kk, None)

        # 3) 기본값 복구
        st.session_state["__sims_open"] = False
        st.session_state["__sims_open_ui"] = False
        st.session_state["__sims_force_open"] = False
        st.session_state["__sims_rendered"] = False
        st.session_state["__sims_was_final"] = False
        st.session_state["__sims_panel_active"] = False
        st.session_state["__sims_run_flag"] = False
        st.session_state["__sims_inner_submit"] = False
        st.session_state.setdefault("__sims_selected_snapshot", {})

        log.info("[sims.reset] apply done")
        st.rerun()

    # (B0) SIMS→채팅 동기화(인박스 드레인 + 컨텍스트 주입)
    wire_chat_context()

    # === SIMS 실행 프리패스: 클릭 신호를 메인 컨테이너가 즉시 감지하도록 보강 ===
    if st.session_state.get("__sims_run_flag") or st.session_state.get("__sims_inner_submit"):
        st.session_state["__sims_force_open"] = True
        st.session_state["__sims_panel_active"] = True
        st.session_state["__sims_rendered"] = False
        st.session_state["__sims_was_final"] = False

        log.info(
            "[prepass] submit detected → run_flag=%s, inner_submit=%s, open=%s, panel_active=%s",
            st.session_state.get("__sims_run_flag"),
            st.session_state.get("__sims_inner_submit"),
            (st.session_state.get("__sims_open") or st.session_state.get("__sims_force_open")),
            st.session_state["__sims_panel_active"],
        )

    # ✅ SIMS 세션키 선초기화 (어디서든 안전하게 사용)
    # 다시 한 번 안전 장치(다중 경로에서 호출돼도 일관 유지)
    st.session_state.setdefault("__sims_open", False)        # 토글 위젯 전용 (건드리지 말 것)
    st.session_state.setdefault("__sims_open_ui", False)     # (선택) UI 토글 별도키 쓰고싶다면
    st.session_state.setdefault("__sims_force_open", False)  # 프로그램적으로 "열림 간주" 플래그
    st.session_state.setdefault("__sims_q", "")
    st.session_state.setdefault("__sims_panel_active", False)
    st.session_state.setdefault("__sims_run_flag", False)
    st.session_state.setdefault("__sims_form_id", 0)
    st.session_state.setdefault("__sims_run_seq", 0)
    st.session_state.setdefault("__sims_last_action", "")
    st.session_state.setdefault("__sims_rendered", False)

    # 메시지 병합(rooms 전용)
    merged_msgs = list(current_room.get("messages", [])) + list(current_room.get("history", []))

    # id 중복 제거
    _seen = set(); _dedup = []
    for _m in merged_msgs:
        _id = _m.get("id") or f"{_m.get('role')}|{_m.get('time')}|{hash((_m.get('content') or ''))}"
        if _id in _seen: 
            continue
        _seen.add(_id)
        _dedup.append(_m)
    merged_msgs = _dedup

    # 시간/순서 정렬
    try:
        # seq 없는(표/시스템 payload) 항목이 0으로 처리되어 맨 위로 가는 문제 방지
        for i, _m in enumerate(merged_msgs):
            _m.setdefault("__merge_idx", i)
        def _sort_key(m):
            s = m.get("seq")
            if isinstance(s, (int, float)):
                return (0, s, m.get("__merge_idx", 0))
            return (1, 0, m.get("__merge_idx", 0))  # seq 없음 → 뒤로
        merged_msgs.sort(key=_sort_key)
    except Exception:
        pass

    # SIMS 표 버블 전용 렌더
    def _render_message(m: dict) -> bool:
        meta = (m.get("meta") or {})

        if meta.get("kind") == "table" and meta.get("table_key"):
            table_key = meta.get("table_key")
            df = (st.session_state.get("sims_tables") or {}).get(table_key)

            title_text = (
                m.get("title")
                or m.get("content")
                or meta.get("action")
                or "SIMS 결과"
            )

            if df is not None:
                # 저장된 SIMS 표도 chat_middleware의 기존 렌더러로 다시 그린다.
                # 그래야 조회조건/헤더/summary_md/음수 빨간색/부족등급 굵게/다운로드/LLM 버튼이 유지된다.
                table_item = dict(m)
                table_item["type"] = "table"
                table_item["role"] = "assistant"
                table_item["data"] = df
                table_item["meta"] = meta
                table_item["title"] = title_text
                table_item["action"] = (
                    m.get("action")
                    or meta.get("action")
                    or table_item["title"]
                )
                table_item["params"] = (
                    m.get("params")
                    or meta.get("params")
                    or {}
                )

                render_sims_chat_item(table_item)

            else:
                with st.chat_message("assistant"):
                    st.markdown(m.get("content") or "📊 SIMS 결과")
                    st.info("표 데이터가 세션에서 만료되었습니다. 같은 조회를 다시 실행해 주세요.")

            return True

        # SIMS text 안내 payload 렌더
        # 예: 제품 후보 선택 취소, 0건 조회 메시지
        if (
            m.get("type") in {"text", "object"}
            and (
                m.get("message")
                or m.get("data")
                or m.get("content")
                or meta.get("summary_md")
            )
            and (
                meta.get("nlq")
                or m.get("action")
                or m.get("title")
            )
        ):
            text_item = dict(m)
            text_item["role"] = "assistant"
            text_item.setdefault("title", m.get("title") or m.get("action") or "SIMS 안내")
            text_item.setdefault("action", m.get("action") or text_item["title"])
            text_item.setdefault(
                "message",
                m.get("message")
                or m.get("content")
                or m.get("data")
                or meta.get("summary_md")
                or ""
            )
            text_item.setdefault("data", text_item.get("message") or "")
            text_item["meta"] = meta

            render_sims_chat_item(text_item)
            return True


    # (A) 이번 rerun에서 점프할 앵커(검색 패널에서 set) 한 번만 소비
    _jump_to = st.session_state.pop("__scroll_to_msg", None)

    # 이번 rerun에서 history 영역에 이미 렌더되는 메시지 ID.
    # render_pending_chat_items()에서 같은 SIMS 표를 한 번 더 그리지 않도록 사용한다.
    try:
        st.session_state["__chat_rendered_ids_this_run"] = [
            m.get("id")
            for m in merged_msgs
            if isinstance(m, dict) and m.get("id")
        ]
    except Exception:
        st.session_state["__chat_rendered_ids_this_run"] = []

    # ------------------------------------------------------------
    # SIMS LLM 분석 fragment runner
    # ------------------------------------------------------------
    # chat_middleware.py의 LLM 분석 버튼은 st.fragment 안에서 실행된다.
    # 따라서 전체 앱 rerun 없이, 현재 화면은 그대로 둔 채 LLM 분석만 수행한다.
    def _run_sims_llm_analysis_from_fragment(prompt: str) -> None:
        try:
            prompt = (prompt or "").strip()
            if not prompt:
                prompt = "현재 조회 결과를 핵심 요약, 주요 수치, 주의할 점, 다음 조회 제안 순서로 분석해줘"

            msgs = build_messages_with_system([], user_text=prompt)

            try:
                has_sims = any(
                    (m.get("role") == "system") and ("SIMS_JSON" in str(m.get("content", "")))
                    for m in msgs
                )
                log.debug(
                    "[chat.fragment] build_messages_with_system OK → msgs=%d, has_sims=%s",
                    len(msgs),
                    has_sims,
                )
            except Exception:
                pass

            stream_and_append_assistant(
                messages_for_ai=msgs,
                room=current_room,
                temperature=0.2,
                history_channel="sims_messages",
            )

            log.debug("[chat.fragment] stream_and_append_assistant finished")

        except Exception:
            log.exception("[chat.fragment] LLM analysis runner failed")
            st.error("LLM 분석 실행 중 오류가 발생했습니다.")

    st.session_state["__sims_llm_analysis_runner"] = _run_sims_llm_analysis_from_fragment

    # (B) 채팅 렌더 — 각 메시지 위에 앵커 삽입
    for idx, m in enumerate(merged_msgs):

        anchor_core = m.get("id") or f"seq-{m.get('seq', idx)}"
        # 검색 결과에서 저장하는 키와 1:1로 맞춘다 → 'jump-' 접두어 포함
        st.markdown(f"<div id='jump-{anchor_core}'></div>", unsafe_allow_html=True)

        if _render_message(m):
            continue

        role = m.get("role") if m.get("role") in ("user","assistant") else "assistant"
        content = (m.get("content") or "")
        ts = normalize_ts(m.get("time",""))

        with st.chat_message(role):
            st.markdown(content)
            if ts:
                st.caption(ts)

    # (C) 채팅 바로 아래에 '이번 턴 답변' 표시 영역 예약 (✅ 딱 1곳에서만 생성/지정)
    pending_area = st.container()

    # ✅ (1) NLQ/백엔드 푸시 테이블 렌더 타겟 지정 (반드시 wire_chat_context()보다 먼저)
    try:
        from app.ui.chat_middleware import set_chat_render_target, render_pending_chat_items,render_sims_chat_item
        set_chat_render_target(pending_area)
        render_pending_chat_items(pending_area, room=current_room)
    except Exception:
        log.exception("[ui] render_pending_chat_items failed")

    # ✅ (2) 기존 앵커도 동일 위치로 (있다면 같이 유지)
    set_chat_render_anchor(pending_area)

    # ✅ 다음 메시지 입력창
    # - 위치: 채팅 결과 바로 아래 / 파일 첨부 바로 위
    # - Enter 전송
    # - 전송 버튼 없음
    # - 전송 후 입력창 자동 비움
    # - 기존 text_area/form 방식 잔여 상태 제거
    st.session_state.pop("__chat_inline_input_text", None)

    def _submit_inline_chat_input() -> None:
        text = str(st.session_state.get("__chat_inline_text") or "").strip()
        if not text:
            return

        st.session_state["__sims_auto_user_input"] = text
        st.session_state["__chat_inline_text"] = ""

    st.text_input(
        "다음 메시지",
        key="__chat_inline_text",
        placeholder="메시지를 입력하세요... Enter 전송",
        label_visibility="collapsed",
        disabled=st.session_state.get("__an_busy", False),
        on_change=_submit_inline_chat_input,
    )

    # ✅ 스크롤 타겟(바닥)
    st.markdown("<div id='__chat_bottom'></div>", unsafe_allow_html=True)
 
    # ✅ 검색 점프 / 입력 후 바닥 유지 (Streamlit rerun 스크롤 튐 보정)
    try:
        if _jump_to:
            stc.html(
                f"""
<script>
(function(){{
  const el = window.parent.document.getElementById({json.dumps(_jump_to)});
  if (el) el.scrollIntoView({{ behavior: "smooth", block: "start" }});
}})();
</script>
""",
                height=0,
            )
        elif st.session_state.pop("__did_user_input", False):
            stc.html(
                """
<script>
(function(){
  const el = window.parent.document.getElementById('__chat_bottom');
  if (el) el.scrollIntoView({ behavior: "smooth", block: "end" });
})();
</script>
""",
                height=0,
            )
    except Exception:
        log.exception("[ui] scrollIntoView failed")

    # (D) 파일 첨부 + SIMS 섹션
    with st.container(): 
        st.divider()
        # SIMS 토글/쿼리 기본값 방어 (첫 렌더 대비)
        st.session_state.setdefault("__sims_open", False)
        st.session_state.setdefault("__sims_q", "")


        can_upload_file = require_permission("UPLOAD_FILE", show_error=False)

        if not can_upload_file:
            st.info(_upload_unavailable_message())

        uploaded_files = st.file_uploader(
            "📂 파일 첨부 (PDF, Excel, CSV, TXT, DOCX, 이미지 등)",
            type=["pdf", "csv", "xlsx", "xls", "txt", "docx", "png", "jpg", "jpeg"],
            accept_multiple_files=True,
            key="file_upload_below_input",
            disabled=(not can_upload_file) or st.session_state.get("__an_busy", False),
            help=_upload_unavailable_help() if not can_upload_file else None,
        )

        # 2) SIMS 결과(모드별 실행 방식)
        # ✅ 렌더 조건은 로컬 변수로만 판정한다.
        # __sims_open  : 사용자 토글
        # __sims_open_ui: (선택) UI 별도 토글
        # __sims_force_open: 실행 직후 결과만 강제 노출하고 싶을 때 사용(상태키에 대입하지 않음)
        # 사용자가 사이드바 토글을 끄면 SIMS 패널은 반드시 닫힌다.
        # __sims_force_open / __sims_open_ui 는 실행 직후 보조 플래그일 뿐,
        # 사용자 닫기 상태를 덮어쓰지 않는다.
        _user_sims_open = bool(st.session_state.get("__sims_open"))
        if not _user_sims_open:
            st.session_state["__sims_force_open"] = False
            st.session_state["__sims_open_ui"] = False
        _open = _user_sims_open
        if _open:
            st.markdown('<a id="__sims_result"></a>', unsafe_allow_html=True)
            st.markdown("## 🧩 SIMS 결과")

#           st.caption("SIMS 관련 질문에 답변이 도착하면 이 영역에 결과가 표시됩니다.")
 
            # 중복 렌더 가드(프레임 단위)
            st.session_state.setdefault("__sims_rendered", False)
            st.session_state.setdefault("__sims_was_final", False)
            st.session_state.setdefault("__sims_panel_active", False)
            st.session_state.setdefault("__sims_last_render_run_seq", -1)

            st.session_state.setdefault("__sims_mode", DEFAULT_MODE)
            mode = st.session_state.get("__sims_mode") or DEFAULT_MODE
            run_flag = bool(st.session_state.get("__sims_run_flag"))
            inner_submit = bool(st.session_state.get("__sims_inner_submit"))

            selected_now = st.session_state.get("__sims_selected") or {}
            has_selection = bool(selected_now.get("category") and selected_now.get("action"))

            run_seq = int(st.session_state.get("__sims_run_seq") or 0)
            already_rendered_this_run = (
                run_seq > 0
                and int(st.session_state.get("__sims_last_render_run_seq", -1)) == run_seq
            )

            log.info(
                "[prepass.dbg] run_flag=%s inner_submit=%s panel_active=%s run_seq=%s selected=%r",
                st.session_state.get("__sims_run_flag"),
                st.session_state.get("__sims_inner_submit"),
                st.session_state.get("__sims_panel_active"),
                st.session_state.get("__sims_run_seq"),
                st.session_state.get("__sims_selected"),
            ) 

            if run_flag or inner_submit:
                st.session_state["__sims_panel_active"] = True
                st.session_state["__sims_rendered"] = False
                st.session_state["__sims_was_final"] = False

            if str(mode).startswith("Panel"):
                panel_active = bool(st.session_state.get("__sims_panel_active"))

                # 채팅창 SSOT 구조:
                # - 패널은 조건 입력/조회 실행 전용으로 계속 렌더한다.
                # - DB 조회 재실행 방지는 "채팅 입력/후속분석 rerun"처럼 명확한 경우에만
                #   user_input 처리부에서 __sims_panel_skip_view_once를 세워서 처리한다.
                # - 여기서 idle 상태마다 skip_view를 자동 설정하면 조건 입력 화면까지 사라진다.

                # 폼 기반 화면은 결과가 나온 뒤에도 같은 액션에서 연속 조회가 가능해야 함
                should_render = has_selection and (
                    panel_active
                    or ((run_flag or inner_submit) and not already_rendered_this_run)
                )

                if should_render:
                    st.session_state["__sims_last_render_run_seq"] = run_seq

                    log.info(
                        "[app.main] SIMS about to render: mode=%r, run_seq=%s, form_id=%s, selected=%r",
                        st.session_state.get("__sims_mode"),
                        st.session_state.get("__sims_run_seq"),
                        st.session_state.get("__sims_form_id"),
                        st.session_state.get("__sims_selected"),
                    )

                    if st.session_state.get("__sims_selected_snapshot"):
                        st.session_state["__sims_selected"] = dict(
                            st.session_state["__sims_selected_snapshot"]
                        )

                    selected_for_render = st.session_state.get("__sims_selected") or {}

                    if not _check_sims_action_permission(selected_for_render):
                        st.session_state["__sims_panel_active"] = False
                        st.session_state["__sims_run_flag"] = False
                        st.session_state["__sims_inner_submit"] = False
                        st.stop()

                    render_sims_main(
                        selected=selected_for_render
                    )

                    # SIMS 패널 직접 조회 결과도 현재 채팅방 정책에 맞춰 저장한다.
                    # - 방 선택 없음: 새 대화 대기방을 정식 방으로 전환 후 저장
                    # - 기존 방 선택: 선택 방에 저장
                    if st.session_state.get("__sims_was_final"):
                        try:
                            pushed_panel_result = _push_panel_result_to_current_chat(
                                current_room,
                                selected_for_render=selected_for_render,
                                run_seq=run_seq,
                            )

                            if pushed_panel_result:
                                # 결과는 채팅창에 1회 저장한다.
                                # 채팅 영역은 이미 위에서 한 번 렌더되었으므로,
                                # app 전체 rerun을 걸지 않고 같은 pending_area에 즉시 1회 더 렌더한다.
                                # 이렇게 해야 조건 입력 화면이 사라지지 않고, DB 재조회도 발생하지 않는다.
                                st.session_state["__sims_run_flag"] = False
                                st.session_state["__sims_inner_submit"] = False
                                st.session_state["__sims_force_open"] = False
                                st.session_state["__sims_panel_active"] = True

                                # v2 자동 닫기/이전 실험 플래그 제거
                                st.session_state.pop("__sims_close_after_push", None)
                                st.session_state.pop("__sims_keep_open_after_push", None)
                                st.session_state.pop("__sims_panel_skip_view_once", None)
                                st.session_state.pop("__sims_panel_skip_view_reason", None)

                                try:
                                    from app.ui.chat_middleware import render_pending_chat_items as _render_pending_sims_items
                                    _render_pending_sims_items(pending_area, room=current_room)
                                    log.info(
                                        "[chat.panel.push] rendered immediately in chat pending area action=%s run_seq=%s",
                                        (selected_for_render or {}).get("action"),
                                        run_seq,
                                    )
                                except Exception:
                                    log.exception("[chat.panel.push] immediate pending render failed")


                        except Exception:
                            log.exception("[chat.panel.push] failed")

                    # ------------------------------------------------------------
                    # 조회 화면에서 "현재표 ..." 질문이 먼저 들어온 경우,                
                    # panel render 이후 현재표 stash가 완료되므로 여기서 재시도한다.
                    # ------------------------------------------------------------
                    deferred = st.session_state.get("__deferred_current_table_followup")
                    if isinstance(deferred, dict):
                        deferred_query = str(deferred.get("user_input") or "").strip()
                        retry = int(deferred.get("retry") or 0)

                        if deferred_query and retry <= 1:
                            try:
                                log.debug(
                                    "[chat.followup_table] retry deferred current-table followup after panel render: %r",
                                    deferred_query[:80],
                                )

                                handled_deferred = _try_handle_current_table_dataframe_followup(
                                    deferred_query,
                                    room=current_room,
                                    make_ts=make_ts,
                                    next_seq=_next_seq,
                                )

                                if handled_deferred:
                                    st.session_state.pop("__deferred_current_table_followup", None)
                                    st.session_state["__queue_ai"] = False

                                    # deferred 현재표 후속표도 결과 표시 후 패널 열린 상태를 유지한다.
                                    try:
                                        if st.session_state.get("__sims_open") or st.session_state.get("__sims_panel_active"):
                                            st.session_state["__sims_panel_active"] = True
                                            st.session_state["__sims_force_open"] = False
                                            st.session_state["__sims_keep_open_after_push"] = True
                                            st.session_state.pop("__sims_close_after_push", None)
                                    except Exception:
                                        pass

                                    save_chat_rooms()
                                    log.debug(
                                        "[chat.followup_table] deferred current-table followup handled after panel render → rerun"
                                    )
                                    st.rerun()


                                else:
                                    # 패널 렌더 후에도 전용 pandas 후속분석이 처리하지 못한 경우:
                                    # 1) 현재표 source가 없으면 절대 LLM으로 넘기지 않는다.
                                    #    이전 SIMS_ANALYSIS_CONTEXT가 남아 있으면 엉뚱한 이전표를 분석할 수 있다.
                                    # 2) TOP/목록/표/월별/조건 필터 같은 정형 표 요청도 LLM으로 넘기지 않는다.
                                    # 3) 그 외 서술형 분석 질문만 기존 LLM fallback을 허용한다.
                                    st.session_state.pop("__deferred_current_table_followup", None)

                                    df_now, table_key_now = _current_table_get_latest_df()
                                    no_current_source = not isinstance(df_now, pd.DataFrame) or df_now.empty
                                    hard_table_request = _current_table_should_block_llm_fallback(deferred_query)

                                    if no_current_source:
                                        _push_no_current_table_notice(deferred_query)
                                        st.session_state["__queue_ai"] = False
                                        save_chat_rooms()
                                        log.debug(
                                            "[chat.followup_table] deferred current-table followup has no source → notice only query=%r",
                                            deferred_query[:80],
                                        )
                                        st.rerun()

                                    elif hard_table_request:
                                        _current_table_push_notice(
                                            title="현재표 후속분석 미지원",
                                            action="현재표 후속분석 미지원",
                                            message=(
                                                "요청하신 현재표 후속표는 아직 지원하지 않는 형식입니다.\n\n"
                                                "TOP/목록/표/월별/일자별/조건 필터 요청은 LLM이 임의로 작성하지 않고, "
                                                "pandas handler에서 지원되는 경우에만 표로 생성합니다."
                                            ),
                                            query_summary="현재표 / 후속분석 미지원 / LLM fallback 차단",
                                            source_query=deferred_query,
                                        )
                                        st.session_state["__queue_ai"] = False
                                        save_chat_rooms()
                                        log.debug(
                                            "[chat.followup_table] deferred hard current-table request not handled → notice only table_key=%s query=%r",
                                            table_key_now,
                                            deferred_query[:80],
                                        )
                                        st.rerun()

                                    else:
                                        st.session_state["__queue_ai"] = True
                                        st.session_state.pop("__sims_last_push_sig", None)
                                        log.debug(
                                            "[chat.followup_table] deferred current-table analysis not handled → fallback to LLM query=%r",
                                            deferred_query[:80],
                                        )


                            except Exception:
                                st.session_state.pop("__deferred_current_table_followup", None)
                                log.exception("[chat.followup_table] deferred current-table followup retry failed")

                    st.session_state["__sims_run_flag"] = False
                    st.session_state["__sims_inner_submit"] = False

                    # 결과가 이미 채팅창으로 넘어간 뒤에도 SIMS 패널은 닫지 않는다.
                    # 단, run_flag/inner_submit은 이미 내려갔으므로 같은 조회가 재실행되지는 않는다.
                    st.session_state["__sims_rendered"] = False
                    if st.session_state.get("__sims_was_final"):
                        st.session_state["__sims_panel_active"] = True
                        st.session_state["__sims_force_open"] = False
                        st.session_state["__sims_keep_open_after_push"] = True
                    else:
                        st.session_state["__sims_panel_active"] = True
                else:
                    if not has_selection:
                        st.caption("좌측 ‘SIMS 옵션’에서 카테고리/작업을 선택한 뒤 ‘SIMS 작업 열기’를 눌러주세요.")

            else:
                # Hub(B): 메인에서 허브 UI를 항상 렌더
                if not st.session_state["__sims_rendered"]:
                    log.info("[app.main] call render_sims_main() Hub(B) mode=%r", mode)

                    selected_for_render = st.session_state.get("__sims_selected") or {}

                    if not _check_sims_action_permission(selected_for_render):
                        st.session_state["__sims_rendered"] = True
                        st.stop()

                    render_sims_main(
                        selected=selected_for_render
                    )                    
                    log.info("[app.main] after render_sims_main() Hub(B)")
                    if st.session_state.get("__sims_was_final"):
                        st.session_state["__sims_rendered"] = True

            # 컨텍스트 안내(있으면 노출)
            note = st.session_state.get("__sims_context_note")
            if note:
                st.info("🧠 " + note)

    # 등급에 따른 감추기 KWG
    if _can_show_admin_diagnostics_sidebar():
        # (선택) 디버그 보기
        with st.sidebar.expander("🔍 디버그 보기", expanded=False):
            show_console = st.checkbox("콘솔/로그 요약 보기", key="__show_dbg")
            if show_console:
                try:
                    import pathlib, itertools
                    log_path = pathlib.Path(os.getenv("LOG_FILE") or "logs/app.log")
                    if log_path.exists():
                        tail = "\n".join(log_path.read_text(encoding="utf-8", errors="ignore").splitlines()[-60:])
                        st.text_area("app.log (tail 60)", value=tail, height=220, label_visibility="collapsed")
                    else:
                        st.caption("app.log 파일이 아직 없습니다.")
                except Exception as e:
                    st.caption(f"디버그 표시 오류: {e}")        # ⚠️ SIMS 옵션/실행 UI는 오롯이 '사이드바'에서만 렌더합니다.

            # 🔹 결과 도착 시 간단 스크롤 유도(링크)
            if st.session_state.get("__sims_last_push_sig"):
                st.markdown("[결과 위치로 이동](#__sims_result)")

            # 🔹 SIMS 임시 플래시 (다음 액션 전까지 유지: 자동 접기 지원)
            st.session_state.setdefault("__sims_flash_autoclose", True)       # 자동 접기 on/off
            st.session_state.setdefault("__sims_flash_autoclose_ms", 8000)    # 접기까지 대기(ms)
            st.session_state.setdefault("__sims_flash_auto_remove", False)    # 접힌 뒤 제거할지

            _flash = st.session_state.get("__sims_flash")

            if isinstance(_flash, dict):
                # 최초 표출 시각 보강(이전 버전 호환)
                if "ts" not in _flash:
                    _flash["ts"] = time.time() * 1000.0

                # 자동 접기/제거 판정
                now_ms = time.time() * 1000.0
                auto    = bool(st.session_state["__sims_flash_autoclose"])
                ttl_ms  = int(st.session_state["__sims_flash_autoclose_ms"])
                expired = auto and ((now_ms - float(_flash["ts"])) >= ttl_ms)

                # 제거 모드면 expander 만들지 않고 바로 제거
                if expired and bool(st.session_state["__sims_flash_auto_remove"]):
                    st.session_state.pop("__sims_flash", None)
                else:
                    _df = (st.session_state.get("sims_tables") or {}).get(_flash.get("table_key"))
                    title = _flash.get("title", "SIMS 결과(임시)")

                    # ▶▶ 여기에서 파일명 준비 ( _flash 가 확실히 존재하는 시점 )
                    def _slug_name(s: str) -> str:
                        s = re.sub(r"[^0-9A-Za-z가-힣_\-]+", "_", str(s or "")).strip("_")
                        return s[:60] or "sims_result"
                    _action_guess = (
                        _flash.get("action")                       # 플래시에 액션명을 저장했다면 최우선
                        or st.session_state.get("__sims_selected_action")
                        or st.session_state.get("__sims_action")
                        or st.session_state.get("__sims_hub_action")
                        or title
                        or "SIMS_결과"
                    )
                    base  = _slug_name(_action_guess)
                    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    csv_name  = f"{base}_{stamp}.csv"
                    xlsx_name = f"{base}_{stamp}.xlsx"

                    # 만료 시 expanded=False 로 자동 접힘
                    expanded_now = not expired

                    if _df is not None:
                        with st.expander(title, expanded=expanded_now):
                            st.dataframe(_df, use_container_width=True)

                            # CSV / Excel 다운로드 권한 체크
                            can_export_excel = require_permission("EXPORT_EXCEL", show_error=False)

                            if not can_export_excel:
                                st.warning("다운로드 권한이 없습니다. 필요 권한: EXPORT_EXCEL (엑셀/CSV 다운로드)")
                            else:
                                # CSV
                                try:
                                    csv_bytes = _df.to_csv(index=False).encode("utf-8-sig")
                                    st.download_button(
                                        "CSV 다운로드",
                                        data=csv_bytes,
                                        file_name=csv_name,
                                        mime="text/csv",
                                        key="__sims_flash_csv",
                                    )
                                except Exception:
                                    pass

                                # Excel (xlsxwriter → openpyxl 폴백)
                                try:
                                    import io
                                    bio = io.BytesIO()

                                    try:
                                        with pd.ExcelWriter(bio, engine="xlsxwriter") as w:
                                            _df.to_excel(w, index=False, sheet_name="result")
                                    except Exception:
                                        bio = io.BytesIO()
                                        with pd.ExcelWriter(bio, engine="openpyxl") as w:
                                            _df.to_excel(w, index=False, sheet_name="result")

                                    bio.seek(0)
                                    st.download_button(
                                        "Excel 다운로드",
                                        data=bio.getvalue(),
                                        file_name=xlsx_name,
                                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                        key="__sims_flash_xlsx",
                                    )
                                except Exception:
                                    st.caption("⚠️ Excel 엔진(xlsxwriter/openpyxl)이 없어 CSV만 제공합니다.")
                    else:
                        st.info("SIMS 임시 결과 표 데이터가 세션에서 만료되었습니다.")

                    # 수동 닫기 버튼(항상 제공)
                    if st.button("닫기", key="__sims_flash_close"):
                        st.session_state.pop("__sims_flash", None)

# @@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@
    # (E) 🔎 채팅 검색 결과(옵션) — 사이드바 토글과 독립
    if st.session_state.get("__search_open"):
        q          = (st.session_state.get("__search_q") or "").strip()
        regex_mode = bool(st.session_state.get("__search_regex", False))
        case_sens  = bool(st.session_state.get("__search_case", False))
        roles      = set(st.session_state.get("__search_roles") or ["user", "assistant"])
        max_show   = int(st.session_state.get("__search_limit", 50))

        st.divider()
        st.subheader("🔎 검색 결과")

        if not isinstance(current_room, dict) or "messages" not in current_room:
            st.info("대화방이 아직 없습니다. 먼저 대화를 시작하세요.")
        elif not q:
            st.info("키워드를 입력하면 결과가 여기 표시됩니다.")
        else:
            try:
                pat = _mk_search_pattern(q, regex=regex_mode, case_sensitive=case_sens)
            except re.error as e:
                st.error(f"정규식 에러: {e}")
                pat = None

            if pat is not None:
                hits = search_messages_in_room(current_room, pat, roles)
                hits = list(reversed(hits))  # 최신 우선
                st.caption(f"총 {len(hits)}건 중 상위 {min(len(hits), max_show)}건 표시")

                if not hits:
                    st.info("일치하는 결과가 없습니다.")
                else:
                    for i, h in enumerate(hits[:max_show]):
                        st.markdown(f"**{h['role']}** · {h['time'] or ''}", help=f"index={h['idx']}")
                        st.markdown(h["snippet_html"], unsafe_allow_html=True)
                        with st.expander(f"전체 보기 · #{i+1}", expanded=False):
                            st.markdown(h["content"])

                        # 원문으로 이동
                        if st.button("원문으로 이동", key=f"__jump_main_{i}"):
                            msg = (current_room.get("messages") or [])[h["idx"]]
                            anchor_core = msg.get("id") or f"seq-{msg.get('seq', h['idx'])}"
                            st.session_state["__scroll_to_msg"] = f"jump-{anchor_core}"
                            st.rerun()

                    if len(hits) > max_show:
                        c1, c2 = st.columns(2)
                        with c1:
                            if st.button("더 보기 +50", key="__search_more_main"):
                                st.session_state["__search_limit"] = min(max_show + 50, 500)
                                st.rerun()
                        with c2:
                            if st.button("최대(500)", key="__search_more_max"):
                                st.session_state["__search_limit"] = 500
                                st.rerun()

    # (F) ▶ 메시지들이 모두 그려진 '이후'에 스크롤 실행 (1회)
    if _jump_to:
        _scroll_to_anchor_js(_jump_to, center=True)

    # =========================
    # 3) AI 답변 생성 (같은 런에서 표시, rerun 없음)
    # =========================
    if st.session_state.get("__queue_ai", False):
        st.session_state["__queue_ai"] = False   # pop 대신 안전하게 리셋
        log.debug("[chat] __queue_ai detected → AI block start")
        
        # ✅ 2-채널 히스토리: LLM에 넣을 히스토리 선택
        # - SIMS 패널 ON + SIMS 질문이면 sims_messages
        # - 그 외는 gen_messages
        history_msgs_all = current_room.get("messages", [])
        last_user_text = ""
        for mm in reversed(history_msgs_all):
            if mm.get("role") == "user":
                last_user_text = str(mm.get("content", "") or "")
                break
        last_user_is_sims = (
            is_sims_related_question(last_user_text)
            or is_sims_result_followup_question(last_user_text)
        )
        use_sims_hist = bool(st.session_state.get("__sims_open", False) and last_user_is_sims)

        history_msgs = current_room.get("sims_messages" if use_sims_hist else "gen_messages", []) or []

        try:
            msgs = build_messages_with_system(history_msgs, user_text=last_user_text)

            try:
                has_sims = any(
                    (m.get("role") == "system") and ("SIMS_JSON" in str(m.get("content", "")))
                    for m in msgs
                )
                log.debug("[chat] build_messages_with_system OK → msgs=%d, has_sims=%s",
                        len(msgs), has_sims)
            except Exception:
                pass
        except Exception:
            log.exception("[chat] build_messages_with_system failed, fallback to simple history")
            # fallback: system 없이 user/assistant 만 구성
            msgs = []
            for mm in history_msgs[-50:]:
                r = mm.get("role", "user")
                c = str(mm.get("content", ""))[:4000]
                if r in ("user", "assistant") and c.strip():
                    msgs.append({"role": r, "content": c})
            if not msgs:
                msgs = [{"role": "user", "content": "You are a helpful assistant."}]

        # ✅ 스트리밍 + 저장까지 내부에서 처리
        with pending_area:
            stream_and_append_assistant(
                messages_for_ai=msgs,
                room=current_room,
                temperature=0.2,
                history_channel=("sims_messages" if use_sims_hist else "gen_messages"),
            )
            log.debug("[chat] stream_and_append_assistant finished")
# =========================

st.write("")  # 시각적 여백

# =========================
# 4) 파일 분석 → 요약 → 즉시 AI 요청 (항상 자동)
# =========================
st.session_state.setdefault("__attach_keep_raw", False)

can_upload_file = require_permission("UPLOAD_FILE", show_error=False)

with st.expander("📎 첨부 처리 옵션", expanded=False):
    if not can_upload_file:
        st.info(_upload_unavailable_message())

    summary_target = st.number_input(
        "요약 목표 길이(문자)",
        min_value=300,
        max_value=4000,
        value=1200,
        step=100,
        help="긴 텍스트는 청크 요약 → 통합 요약으로 압축해 이 길이 내외로 정리합니다.",
        disabled=(not can_upload_file) or st.session_state.get("__an_busy", False),
    )

    st.checkbox(
        "원문 메시지도 함께 남기기",
        key="__attach_keep_raw",
        help=_upload_unavailable_help() if not can_upload_file else "체크하지 않으면 대화에는 요약만 남깁니다(원문은 파일 보관).",
        disabled=(not can_upload_file) or st.session_state.get("__an_busy", False),
    )



# ✔ 체크박스 값은 세션에서 읽어 사용
keep_raw = st.session_state.get("__attach_keep_raw", False)

# 파일 '분석하기'
if uploaded_files:
    if not can_upload_file:
        st.info(_upload_unavailable_message())
        st.button(
            "🔒 파일 분석",
            disabled=True,
            key="__btn_analyze_disabled",
            use_container_width=True,
            help=_upload_unavailable_help(),
        )
    else:
        st.info(f"{len(uploaded_files)}개 파일이 첨부되었습니다.")

    # 이미 분석 중이면 취소 버튼만 노출
    if st.session_state.get("__an_busy", False):
        st.warning("현재 파일 분석이 진행 중입니다…")
        st.button(
            "⏹️ 분석 취소",
            key="__an_cancel_btn",
            use_container_width=True,
            on_click=lambda: st.session_state.__setitem__("__an_cancel", True),
        )

    # 분석 시작 버튼
    elif st.button("🔍 파일 분석하기", use_container_width=True, key="__btn_analyze"):
        # ── 작업 시작 플래그/메타 ─────────────────────────────────
        st.session_state["__an_busy"] = True
        st.session_state["__an_cancel"] = False
        st.session_state["__an_job"] = {
            "id": str(uuid.uuid4()),
            "room_id": current_room.get("id"),
            "sig": _sig_of_uploads(uploaded_files),
            "started_at": time.time(),
        }

        attached_texts = []
        attached_saved_names = []
        progress = st.progress(0.0, text="파일 분석 시작…")
        total = len(uploaded_files)

        # ── 파일 루프 ───────────────────────────────────────────
        for idx, uf in enumerate(uploaded_files, start=1):
            # 사용자가 취소를 누르면 중단
            if st.session_state.get("__an_cancel"):
                st.info("⛔ 분석이 취소되었습니다.")
                break

            # 저장/해시
            try:
                file_hash = _sha256_of_filelike(uf)
            except Exception:
                file_hash = uuid.uuid4().hex
                
            save_path = _make_upload_save_path(uf.name, file_hash)

            if not save_path.exists():
                with save_path.open("wb") as f:
                    f.write(uf.getbuffer())

            attached_saved_names.append(save_path.name)
            cleanup_uploads(upload_dir=save_path.parent)
            
            # 텍스트 추출(미리보기 길이 보호)
            try:
                uf.seek(0)
            except Exception:
                pass
            extracted = process_file(uf, preview=True)
            attached_texts.append(f"### {uf.name}\n{extracted}")

            progress.progress(idx/total, text=f"분석 중... ({idx}/{total})")

        progress.empty()

        # ── 취소가 아니면 요약/전송 진행 ─────────────────────────
        if not st.session_state.get("__an_cancel"):
            combined_input = "\n\n---\n\n".join(attached_texts) or "(첨부 파일 분석 요청)"

            # 옵션: 원문 메시지도 남길지 여부
            if keep_raw:
                current_room.setdefault("messages", []).append({
                    "id": str(uuid.uuid4()),
                    "role": "user",
                    "content": combined_input,
                    "time": make_ts(),
                    "seq": _next_seq(),
                    "files": attached_saved_names,
                })

            # 요약 생성
            with st.spinner("첨부 요약 중..."):
                try:
                    summary = summarize_text_long(
                        combined_input,
                        target_chars=int(summary_target),
                    )
                except Exception as e:
                    summary = _clip_for_model(combined_input, limit=12000)
                    summary += f"\n\n(요약 실패로 원문 일부를 사용했습니다: {e})"

            # 파일 분석 결과는 추가 AI 호출 없이 assistant 메시지로 1회만 저장한다.
            # 이유:
            # - summarize_text_long() 단계에서 이미 LLM 요약이 수행됨
            # - 다시 __queue_ai=True로 일반 채팅 답변을 생성하면 답변이 2번 나온 것처럼 보임
            # - 첨부 문서가 시스템 프롬프트/역할 설명 문서일 경우 두 번째 답변이 자기소개처럼 나올 수 있음

            file_list_md = "\n".join([f"- {name}" for name in attached_saved_names]) or "- (저장 파일 없음)"

            request_msg = (
                f"📎 첨부 파일 분석 요청({len(attached_saved_names)}개)\n\n"
                f"{file_list_md}"
            )

            result_msg = (
                f"📎 첨부 파일 분석 결과({len(attached_saved_names)}개)\n\n"
                f"{summary}"
            )

            user_item = {
                "id": str(uuid.uuid4()),
                "role": "user",
                "content": request_msg,
                "time": make_ts(),
                "seq": _next_seq(),
                "files": attached_saved_names,
                **_message_meta("file_analysis_request"),
            }

            assistant_item = {
                "id": str(uuid.uuid4()),
                "role": "assistant",
                "content": result_msg,
                "time": make_ts(),
                "seq": _next_seq(),
                "files": attached_saved_names,
                **_message_meta("file_analysis_result"),
            }

            current_room.setdefault("messages", []).append(user_item)
            current_room.setdefault("messages", []).append(assistant_item)

            # 일반 대화 히스토리에도 저장해서 이후 사용자가 "방금 첨부파일 내용으로..."라고 물을 수 있게 한다.
            current_room.setdefault("gen_messages", []).append(user_item.copy())
            current_room.setdefault("gen_messages", []).append(assistant_item.copy())

            _sync_room_meta(current_room, materialize=True)
            save_chat_rooms()

            log.info(
                "[file.analysis] saved result %s files=%s summary_len=%s",
                _chat_log_kv(current_room),
                len(attached_saved_names),
                len(summary or ""),
            )

            try:
                user = get_current_user()
                company = get_selected_company() or {}

                safe_log_audit_event(
                    event_type="FILE_ANALYSIS",
                    action_result="SUCCESS",
                    actor_user_id=int(getattr(user, "user_id", 0) or 0) if user else None,
                    actor_login_id=str(getattr(user, "login_id", "") or "") if user else None,
                    company_id=int(company.get("company_id")) if company.get("company_id") else None,
                    target_company_id=int(company.get("company_id")) if company.get("company_id") else None,
                    message="파일 분석 실행",
                    details={
                        "room_id": current_room.get("id"),
                        "chat_file": str(_effective_chat_file()),
                        "file_count": len(attached_saved_names),
                        "files": attached_saved_names,
                        "summary_len": len(summary or ""),
                    },
                )
            except Exception as e:
                log.warning("[file.analysis.audit] failed reason=%s", e)

            # 추가 AI 답변 생성 금지
            st.session_state["__queue_ai"] = False
            st.session_state["__an_busy"] = False
            st.session_state["__an_job"] = None
            st.session_state["__an_cancel"] = False
            st.rerun()

        # ── 작업 종료/정리 ─────────────────────────────────────
        # 취소 경로 등 rerun을 호출하지 않은 경우에 대비한 안전망
        st.session_state["__an_busy"] = False
        st.session_state["__an_job"] = None
        st.session_state["__an_cancel"] = False

# 페이지 끝
